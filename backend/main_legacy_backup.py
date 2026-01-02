from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List, Tuple, Dict
import uvicorn
import os
import requests
import json
import io
import base64
from dotenv import load_dotenv
from datetime import datetime, timedelta
import uuid
from collections import defaultdict
import hashlib
import google.generativeai as genai

# Try to import optional libraries
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from PIL import Image
    import pytesseract
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_SUPPORT = True
except ImportError:
    PDF2IMAGE_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# Import PDF Summarizer (from TextSummarizer folder - using transformers)
try:
    from pdf_summarizer import PDFSummarizer
    PDF_SUMMARIZER_SUPPORT = True
except ImportError as e:
    PDF_SUMMARIZER_SUPPORT = False
    print(f"PDF Summarizer not available: {e}")
    print("Install required: pip install transformers torch sentencepiece")

# Import DOC Summarizer (from TextSummarizer folder - using transformers)
try:
    from doc_summarizer import DOCSummarizer
    DOC_SUMMARIZER_SUPPORT = True
except ImportError:
    DOC_SUMMARIZER_SUPPORT = False
    print("DOC Summarizer not available. Install: pip install transformers torch sentencepiece")

# Load environment variables
load_dotenv()

# Configuration
API_TITLE = os.getenv("API_TITLE", "Gurukul Backend API")
HOST = os.getenv("HOST", "0.0.0.0")
PORT = int(os.getenv("PORT", "3000"))
RELOAD = os.getenv("RELOAD", "True").lower() == "true"

# API Keys
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_API_ENDPOINT = os.getenv("GROQ_API_ENDPOINT", "https://api.groq.com/openai/v1/chat/completions")
GROQ_MODEL_NAME = os.getenv("GROQ_MODEL_NAME", "llama-3.3-70b-versatile")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL_PRIMARY = os.getenv("OLLAMA_MODEL_PRIMARY", "llama2")
LOCAL_LLAMA_API_URL = os.getenv("LOCAL_LLAMA_API_URL", "http://localhost:8080/v1/chat/completions")

YOUTUBE_API_KEY = os.getenv("YOUTUBE_API_KEY")
YOUTUBE_API_BASE_URL = "https://www.googleapis.com/youtube/v3/search"

# Initialize FastAPI
app = FastAPI(title=API_TITLE)

@app.on_event("startup")
async def startup_event():
    # Try to download models if they don't exist (for Render deployment)
    try:
        from download_models_on_startup import ensure_models_exist
        ensure_models_exist()
    except Exception as e:
        print(f"[Startup] Model download check failed: {e}")
    
    print("\n" + "="*50)
    print(f"API Started at http://{HOST}:{PORT}")
    print(f"PDF Support (PyPDF2): {PDF_SUPPORT}")
    print(f"PDF Summarizer Support (Local LED Transformer): {PDF_SUMMARIZER_SUPPORT}")
    print(f"DOC Summarizer Support: {DOC_SUMMARIZER_SUPPORT}")
    print(f"GEMINI_API_KEY Present: {'Yes' if GEMINI_API_KEY else 'No'}")
    print("="*50 + "\n")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Conversation Memory Storage (In-memory, can be upgraded to database)
# Structure: {conversation_id: {"messages": [...], "created_at": "...", "updated_at": "..."}}
conversation_store: Dict[str, Dict] = {}

# RAG Knowledge Store (Simple vector store for context retrieval)
# Structure: {hash: {"text": "...", "metadata": {...}}}
rag_knowledge_store: Dict[str, Dict] = {}

# Saved Summaries Store
# Structure: {summary_id: {"title": "...", "content": "...", "source": "...", "created_at": "...", "questions": [...]}}
saved_summaries_store: Dict[str, Dict] = {}

# Questions Store (linked to summaries)
# Structure: {question_id: {"summary_id": "...", "type": "...", "question": "...", "answer": "...", "options": [...], "review_schedule": {...}}}
questions_store: Dict[str, Dict] = {}

# Quiz Store
# Structure: {quiz_id: {"subject": "...", "topic": "...", "questions": [...], "created_at": "...", "provider": "..."}}
quiz_store: Dict[str, Dict] = {}

# Agent Simulator Stores
# Financial Profiles Store
# Structure: {profile_id: {"name": "...", "income": ..., "expenses": ..., "goal": "...", "created_at": "..."}}
financial_profiles_store: Dict[str, Dict] = {}

# Wellness Profiles Store
# Structure: {profile_id: {"emotional_score": ..., "financial_score": ..., "mood": ..., "stress": ..., "created_at": "..."}}
wellness_profiles_store: Dict[str, Dict] = {}

# Pydantic Models
class YouTubeVideo(BaseModel):
    title: str
    video_id: str
    url: str
    thumbnail: str
    channel_title: str
    duration: Optional[str] = None
    view_count: Optional[str] = None

class SubjectExplorerRequest(BaseModel):
    subject: str
    topic: str
    provider: Optional[str] = "groq"  # groq, ollama, llama
    use_knowledge_store: Optional[bool] = False

class SubjectExplorerResponse(BaseModel):
    subject: str
    topic: str
    notes: str
    provider: str
    youtube_recommendations: List[YouTubeVideo]
    success: bool




class SummarizerRequest(BaseModel):
    text: Optional[str] = None
    provider: Optional[str] = "groq"  # groq, ollama, llama, textrank
    summary_type: Optional[str] = "concise"  # concise, detailed, bullet_points

class SummarizerResponse(BaseModel):
    original_length: int
    summary: str
    summary_length: int
    provider: str
    summary_type: str
    success: bool

class PDFPageSummary(BaseModel):
    page_number: int
    summary: str
    original_length: int
    summary_length: int

class PDFSummarizerResponse(BaseModel):
    total_pages: int
    pages_summarized: int
    page_summaries: List[PDFPageSummary]
    overall_summary: str
    summary_type: str
    provider: str
    success: bool

class DOCSectionSummary(BaseModel):
    section_number: int
    summary: str
    original_length: int
    summary_length: int

class DOCSummarizerResponse(BaseModel):
    total_sections: int
    sections_summarized: int
    section_summaries: List[DOCSectionSummary]
    overall_summary: str
    summary_type: str
    provider: str
    success: bool

@app.post("/summarize-pdf", response_model=PDFSummarizerResponse)
async def summarize_pdf(
    file: UploadFile = File(...),
    summary_type: str = Form("detailed"),  # detailed, concise, comprehensive
    improve_grammar: bool = Form(False),
    save_summary: bool = Form(False),
    summary_title: Optional[str] = Form(None)
):
    """
    Summarize an uploaded PDF file using local transformer model.
    """
    print(f"\n[API] Received /summarize-pdf request for file: {file.filename}")
    print(f"[API] Type: {summary_type}, Save: {save_summary}")

    if not PDF_SUMMARIZER_SUPPORT:
        raise HTTPException(status_code=500, detail="PDF Summarizer not available on server.")
    
    # 1. Extract Text
    try:
        pages = await extract_pages_from_pdf(file)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {str(e)}")
        
    if not pages:
        raise HTTPException(status_code=400, detail="No text could be extracted from this PDF.")

    # 2. Initialize Summarizer
    try:
        summarizer = PDFSummarizer()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to initialize summarizer model: {str(e)}")

    # 3. Generate Summary
    try:
        result = summarizer.summarize_all_pages(pages, summary_type=summary_type, improve_grammar=improve_grammar)
        
        # Add success flag if not present
        if "success" not in result:
            result["success"] = True
            
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Summarization failed: {str(e)}")

# Chatbot Models
class ChatMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str
    timestamp: Optional[str] = None

class ChatRequest(BaseModel):
    message: str
    conversation_id: Optional[str] = None  # If None, creates new conversation
    provider: Optional[str] = "auto"  # auto, groq, ollama, llama
    use_rag: Optional[bool] = True  # Enable RAG for context retrieval
    max_history: Optional[int] = 10  # Maximum number of previous messages to include

class ChatResponse(BaseModel):
    response: str
    conversation_id: str
    provider: str
    message_count: int
    timestamp: str
    success: bool

class ConversationHistory(BaseModel):
    conversation_id: str
    messages: List[ChatMessage]
    created_at: str
    updated_at: str
    message_count: int

# Saved Summarizer Models
class SaveSummaryRequest(BaseModel):
    title: str
    content: str  # Summary content
    source: Optional[str] = None  # Original file name or source
    source_type: Optional[str] = None  # "pdf", "doc", "text", etc.
    metadata: Optional[Dict] = None

class SavedSummary(BaseModel):
    summary_id: str
    title: str
    content: str
    source: Optional[str]
    source_type: Optional[str]
    created_at: str
    updated_at: str
    question_count: int
    metadata: Optional[Dict]

class QuestionGenerationRequest(BaseModel):
    summary_id: str
    question_types: Optional[List[str]] = ["qa", "mcq", "flashcard", "case_based"]  # Types to generate
    num_questions: Optional[int] = 5  # Number of questions per type
    difficulty: Optional[str] = "medium"  # easy, medium, hard

class MCQOption(BaseModel):
    option: str
    is_correct: bool

class Question(BaseModel):
    question_id: str
    summary_id: str
    question_type: str  # "qa", "mcq", "flashcard", "case_based"
    question: str
    answer: str
    options: Optional[List[MCQOption]] = None  # For MCQs
    explanation: Optional[str] = None
    difficulty: str
    created_at: str

class QuestionGenerationResponse(BaseModel):
    summary_id: str
    questions: List[Question]
    total_questions: int
    success: bool

class ReviewSchedule(BaseModel):
    question_id: str
    summary_id: str
    first_review: str  # Same day
    second_review: str  # Next day
    third_review: str  # 1 week later
    last_review_date: Optional[str] = None
    next_review_date: str
    review_count: int = 0
    mastery_level: str = "learning"  # learning, reviewing, mastered

class ReviewRequest(BaseModel):
    question_id: str
    user_answer: Optional[str] = None
    is_correct: bool
    confidence: Optional[str] = "medium"  # low, medium, high

class ReviewResponse(BaseModel):
    question_id: str
    correct_answer: str
    is_correct: bool
    next_review_date: str
    review_count: int
    mastery_level: str
    message: str

class PendingReview(BaseModel):
    question_id: str
    summary_id: str
    question: str
    question_type: str
    answer: str
    next_review_date: str
    days_until_review: int

# Quiz/Test Models
class QuizRequest(BaseModel):
    subject: str
    topic: str
    provider: Optional[str] = "auto"  # auto, groq, ollama, llama
    difficulty: Optional[str] = "medium"  # easy, medium, hard

class QuizQuestion(BaseModel):
    question_id: str
    question_number: int
    question: str
    options: List[str]  # For MCQs
    question_type: str  # "mcq" or "qa"

class QuizResponse(BaseModel):
    quiz_id: str
    subject: str
    topic: str
    questions: List[QuizQuestion]
    total_questions: int
    created_at: str
    provider: str

class QuizSubmission(BaseModel):
    quiz_id: str
    answers: Dict[str, str]  # {question_id: user_answer}

class QuizResult(BaseModel):
    quiz_id: str
    total_questions: int
    correct_answers: int
    wrong_answers: int
    score_percentage: float
    results: List[Dict]  # Detailed results for each question
    submitted_at: str

# Agent Simulator Models
# EduMentor Models
class EduMentorRequest(BaseModel):
    subject: str
    topic: str
    include_wikipedia: Optional[bool] = False
    use_knowledge_store: Optional[bool] = False
    use_orchestration: Optional[bool] = False
    provider: Optional[str] = "auto"  # auto, groq, ollama, llama

class EduMentorResponse(BaseModel):
    subject: str
    topic: str
    lesson_content: str
    wikipedia_sources: Optional[List[Dict]] = None
    knowledge_store_used: bool
    orchestration_used: bool
    provider: str
    created_at: str
    success: bool

# Financial Agent Models
class ExpenseCategory(BaseModel):
    name: str
    amount: float

class FinancialProfileRequest(BaseModel):
    name: str
    monthly_income: float
    monthly_expenses: float
    expense_categories: List[ExpenseCategory]
    financial_goal: str
    financial_type: str  # Conservative, Moderate, Aggressive
    risk_level: str  # Low, Moderate, High

class FinancialAdviceResponse(BaseModel):
    name: str
    monthly_income: float
    monthly_expenses: float
    monthly_savings: float
    expense_breakdown: List[Dict]
    financial_advice: str
    recommendations: List[str]
    goal_analysis: Dict
    provider: str
    created_at: str

# Wellness Bot Models
class WellnessSupportRequest(BaseModel):
    emotional_wellness_score: int  # 1-10 (1=bad, 10=good)
    financial_wellness_score: int  # 1-10 (1=bad, 10=good)
    current_mood_score: int  # 1-10 (1=bad, 10=good)
    stress_level: int  # 1-10 (1=bad, 10=good)
    concerns: Optional[str] = None

class WellnessSupportResponse(BaseModel):
    emotional_support: str
    motivational_message: str
    life_importance: str
    study_importance: str
    goal_importance: str
    positive_affirmations: List[str]
    recommendations: List[str]
    overall_assessment: str
    provider: str
    created_at: str

# --- Flashcards & Summary Models ---
class SavedSummary(BaseModel):
    title: str
    content: str
    date: str

class Flashcard(BaseModel):
    question_id: str
    question: str
    answer: str
    question_type: str = "conceptual"
    days_until_review: int = 0
    confidence: float = 0.0

# In-memory storage
saved_summaries_db = []
flashcards_db = []

# --- Summary Endpoints ---
@app.post("/summaries/save")
async def save_summary_endpoint(summary: SavedSummary):
    """Save a summary to the database"""
    print(f"[Summary] Saving summary: {summary.title}")
    saved_summaries_db.append(summary.dict())
    return {"message": "Summary saved successfully", "count": len(saved_summaries_db)}

@app.get("/summaries")
async def get_summaries():
    """Get list of saved summaries"""
    return saved_summaries_db

# --- Flashcard Endpoints ---
class GenerateFlashcardsRequest(BaseModel):
    title: str
    content: str
    date: str
    question_type: str = "conceptual" # mixed, fill_in_blanks, short_answer, mcq

# ... (SavedSummary remains for save endpoint) ...

from fpdf import FPDF
import json

@app.post("/flashcards/generate")
async def generate_flashcards(request: GenerateFlashcardsRequest):
    """Generate flashcards from summary using LLM"""
    print(f"[Flashcards] Generating {request.question_type} from: {request.title}")
    
    # Construct prompt for Groq/Ollama
    type_instruction = ""
    if request.question_type == "fill_in_blanks":
        type_instruction = "Create Fill-in-the-Blank style questions. Output Format: {'question': 'The capital of France is ______.', 'answer': 'Paris', 'type': 'fill_in_blanks'}"
    elif request.question_type == "short_answer":
        type_instruction = "Create Short Answer (One Line) questions."
    elif request.question_type == "mcq":
        type_instruction = "Create Multiple Choice Questions (MCQs). Format the question to include options like 'A) ... B) ...' and provide the correct option in answer."
    else:
        type_instruction = "Create Conceptual/Mixed questions."

    prompt = f"""
    Create 5 flashcards based on this summary:
    "{request.content[:3000]}..."
    
    INSTRUCTION: {type_instruction}
    
    Format the output as a PURE JSON list (no markdown, no backticks).
    Example:
    [
        {{"question": "Question 1 text...", "answer": "Answer 1...", "type": "{request.question_type}"}},
        ...
    ]
    """
    
    new_cards = []
    generated_text = ""
    
    try:
        # 1. Try Groq
        if GROQ_API_KEY:
            try:
                print("[Flashcards] Using Groq...")
                headers = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
                payload = {
                    "model": GROQ_MODEL_NAME, 
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.5
                }
                resp = requests.post(GROQ_API_ENDPOINT, headers=headers, json=payload, timeout=20)
                if resp.status_code == 200:
                    generated_text = resp.json()["choices"][0]["message"]["content"]
            except Exception as e:
                print(f"[Flashcards] Groq failed: {e}")

        # 2. Try Gemini if Groq failed
        if not generated_text and GEMINI_API_KEY:
            try:
                print("[Flashcards] Using Gemini...")
                genai.configure(api_key=GEMINI_API_KEY)
                model = genai.GenerativeModel('gemini-pro')
                response = model.generate_content(prompt)
                generated_text = response.text
            except Exception as e:
                print(f"[Flashcards] Gemini failed: {e}")

        # 3. Fallback to Ollama if Groq/Gemini failed
            try:
                resp = requests.post(f"{OLLAMA_BASE_URL}/api/generate", json={
                    "model": OLLAMA_MODEL_PRIMARY,
                    "prompt": prompt,
                    "stream": False
                }, timeout=30)
                if resp.status_code == 200:
                    generated_text = resp.json().get("response", "")
            except Exception as e:
                print(f"[Flashcards] Ollama failed: {e}")

        # 4. Parse JSON
        if generated_text:
            # Clean markdown if present
            clean_text = generated_text.replace("```json", "").replace("```", "").strip()
            try:
                data = json.loads(clean_text)
                import uuid
                for item in data:
                    new_cards.append({
                        "question_id": str(uuid.uuid4()),
                        "question": item.get("question", "Unknown"),
                        "answer": item.get("answer", "Unknown"),
                        "question_type": item.get("type", request.question_type),
                        "days_until_review": 0
                    })
            except json.JSONDecodeError:
                print(f"[Flashcards] JSON Decode Failed. Raw: {clean_text}")
                # Fallback to mock if parsing fails
                
        # 5. If AI failed completely, use Mock
        if not new_cards:
            print("[Flashcards] AI generation failed or returned invalid JSON. Using Mock Fallback.")
            import uuid
            for i in range(1, 6):
                new_cards.append({
                    "question_id": str(uuid.uuid4()),
                    "question": f"Mock Question {i} ({request.question_type})",
                    "answer": "Mock Answer",
                    "question_type": request.question_type,
                    "days_until_review": 0
                })
        
        flashcards_db.extend(new_cards)
        return {"message": f"Generated {len(new_cards)} cards", "cards": new_cards}
        
    except Exception as e:
        print(f"[Flashcards] Critical Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/flashcards/download_pdf")
async def download_flashcards_pdf():
    """Generate and return PDF of all flashcards"""
    if not flashcards_db:
        raise HTTPException(status_code=404, detail="No flashcards found")
        
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        pdf.cell(200, 10, txt="Flashcard Questions", ln=1, align='C')
        pdf.ln(10)
        
        for i, card in enumerate(flashcards_db, 1):
            # Safe unicode handling
            q = card['question'].encode('latin-1', 'replace').decode('latin-1')
            a = card['answer'].encode('latin-1', 'replace').decode('latin-1')
            
            pdf.set_font("Arial", 'B', 12)
            pdf.multi_cell(0, 10, f"Q{i}: {q}")
            pdf.set_font("Arial", '', 11)
            pdf.multi_cell(0, 10, f"A:  {a}")
            pdf.ln(5)
            
        filename = "flashcards.pdf"
        # In a real app, save to temp file. Here, return bytes directly.
        # FPDF output() returns string in Py2, bytes in Py3 depending on args.
        # output(dest='S') returns string/bytes.
        # But fastapi Response needs bytes.
        
        # We will save to a temp file then stream it
        import tempfile
        from fastapi.responses import FileResponse
        
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        pdf.output(tmp.name)
        tmp.close()
        
        return FileResponse(tmp.name, filename="practice_questions.pdf", media_type='application/pdf')
        
    except Exception as e:
        print(f"PDF Error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate PDF")

@app.get("/reviews/pending")
async def get_pending_reviews():
    """Get pending flashcards for frontend"""
    return flashcards_db

@app.get("/reviews/stats")
async def get_review_stats():
    """Get stats"""
    return {
        "total_questions": len(flashcards_db),
        "pending_reviews": len(flashcards_db),
        "learning": 0,
        "reviewing": 0,
        "mastered": 0
    }

# Helper function to create teacher-like prompt
def create_teaching_prompt(subject: str, topic: str) -> str:
    return f"""You are an expert teacher specializing in {subject}. Your task is to teach the topic "{topic}" to a student.

Please provide comprehensive, clear, and engaging educational notes that include:

1. **Introduction**: A brief overview of the topic and why it's important
2. **Key Concepts**: Break down the main concepts in simple, understandable terms
3. **Detailed Explanation**: Provide a thorough explanation with examples
4. **Real-world Applications**: Show how this topic applies in real life
5. **Summary**: A concise recap of the key points
6. **Practice Questions**: 2-3 questions to help reinforce learning

Write in a friendly, encouraging tone as if you're speaking directly to a student. Use simple language, analogies, and examples to make complex concepts easy to understand. Structure the content with clear headings and bullet points for better readability.

Subject: {subject}
Topic: {topic}

Now, provide the educational notes:"""

# Groq API call
async def call_groq_api(subject: str, topic: str) -> str:
    if not GROQ_API_KEY:
        raise HTTPException(status_code=500, detail="Groq API key not configured")
    
    prompt = create_teaching_prompt(subject, topic)
    
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": GROQ_MODEL_NAME,
        "messages": [
            {
                "role": "system",
                "content": "You are an expert teacher who explains concepts clearly and simply."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.7,
        "max_tokens": 2000
    }
    
    try:
        response = requests.post(GROQ_API_ENDPOINT, headers=headers, json=payload, timeout=30)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Groq API error: {str(e)}")

# Ollama API call
async def call_ollama_api(subject: str, topic: str) -> str:
    prompt = create_teaching_prompt(subject, topic)
    
    url = f"{OLLAMA_BASE_URL}/api/generate"
    
    payload = {
        "model": OLLAMA_MODEL_PRIMARY,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.7,
            "num_predict": 2000
        }
    }
    
    try:
        response = requests.post(url, json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
        return data.get("response", "Error generating response from Ollama")
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Ollama service not available. Make sure Ollama is running on localhost:11434")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ollama API error: {str(e)}")

# Local LLaMA API call
async def call_llama_api(subject: str, topic: str) -> str:
    prompt = create_teaching_prompt(subject, topic)
    
    headers = {
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "llama3.1",
        "messages": [
            {
                "role": "system",
                "content": "You are an expert teacher who explains concepts clearly and simply."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.7,
        "max_tokens": 2000
    }
    
    try:
        response = requests.post(LOCAL_LLAMA_API_URL, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Local LLaMA service not available. Make sure the local LLaMA server is running")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLaMA API error: {str(e)}")

# YouTube API call to get recommendations
async def get_youtube_recommendations(subject: str, topic: str, max_results: int = 5) -> List[YouTubeVideo]:
    """Get YouTube video recommendations based on subject and topic"""
    if not YOUTUBE_API_KEY:
        return []  # Return empty list if API key is not configured
    
    # Create search query
    search_query = f"{subject} {topic} tutorial explanation"
    
    params = {
        "part": "snippet",
        "q": search_query,
        "type": "video",
        "maxResults": max_results,
        "key": YOUTUBE_API_KEY,
        "order": "relevance",
        "videoCategoryId": "27",  # Education category
        "safeSearch": "strict"
    }
    
    try:
        response = requests.get(YOUTUBE_API_BASE_URL, params=params, timeout=10)
        response.raise_for_status()
        data = response.json()
        
        videos = []
        for item in data.get("items", []):
            video_id = item["id"]["videoId"]
            snippet = item["snippet"]
            
            video = YouTubeVideo(
                title=snippet["title"],
                video_id=video_id,
                url=f"https://www.youtube.com/watch?v={video_id}",
                thumbnail=snippet["thumbnails"]["high"]["url"],
                channel_title=snippet["channelTitle"]
            )
            videos.append(video)
        
        return videos
    except requests.exceptions.RequestException as e:
        # If YouTube API fails, return empty list (don't break the main functionality)
        print(f"YouTube API error: {str(e)}")
        return []
    except Exception as e:
        print(f"Error fetching YouTube recommendations: {str(e)}")
        return []

# Text extraction functions
async def extract_pages_from_pdf(file: UploadFile) -> List[str]:
    """
    Extract text from PDF file with multi-page support.
    Returns list of page texts (one per page).
    Tries text extraction first, falls back to OCR for scanned PDFs.
    """
    if not PDF_SUPPORT:
        raise HTTPException(status_code=400, detail="PDF support not available. Install PyPDF2: pip install PyPDF2")
    
    try:
        content = await file.read()
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        print(f"[PDF] Processing PDF with {len(pdf_reader.pages)} pages...")
        
        # Try text extraction first
        text_pages = []
        for page_num, page in enumerate(pdf_reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text and len(page_text.strip()) > 10:
                    text_pages.append(page_text)
                    print(f"[PDF] Page {page_num}: Extracted {len(page_text)} characters")
                else:
                    print(f"[PDF] Page {page_num}: No text found, will use OCR")
                    text_pages.append(None)  # Mark for OCR
            except Exception as e:
                print(f"[PDF] Page {page_num}: Text extraction failed: {str(e)}, will use OCR")
                text_pages.append(None)
        
        # Check if we need OCR (scanned PDFs)
        needs_ocr = any(page is None for page in text_pages)
        
        if needs_ocr and OCR_SUPPORT and PDF2IMAGE_SUPPORT:
            print("[PDF] Using OCR for scanned PDF pages...")
            # Convert PDF pages to images and use OCR
            images = convert_from_bytes(content, dpi=300)
            
            ocr_text_pages = []
            for page_num, image in enumerate(images, 1):
                try:
                    page_text = pytesseract.image_to_string(image)
                    if page_text and len(page_text.strip()) > 10:
                        ocr_text_pages.append(page_text)
                        print(f"[PDF] Page {page_num} (OCR): Extracted {len(page_text)} characters")
                    else:
                        ocr_text_pages.append("")
                except Exception as e:
                    print(f"[PDF] Page {page_num} (OCR): Failed - {str(e)}")
                    ocr_text_pages.append("")
            
            # Combine OCR results with text extraction results
            final_pages = []
            for i, (text_page, ocr_page) in enumerate(zip(text_pages, ocr_text_pages), 1):
                if text_page:
                    final_pages.append(text_page)
                elif ocr_page:
                    final_pages.append(ocr_page)
                else:
                    final_pages.append("")  # Empty page
            
            print(f"[PDF] Total extracted: {len(final_pages)} pages")
            return final_pages
        
        elif needs_ocr:
            raise HTTPException(
                status_code=400, 
                detail="PDF appears to be scanned (image-based). OCR support required. Install: pip install pdf2image pytesseract Pillow"
            )
        else:
            # All pages have text, return as list
            print(f"[PDF] Total extracted: {len(text_pages)} pages")
            return text_pages
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from PDF: {str(e)}")

async def extract_text_from_pdf(file: UploadFile) -> str:
    """
    Extract text from PDF file with multi-page support.
    Tries text extraction first, falls back to OCR for scanned PDFs.
    """
    if not PDF_SUPPORT:
        raise HTTPException(status_code=400, detail="PDF support not available. Install PyPDF2: pip install PyPDF2")
    
    try:
        content = await file.read()
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        print(f"[PDF] Processing PDF with {len(pdf_reader.pages)} pages...")
        
        # Try text extraction first
        text_pages = []
        for page_num, page in enumerate(pdf_reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text and len(page_text.strip()) > 10:
                    text_pages.append(page_text)
                    print(f"[PDF] Page {page_num}: Extracted {len(page_text)} characters")
                else:
                    print(f"[PDF] Page {page_num}: No text found, will use OCR")
                    text_pages.append(None)  # Mark for OCR
            except Exception as e:
                print(f"[PDF] Page {page_num}: Text extraction failed: {str(e)}, will use OCR")
                text_pages.append(None)
        
        # Check if we need OCR (scanned PDFs)
        needs_ocr = any(page is None for page in text_pages)
        
        if needs_ocr and OCR_SUPPORT and PDF2IMAGE_SUPPORT:
            print("[PDF] Using OCR for scanned PDF pages...")
            # Convert PDF pages to images and use OCR
            images = convert_from_bytes(content, dpi=300)
            
            ocr_text_pages = []
            for page_num, image in enumerate(images, 1):
                try:
                    page_text = pytesseract.image_to_string(image)
                    if page_text and len(page_text.strip()) > 10:
                        ocr_text_pages.append(page_text)
                        print(f"[PDF] Page {page_num} (OCR): Extracted {len(page_text)} characters")
                    else:
                        ocr_text_pages.append("")
                except Exception as e:
                    print(f"[PDF] Page {page_num} (OCR): Failed - {str(e)}")
                    ocr_text_pages.append("")
            
            # Combine OCR results with text extraction results
            final_text = ""
            for i, (text_page, ocr_page) in enumerate(zip(text_pages, ocr_text_pages), 1):
                if text_page:
                    final_text += f"\n[Page {i}]\n{text_page}\n"
                elif ocr_page:
                    final_text += f"\n[Page {i} - OCR]\n{ocr_page}\n"
            
            result = final_text.strip()
            print(f"[PDF] Total extracted text: {len(result)} characters from {len(images)} pages")
            return result
        
        elif needs_ocr:
            raise HTTPException(
                status_code=400, 
                detail="PDF appears to be scanned (image-based). OCR support required. Install: pip install pdf2image pytesseract Pillow"
            )
        else:
            # All pages have text, combine them
            final_text = "\n\n".join([f"[Page {i+1}]\n{page_text}" for i, page_text in enumerate(text_pages) if page_text])
            print(f"[PDF] Total extracted text: {len(final_text)} characters from {len(text_pages)} pages")
            return final_text.strip()
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from PDF: {str(e)}")

async def extract_text_from_image(file: UploadFile) -> str:
    """Extract text from image using OCR"""
    if not OCR_SUPPORT:
        raise HTTPException(status_code=400, detail="OCR support not available. Install pytesseract and Pillow")
    
    try:
        content = await file.read()
        image = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from image: {str(e)}")

async def extract_sections_from_docx(file: UploadFile) -> List[str]:
    """Extract text from DOCX file with multi-section support. Returns list of sections."""
    if not DOCX_SUPPORT:
        raise HTTPException(status_code=400, detail="DOCX support not available. Install: pip install python-docx")
    
    try:
        content = await file.read()
        doc_file = io.BytesIO(content)
        doc = Document(doc_file)
        
        print(f"[DOCX] Processing DOCX with {len(doc.paragraphs)} paragraphs...")
        
        # Extract text from all paragraphs as separate sections
        sections = []
        current_section = []
        
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                # Check if this looks like a heading (bold, larger font, etc.)
                is_heading = False
                if para.runs:
                    for run in para.runs:
                        if run.bold or (run.font and run.font.size and run.font.size.pt > 12):
                            is_heading = True
                            break
                
                if is_heading and current_section:
                    # Start new section
                    sections.append("\n".join(current_section))
                    current_section = [para_text]
                else:
                    current_section.append(para_text)
        
        # Add last section
        if current_section:
            sections.append("\n".join(current_section))
        
        # Also extract text from tables as separate sections
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    table_text.append(row_text)
            if table_text:
                sections.append("\n".join(table_text))
        
        # If no sections found, combine all as one
        if not sections:
            all_text = "\n\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
            if all_text:
                sections = [all_text]
        
        print(f"[DOCX] Extracted {len(sections)} sections")
        return sections
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from DOCX: {str(e)}")

async def extract_text_from_docx(file: UploadFile) -> str:
    """Extract text from DOCX file with multi-page support (legacy function)"""
    sections = await extract_sections_from_docx(file)
    return "\n\n".join(sections)

async def improve_grammar_and_clarity(text: str, provider: str = "groq") -> str:
    """
    Improve grammar, fix errors, and enhance clarity of text using AI models.
    Uses Groq by default, with fallback to other models.
    """
    if not text or len(text.strip()) < 10:
        return text
    
    prompt = f"""Please improve the following text by:
1. Fixing all grammar errors
2. Correcting spelling mistakes
3. Improving sentence structure and clarity
4. Ensuring proper punctuation
5. Making sentences more meaningful and coherent
6. Maintaining the original meaning and context

Text to improve:
{text}

Improved text:"""
    
    # Try Groq first
    if provider == "groq" and GROQ_API_KEY:
        try:
            headers = {
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": GROQ_MODEL_NAME,
                "messages": [
                    {
                        "role": "system",
                        "content": "You are an expert editor who improves text quality, fixes grammar errors, and enhances clarity while maintaining the original meaning."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "temperature": 0.3,
                "max_tokens": 2000
            }
            
            response = requests.post(GROQ_API_ENDPOINT, headers=headers, json=payload, timeout=60)
            response.raise_for_status()
            data = response.json()
            improved_text = data["choices"][0]["message"]["content"]
            print(f"[Grammar] Text improved using Groq. Original: {len(text)} chars, Improved: {len(improved_text)} chars")
            return improved_text
        except Exception as e:
            print(f"[Grammar] Groq failed: {str(e)}, trying fallback...")
    
    # Try Ollama as fallback
    try:
        url = f"{OLLAMA_BASE_URL}/api/generate"
        payload = {
            "model": OLLAMA_MODEL_PRIMARY,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": 0.3,
                "num_predict": 2000
            }
        }
        response = requests.post(url, json=payload, timeout=120)
        response.raise_for_status()
        data = response.json()
        improved_text = data.get("response", text)
        print(f"[Grammar] Text improved using Ollama. Original: {len(text)} chars, Improved: {len(improved_text)} chars")
        return improved_text
    except Exception as e:
        print(f"[Grammar] Ollama failed: {str(e)}, trying LLaMA...")
    
    # Try Local LLaMA as last fallback
    try:
        headers = {"Content-Type": "application/json"}
        payload = {
            "model": "llama3.1",
            "messages": [
                {
                    "role": "system",
                    "content": "You are an expert editor who improves text quality, fixes grammar errors, and enhances clarity while maintaining the original meaning."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.3,
            "max_tokens": 2000
        }
        response = requests.post(LOCAL_LLAMA_API_URL, headers=headers, json=payload, timeout=120)
        response.raise_for_status()
        data = response.json()
        improved_text = data["choices"][0]["message"]["content"]
        print(f"[Grammar] Text improved using LLaMA. Original: {len(text)} chars, Improved: {len(improved_text)} chars")
        return improved_text
    except Exception as e:
        print(f"[Grammar] All grammar improvement methods failed: {str(e)}, returning original text")
        return text

def create_summarization_prompt(text: str, summary_type: str) -> str:
    """Create prompt for summarization based on type"""
    base_prompt = f"""Please summarize the following text. The text is:"""
    
    if summary_type == "concise":
        instruction = "Provide a brief, concise summary (2-3 sentences) capturing the main points."
    elif summary_type == "detailed":
        instruction = "Provide a comprehensive, detailed summary covering all important aspects and key points."
    elif summary_type == "bullet_points":
        instruction = "Provide a summary in bullet point format, highlighting the main ideas and key information."
    else:
        instruction = "Provide a clear and comprehensive summary."
    
    return f"""{base_prompt}

{instruction}

Text to summarize:
{text}

Summary:"""

# Summarization functions using different providers
async def summarize_with_groq(text: str, summary_type: str) -> str:
    """Summarize text using Groq API"""
    if not GROQ_API_KEY:
        raise HTTPException(status_code=500, detail="Groq API key not configured")
    
    prompt = create_summarization_prompt(text, summary_type)
    
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": GROQ_MODEL_NAME,
        "messages": [
            {
                "role": "system",
                "content": "You are an expert at summarizing documents and extracting key information."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.3,
        "max_tokens": 1000
    }
    
    try:
        response = requests.post(GROQ_API_ENDPOINT, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Groq API error: {str(e)}")

async def summarize_with_ollama(text: str, summary_type: str) -> str:
    """Summarize text using Ollama API"""
    prompt = create_summarization_prompt(text, summary_type)
    
    url = f"{OLLAMA_BASE_URL}/api/generate"
    
    payload = {
        "model": OLLAMA_MODEL_PRIMARY,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.3,
            "num_predict": 1000
        }
    }
    
    try:
        response = requests.post(url, json=payload, timeout=120)
        response.raise_for_status()
        data = response.json()
        return data.get("response", "Error generating summary from Ollama")
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Ollama service not available. Make sure Ollama is running on localhost:11434")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ollama API error: {str(e)}")

async def summarize_with_llama(text: str, summary_type: str) -> str:
    """Summarize text using Local LLaMA API"""
    prompt = create_summarization_prompt(text, summary_type)
    
    headers = {
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "llama3.1",
        "messages": [
            {
                "role": "system",
                "content": "You are an expert at summarizing documents and extracting key information."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.3,
        "max_tokens": 1000
    }
    
    try:
        response = requests.post(LOCAL_LLAMA_API_URL, headers=headers, json=payload, timeout=120)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Local LLaMA service not available. Make sure the local LLaMA server is running")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLaMA API error: {str(e)}")

async def summarize_with_bart(text: str, summary_type: str) -> str:
    """Summarize text using BART model (from TextSummarizer folder - transformers)"""
    if not PDF_SUMMARIZER_SUPPORT:
        raise HTTPException(status_code=500, detail="BART summarizer not available. Install: pip install transformers torch sentencepiece")
    
    try:
        print(f"[BART] Starting summarization. Text length: {len(text)}, Type: {summary_type}")
        from pdf_summarizer import PDFSummarizer
        summarizer = PDFSummarizer()
        
        # Check if text contains multiple pages (marked with [Page X])
        if "[Page" in text or "\n[Page" in text:
            print("[BART] Detected multi-page document, using multi-page summarization...")
            # Split by page markers
            pages = []
            current_page = ""
            for line in text.split('\n'):
                if line.strip().startswith('[Page'):
                    if current_page:
                        pages.append(current_page.strip())
                    current_page = ""
                else:
                    current_page += line + "\n"
            if current_page:
                pages.append(current_page.strip())
            
            print(f"[BART] Processing {len(pages)} pages...")
            
            # Determine summary length based on type
            if summary_type == "concise":
                max_length = 80
                min_length = 30
            elif summary_type == "detailed":
                max_length = 250
                min_length = 100
            else:
                max_length = 150
                min_length = 40
            
            # Use multi-page summarization
            summary = summarizer.multi_page_summarize(pages, max_length=max_length, min_length=min_length)
        else:
            # Single document summarization
            if summary_type == "concise":
                max_length = 80
                min_length = 30
            elif summary_type == "detailed":
                max_length = 250
                min_length = 100
            else:
                max_length = 150
                min_length = 40
            
            print(f"[BART] Using max_length: {max_length}, min_length: {min_length}")
            summary = summarizer.summarize(text, max_length=max_length, min_length=min_length)
        
        print(f"[BART] Summary generated. Length: {len(summary)}")
        
        # Format based on summary type
        if summary_type == "bullet_points":
            sentences = summary.split('. ')
            summary = '\n• ' + '\n• '.join([s.strip() for s in sentences if s.strip()])
        
        if not summary or len(summary.strip()) == 0:
            raise ValueError("BART returned empty summary")
        
        print(f"[BART] Success! Final summary length: {len(summary)}")
        return summary
    except Exception as e:
        print(f"[BART] ERROR: {str(e)}")
        raise HTTPException(status_code=500, detail=f"BART summarization error: {str(e)}")

async def summarize_with_fallback(text: str, summary_type: str, preferred_provider: str = None) -> Tuple[str, str]:
    """
    Summarize text with fallback logic:
    1. Try BART first (from TextSummarizer folder - transformers) - ALWAYS TRY FIRST
    2. If BART fails or not available, try Groq
    3. If Groq fails, try Ollama
    4. If Ollama fails, try LLaMA
    
    Returns: (summary, provider_used)
    """
    # Priority 1: ALWAYS Try BART first (from TextSummarizer folder - transformers)
    if PDF_SUMMARIZER_SUPPORT:
        try:
            print("Attempting BART summarization (first priority - from TextSummarizer)...")
            summary = await summarize_with_bart(text, summary_type)
            if summary and len(summary.strip()) > 0:
                print("BART summarization successful!")
                return summary, "bart"
            else:
                print("BART returned empty summary, trying fallback...")
        except Exception as e:
            print(f"BART failed: {str(e)}, trying fallback...")
    else:
        print("BART not available, trying fallback...")
    
    # Priority 2: Try Groq (only if TextRank failed)
    if GROQ_API_KEY:
        try:
            print("Attempting Groq summarization (fallback)...")
            summary = await summarize_with_groq(text, summary_type)
            return summary, "groq"
        except Exception as e:
            print(f"Groq failed, trying fallback: {str(e)}")
    
    # Priority 3: Try Ollama
    try:
        print("Attempting Ollama summarization (fallback)...")
        summary = await summarize_with_ollama(text, summary_type)
        return summary, "ollama"
    except Exception as e:
        print(f"Ollama failed, trying fallback: {str(e)}")
    
    # Priority 4: Try LLaMA
    try:
        print("Attempting LLaMA summarization (fallback)...")
        summary = await summarize_with_llama(text, summary_type)
        return summary, "llama"
    except Exception as e:
        raise HTTPException(
            status_code=500, 
            detail=f"All summarization methods failed. BART (first priority) and all fallbacks failed. Last error: {str(e)}"
        )

# Root endpoint
@app.get("/")
async def root():
    return {
        "message": "Welcome to Gurukul Backend API",
        "status": "running"
    }

# Health check
@app.get("/health")
async def health_check():
    return {"status": "healthy"}

# Test BART endpoint
@app.get("/test-bart")
async def test_bart():
    """Test endpoint to verify BART summarizer is working"""
    test_text = """
    Artificial intelligence is transforming the world. Machine learning algorithms can learn from data.
    Deep learning uses neural networks with multiple layers. Natural language processing helps computers understand text.
    Computer vision enables machines to see and interpret images. AI is being used in healthcare, finance, and education.
    The future of AI looks promising with continued research and development.
    """
    
    try:
        if not PDF_SUMMARIZER_SUPPORT:
            return {
                "status": "error",
                "message": "BART not available. Install: pip install transformers torch sentencepiece",
                "bart_support": False
            }
        
        from pdf_summarizer import PDFSummarizer
        summarizer = PDFSummarizer()
        summary = summarizer.summarize(test_text, max_length=100, min_length=30)
        
        return {
            "status": "success",
            "message": "BART is working!",
            "bart_support": True,
            "original_length": len(test_text),
            "summary": summary,
            "summary_length": len(summary),
            "test_passed": len(summary) > 0 and len(summary) < len(test_text)
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"BART test failed: {str(e)}",
            "bart_support": PDF_SUMMARIZER_SUPPORT,
            "error": str(e)
        }

# Subject Explorer endpoint
@app.post("/subject-explorer", response_model=SubjectExplorerResponse)
async def subject_explorer(request: SubjectExplorerRequest):
    """
    Generate educational notes for a subject and topic using AI models.
    Supports Groq, Ollama, and Local LLaMA providers.
    Also provides YouTube video recommendations related to the subject and topic.
    """
    if not request.subject or not request.topic:
        raise HTTPException(status_code=400, detail="Subject and Topic are required")
    
    provider = request.provider.lower()
    
    try:
        # Generate notes
        if provider == "groq":
            notes = await call_groq_api(request.subject, request.topic)
        elif provider == "ollama":
            notes = await call_ollama_api(request.subject, request.topic)
        elif provider == "llama":
            notes = await call_llama_api(request.subject, request.topic)
        else:
            raise HTTPException(status_code=400, detail="Invalid provider. Use 'groq', 'ollama', or 'llama'")
        
        # Get YouTube recommendations
        youtube_videos = await get_youtube_recommendations(request.subject, request.topic)
        
        return SubjectExplorerResponse(
            subject=request.subject,
            topic=request.topic,
            notes=notes,
            provider=provider,
            youtube_recommendations=youtube_videos,
            success=True
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating notes: {str(e)}")

# Summarizer endpoint - Text input
@app.post("/summarize", response_model=SummarizerResponse)
async def summarize_text(request: SummarizerRequest):
    """
    Summarize text input using AI models with fallback logic.
    Priority: BART (from TextSummarizer folder - transformers) -> Groq -> Ollama -> LLaMA
    """
    if not request.text:
        raise HTTPException(status_code=400, detail="Text is required")
    
    if len(request.text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Text must be at least 50 characters long")
    
    provider = request.provider.lower() if request.provider else None
    summary_type = request.summary_type.lower()
    
    try:
        # If specific provider is requested, use it directly
        if provider == "groq":
            summary = await summarize_with_groq(request.text, summary_type)
            used_provider = "groq"
        elif provider == "ollama":
            summary = await summarize_with_ollama(request.text, summary_type)
            used_provider = "ollama"
        elif provider == "llama":
            summary = await summarize_with_llama(request.text, summary_type)
            used_provider = "llama"
        elif provider == "bart":
            summary = await summarize_with_bart(request.text, summary_type)
            used_provider = "bart"
        elif provider is None or provider == "auto":
            # Use fallback logic: BART first, then others
            summary, used_provider = await summarize_with_fallback(request.text, summary_type)
        else:
            raise HTTPException(status_code=400, detail="Invalid provider. Use 'auto', 'bart', 'groq', 'ollama', or 'llama'")
        
        return SummarizerResponse(
            original_length=len(request.text),
            summary=summary,
            summary_length=len(summary),
            provider=used_provider,
            summary_type=summary_type,
            success=True
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating summary: {str(e)}")

# Summarizer endpoint - DOC/DOCX file upload
@app.post("/summarize-doc", response_model=DOCSummarizerResponse)
async def summarize_doc(
    file: UploadFile = File(...),
    summary_type: str = Form("comprehensive"),
    save_summary: bool = Form(False),
    summary_title: Optional[str] = Form(None)
):
    """
    Comprehensive DOC/DOCX summarizer that processes ALL sections individually and creates an overall summary.
    This endpoint is specifically designed for multi-section documents and ensures all points are covered.
    
    Features:
    - Processes each section individually
    - Creates section-by-section summaries
    - Generates comprehensive overall summary covering all sections
    - Uses BART model for contextual, abstractive summarization
    
    Summary types:
    - "concise": Brief summaries (2-3 sentences per section)
    - "detailed": Comprehensive summaries with all key points
    - "comprehensive": Most detailed summaries covering all aspects (default)
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="File is required")
    
    filename_lower = file.filename.lower()
    if not (filename_lower.endswith('.docx') or filename_lower.endswith('.doc')):
        raise HTTPException(status_code=400, detail="Only DOC/DOCX files are supported for this endpoint")
    
    if filename_lower.endswith('.doc'):
        raise HTTPException(status_code=400, detail="Old .doc format not supported. Please convert to .docx")
    
    if not DOCX_SUPPORT:
        raise HTTPException(status_code=400, detail="DOCX support not available. Install: pip install python-docx")
    
    if not DOC_SUMMARIZER_SUPPORT:
        raise HTTPException(
            status_code=400, 
            detail="DOC Summarizer not available. Install: pip install transformers torch sentencepiece"
        )
    
    summary_type = summary_type.lower()
    if summary_type not in ["concise", "detailed", "comprehensive"]:
        summary_type = "comprehensive"
    
    try:
        # Extract sections from DOCX
        sections = await extract_sections_from_docx(file)
        
        if not sections or all(not s or len(s.strip()) < 50 for s in sections):
            raise HTTPException(
                status_code=400, 
                detail="Could not extract sufficient text from DOCX. File may be empty or corrupted."
            )
        
        print(f"[DOC Summarizer] Processing {len(sections)} sections with {summary_type} summarization...")
        
        # Use comprehensive DOC summarizer
        doc_summarizer = DOCSummarizer()
        result = doc_summarizer.summarize_all_sections(sections, summary_type=summary_type)
        
        # Convert to response format
        section_summaries = [
            DOCSectionSummary(
                section_number=ss["section_number"],
                summary=ss["summary"],
                original_length=ss["original_length"],
                summary_length=ss["summary_length"]
            )
            for ss in result["section_summaries"]
        ]
        
        response = DOCSummarizerResponse(
            total_sections=result["total_sections"],
            sections_summarized=result["sections_summarized"],
            section_summaries=section_summaries,
            overall_summary=result["overall_summary"],
            summary_type=summary_type,
            provider="doc_summarizer",
            success=True
        )
        
        # Optionally save summary for later review
        if save_summary:
            title = summary_title or file.filename or f"DOC Summary - {datetime.now().strftime('%Y-%m-%d')}"
            save_summary_to_store(
                title=title,
                content=result["overall_summary"],
                source=file.filename,
                source_type="docx",
                metadata={"summary_type": summary_type, "total_sections": result["total_sections"]}
            )
            print(f"[DOC Summarizer] Summary saved for later review")
        
        return response
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing DOCX: {str(e)}")

# Dedicated PDF Summarizer endpoint - Comprehensive multi-page summarization
@app.post("/summarize-pdf", response_model=PDFSummarizerResponse)
async def summarize_pdf(
    file: UploadFile = File(...),
    summary_type: str = Form("comprehensive"),
    save_summary: bool = Form(False),
    summary_title: Optional[str] = Form(None)
):
    """
    Comprehensive PDF summarizer that processes ALL pages individually and creates an overall summary.
    This endpoint is specifically designed for multi-page PDFs and ensures all points are covered.
    
    Features:
    - Processes each page individually
    - Creates page-by-page summaries
    - Generates comprehensive overall summary covering all pages
    - Uses BART model for contextual, abstractive summarization
    - Supports OCR for scanned PDFs
    
    Summary types:
    - "concise": Brief summaries (2-3 sentences per page)
    - "detailed": Comprehensive summaries with all key points
    - "comprehensive": Most detailed summaries covering all aspects (default)
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="File is required")
    
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported for this endpoint")
    
    if not PDF_SUPPORT:
        raise HTTPException(status_code=400, detail="PDF support not available. Install PyPDF2: pip install PyPDF2")
    
    if not PDF_SUMMARIZER_SUPPORT:
        raise HTTPException(
            status_code=400, 
            detail="PDF Summarizer not available. Install: pip install transformers torch sentencepiece"
        )
    
    summary_type = summary_type.lower()
    if summary_type not in ["concise", "detailed", "comprehensive"]:
        summary_type = "comprehensive"
    
    try:
        # Extract pages from PDF
        pages = await extract_pages_from_pdf(file)
        
        if not pages or all(not p or len(p.strip()) < 50 for p in pages):
            raise HTTPException(
                status_code=400, 
                detail="Could not extract sufficient text from PDF. File may be empty, corrupted, or all pages are images."
            )
        
        print(f"[PDF Summarizer] Processing {len(pages)} pages with {summary_type} summarization...")
        
        # Use comprehensive PDF summarizer
        pdf_summarizer = PDFSummarizer()
        result = pdf_summarizer.summarize_all_pages(pages, summary_type=summary_type, improve_grammar=False)
        
        # Improve grammar of all page summaries (DISABLED to prevent 429 Rate Limits)
        # print("[PDF Summarizer] Improving grammar of page summaries...")
        # for ps in result["page_summaries"]:
        #    if ps["summary_length"] > 0:
        #        try:
        #            ps["summary"] = await improve_grammar_and_clarity(ps["summary"])
        #        except Exception as e:
        #            print(f"[PDF Summarizer] Error improving grammar for page {ps['page_number']}: {str(e)}")
        
        # Improve grammar of overall summary (DISABLED to prevent 429 Rate Limits)
        # if result["overall_summary"]:
        #    try:
        #        result["overall_summary"] = await improve_grammar_and_clarity(result["overall_summary"])
        #    except Exception as e:
        #        print(f"[PDF Summarizer] Error improving overall summary grammar: {str(e)}")
        
        # Convert to response format
        page_summaries = [
            PDFPageSummary(
                page_number=ps["page_number"],
                summary=ps["summary"],
                original_length=ps["original_length"],
                summary_length=ps["summary_length"]
            )
            for ps in result["page_summaries"]
        ]
        
        response = PDFSummarizerResponse(
            total_pages=result["total_pages"],
            pages_summarized=result["pages_summarized"],
            page_summaries=page_summaries,
            overall_summary=result["overall_summary"],
            summary_type=summary_type,
            provider="pdf_summarizer",
            success=True
        )
        
        # Optionally save summary for later review
        if save_summary:
            title = summary_title or file.filename or f"PDF Summary - {datetime.now().strftime('%Y-%m-%d')}"
            save_summary_to_store(
                title=title,
                content=result["overall_summary"],
                source=file.filename,
                source_type="pdf",
                metadata={"summary_type": summary_type, "total_pages": result["total_pages"]}
            )
            print(f"[PDF Summarizer] Summary saved for later review")
        
        return response
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing PDF: {str(e)}")

# ==================== CHATBOT FUNCTIONALITY ====================

# Helper Functions for Chatbot

def get_or_create_conversation(conversation_id: Optional[str] = None) -> str:
    """Get existing conversation or create a new one with the provided ID"""
    if conversation_id:
        if conversation_id in conversation_store:
            return conversation_id
        else:
            # Create new conversation with the provided ID
            conversation_store[conversation_id] = {
                "messages": [],
                "created_at": datetime.now().isoformat(),
                "updated_at": datetime.now().isoformat()
            }
            return conversation_id
    
    # Create new conversation with UUID if no ID provided
    new_id = str(uuid.uuid4())
    conversation_store[new_id] = {
        "messages": [],
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat()
    }
    return new_id

def add_message_to_conversation(conversation_id: str, role: str, content: str):
    """Add a message to conversation history"""
    if conversation_id not in conversation_store:
        get_or_create_conversation(conversation_id)
    
    message = {
        "role": role,
        "content": content,
        "timestamp": datetime.now().isoformat()
    }
    conversation_store[conversation_id]["messages"].append(message)
    conversation_store[conversation_id]["updated_at"] = datetime.now().isoformat()

def get_conversation_history(conversation_id: str, max_messages: int = 10) -> List[Dict]:
    """Get conversation history (limited to max_messages)"""
    if conversation_id not in conversation_store:
        return []
    
    messages = conversation_store[conversation_id]["messages"]
    # Return last max_messages messages
    return messages[-max_messages:] if len(messages) > max_messages else messages

def retrieve_rag_context(query: str, top_k: int = 3) -> str:
    """Retrieve relevant context from RAG knowledge store using simple keyword matching"""
    try:
        if not rag_knowledge_store:
            return ""
        
        if not isinstance(query, str):
            print(f"[RAG] Warning: Query is not a string: {type(query)}")
            return ""

        query_lower = query.lower()
        query_words = set(query_lower.split())
        
        # Simple keyword-based retrieval (can be upgraded to vector similarity)
        scored_items = []
        for item_id, item in rag_knowledge_store.items():
            if not isinstance(item, dict) or "text" not in item:
                continue
                
            text = item.get("text", "")
            if not isinstance(text, str):
                continue

            text_lower = text.lower()
            text_words = set(text_lower.split())
            
            # Calculate simple overlap score
            overlap = len(query_words.intersection(text_words))
            if overlap > 0:
                scored_items.append((overlap, text))
        
        # Sort by score and return top_k
        scored_items.sort(reverse=True, key=lambda x: x[0])
        context_parts = [text for _, text in scored_items[:top_k]]
        
        if context_parts:
            print(f"[RAG] Retrieved {len(context_parts)} context parts for query: '{query[:50]}...'")
            return "\n\n".join([f"Context: {text}" for text in context_parts])
        return ""
    except Exception as e:
        print(f"[RAG] Error in retrieve_rag_context: {str(e)}")
        return ""

def add_to_rag_store(text: str, metadata: Optional[Dict] = None):
    """Add text to RAG knowledge store"""
    try:
        if not text or not isinstance(text, str):
            print(f"[RAG] Warning: Cannot add non-string or empty text to store: {type(text)}")
            return
            
        text_hash = hashlib.md5(text.encode()).hexdigest()
        rag_knowledge_store[text_hash] = {
            "text": text,
            "metadata": metadata or {},
            "added_at": datetime.now().isoformat()
        }
        print(f"[RAG] Added content to store (Hash: {text_hash[:8]})")
    except Exception as e:
        print(f"[RAG] Error in add_to_rag_store: {str(e)}")

async def chat_with_groq(messages: List[Dict], use_rag: bool = False, rag_context: str = "") -> str:
    """Chat with Groq API"""
    if not GROQ_API_KEY:
        raise HTTPException(status_code=500, detail="Groq API key not configured")
    
    # Prepare messages
    system_message = {
        "role": "system",
        "content": "You are a helpful, knowledgeable, and friendly AI assistant. You help users learn, solve problems, and answer questions clearly and accurately."
    }
    
    if use_rag and rag_context:
        system_message["content"] += f"\n\nRelevant context from knowledge base:\n{rag_context}\n\nUse this context to provide accurate and helpful responses."
    
    api_messages = [system_message] + messages
    
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": GROQ_MODEL_NAME,
        "messages": api_messages,
        "temperature": 0.7,
        "max_tokens": 2000
    }
    
    try:
        response = requests.post(GROQ_API_ENDPOINT, headers=headers, json=payload, timeout=45)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Groq API error: {str(e)}")

async def chat_with_ollama(messages: List[Dict], use_rag: bool = False, rag_context: str = "") -> str:
    """Chat with Ollama API"""
    # Convert messages to prompt format
    prompt_parts = []
    
    if use_rag and rag_context:
        prompt_parts.append(f"Relevant context:\n{rag_context}\n\n")
    
    for msg in messages:
        role = msg["role"]
        content = msg["content"]
        if role == "user":
            prompt_parts.append(f"User: {content}")
        elif role == "assistant":
            prompt_parts.append(f"Assistant: {content}")
    
    prompt_parts.append("Assistant:")
    prompt = "\n".join(prompt_parts)
    
    url = f"{OLLAMA_BASE_URL}/api/generate"
    payload = {
        "model": OLLAMA_MODEL_PRIMARY,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.7,
            "num_predict": 2000
        }
    }
    
    try:
        response = requests.post(url, json=payload, timeout=90)
        response.raise_for_status()
        data = response.json()
        return data.get("response", "Error generating response from Ollama")
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Ollama service not available. Make sure Ollama is running on localhost:11434")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ollama API error: {str(e)}")

async def chat_with_llama(messages: List[Dict], use_rag: bool = False, rag_context: str = "") -> str:
    """Chat with Local LLaMA API"""
    system_message = {
        "role": "system",
        "content": "You are a helpful, knowledgeable, and friendly AI assistant."
    }
    
    if use_rag and rag_context:
        system_message["content"] += f"\n\nRelevant context:\n{rag_context}\n\nUse this context to provide accurate responses."
    
    api_messages = [system_message] + messages
    
    headers = {"Content-Type": "application/json"}
    payload = {
        "model": "llama3.1",
        "messages": api_messages,
        "temperature": 0.7,
        "max_tokens": 2000
    }
    
    try:
        response = requests.post(LOCAL_LLAMA_API_URL, headers=headers, json=payload, timeout=90)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Local LLaMA service not available. Make sure the local LLaMA server is running")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLaMA API error: {str(e)}")

async def chat_with_fallback(messages: List[Dict], use_rag: bool = False, rag_context: str = "") -> Tuple[str, str]:
    """Chat with fallback logic: Groq -> Ollama -> LLaMA"""
    errors = []
    
    # Try Groq first
    if GROQ_API_KEY:
        try:
            print("[Chat] Attempting Groq...")
            response = await chat_with_groq(messages, use_rag, rag_context)
            return response, "groq"
        except Exception as e:
            err_msg = f"Groq failed: {str(e)}"
            print(f"[Chat] {err_msg}")
            errors.append(err_msg)
    else:
        errors.append("Groq API key not configured")
    
    # Try Ollama
    try:
        print("[Chat] Attempting Ollama...")
        response = await chat_with_ollama(messages, use_rag, rag_context)
        return response, "ollama"
    except Exception as e:
        err_msg = f"Ollama failed: {str(e)}"
        print(f"[Chat] {err_msg}")
        errors.append(err_msg)
    
    # Try LLaMA
    try:
        print("[Chat] Attempting LLaMA...")
        response = await chat_with_llama(messages, use_rag, rag_context)
        return response, "llama"
    except Exception as e:
        err_msg = f"LLaMA failed: {str(e)}"
        print(f"[Chat] {err_msg}")
        errors.append(err_msg)
    
    error_detail = " | ".join(errors)
    raise HTTPException(status_code=500, detail=f"All chat providers failed. Details: {error_detail}")

# Chatbot Endpoint
@app.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """
    Chatbot endpoint with conversation memory and RAG support.
    """
    print(f"\n[Chat] New request: '{request.message[:50]}...' (RAG: {request.use_rag}, Provider: {request.provider})")
    
    if not request.message or len(request.message.strip()) < 1:
        raise HTTPException(status_code=400, detail="Message is required")
    
    try:
        # Get or create conversation
        conv_start = datetime.now()
        conversation_id = get_or_create_conversation(request.conversation_id)
        
        # Get conversation history
        history = get_conversation_history(conversation_id, request.max_history)
        
        # Convert history to API format
        api_messages = []
        for msg in history:
            api_messages.append({
                "role": msg["role"],
                "content": msg["content"]
            })
        
        # Add current user message
        api_messages.append({
            "role": "user",
            "content": request.message
        })
        
        # Retrieve RAG context if enabled
        rag_context = ""
        if request.use_rag:
            try:
                rag_context = retrieve_rag_context(request.message)
                print(f"[Chat] RAG context retrieved: {len(rag_context)} characters")
            except Exception as e:
                print(f"[Chat] RAG retrieval failed (non-fatal): {str(e)}")
        
        # Store user message
        add_message_to_conversation(conversation_id, "user", request.message)
        
        # Get response from AI provider
        provider = request.provider.lower() if request.provider else "auto"
        print(f"[Chat] Processing with provider: {provider}")
        
        if provider == "groq":
            response_text = await chat_with_groq(api_messages, request.use_rag, rag_context)
            used_provider = "groq"
        elif provider == "ollama":
            response_text = await chat_with_ollama(api_messages, request.use_rag, rag_context)
            used_provider = "ollama"
        elif provider == "llama":
            response_text = await chat_with_llama(api_messages, request.use_rag, rag_context)
            used_provider = "llama"
        elif provider == "auto":
            response_text, used_provider = await chat_with_fallback(api_messages, request.use_rag, rag_context)
        else:
            raise HTTPException(status_code=400, detail="Invalid provider. Use 'auto', 'groq', 'ollama', or 'llama'")
        
        # Store assistant response
        add_message_to_conversation(conversation_id, "assistant", response_text)
        
        # Add conversation to RAG store (for future context retrieval)
        if request.use_rag:
            try:
                conversation_text = f"User: {request.message}\nAssistant: {response_text}"
                add_to_rag_store(conversation_text, {"conversation_id": conversation_id, "type": "conversation"})
            except Exception as e:
                print(f"[Chat] Failed to add to RAG store (non-fatal): {str(e)}")
        
        msg_count = len(conversation_store[conversation_id]["messages"])
        duration = (datetime.now() - conv_start).total_seconds()
        print(f"[Chat] Success! Response length: {len(response_text)}, Duration: {duration:.2f}s, Msg Count: {msg_count}")
        
        return ChatResponse(
            response=response_text,
            conversation_id=conversation_id,
            provider=used_provider,
            message_count=msg_count,
            timestamp=datetime.now().isoformat(),
            success=True
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"[Chat] CRITICAL ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error in chat: {str(e)}")

# Get Conversation History Endpoint
@app.get("/chat/history/{conversation_id}", response_model=ConversationHistory)
async def get_conversation_history_endpoint(conversation_id: str):
    """Get full conversation history for a given conversation_id"""
    if conversation_id not in conversation_store:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    conv_data = conversation_store[conversation_id]
    messages = [ChatMessage(**msg) for msg in conv_data["messages"]]
    
    return ConversationHistory(
        conversation_id=conversation_id,
        messages=messages,
        created_at=conv_data["created_at"],
        updated_at=conv_data["updated_at"],
        message_count=len(messages)
    )

# Delete Conversation Endpoint
@app.delete("/chat/history/{conversation_id}")
async def delete_conversation(conversation_id: str):
    """Delete a conversation and its history"""
    if conversation_id not in conversation_store:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    del conversation_store[conversation_id]
    return {"message": "Conversation deleted successfully", "conversation_id": conversation_id}

# RAG Knowledge Store Endpoint (renamed from /chat/knowledge)
# RAG Knowledge Store Endpoint (renamed from /chat/knowledge)
class KnowledgeRequest(BaseModel):
    text: str
    metadata: Optional[Dict] = None

@app.post("/rag/knowledge")
async def add_knowledge(request: KnowledgeRequest):
    """Add knowledge to RAG store for context retrieval in chat"""
    if not request.text or len(request.text.strip()) < 10:
        raise HTTPException(status_code=400, detail="Text must be at least 10 characters")
    
    add_to_rag_store(request.text, request.metadata)
    return {"message": "Knowledge added to RAG store successfully", "text_length": len(request.text)}

# ==================== SAVED SUMMARIZER WITH QUESTION GENERATION ====================

# Helper Functions for Saved Summarizer

def save_summary_to_store(title: str, content: str, source: Optional[str] = None, 
                 source_type: Optional[str] = None, metadata: Optional[Dict] = None) -> str:
    """Save a summary to the store"""
    summary_id = str(uuid.uuid4())
    saved_summaries_store[summary_id] = {
        "title": title,
        "content": content,
        "source": source,
        "source_type": source_type,
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat(),
        "question_count": 0,
        "metadata": metadata or {}
    }
    return summary_id

def get_summary(summary_id: str) -> Optional[Dict]:
    """Get a saved summary"""
    return saved_summaries_store.get(summary_id)

def create_review_schedule() -> Dict:
    """Create spaced repetition schedule: same day, next day, 1 week later"""
    now = datetime.now()
    return {
        "first_review": now.isoformat(),  # Same day
        "second_review": (now + timedelta(days=1)).isoformat(),  # Next day
        "third_review": (now + timedelta(days=7)).isoformat(),  # 1 week later
        "next_review_date": now.isoformat(),  # Start with first review
        "last_review_date": None,
        "review_count": 0,
        "mastery_level": "learning"
    }

def update_review_schedule(question_id: str, is_correct: bool, confidence: str = "medium"):
    """Update review schedule based on performance"""
    if question_id not in questions_store:
        return
    
    question_data = questions_store[question_id]
    schedule = question_data.get("review_schedule", create_review_schedule())
    
    now = datetime.now()
    schedule["last_review_date"] = now.isoformat()
    schedule["review_count"] += 1
    
    # Calculate next review based on performance and current review count
    if schedule["review_count"] == 1:
        # After first review, schedule next day
        schedule["next_review_date"] = (now + timedelta(days=1)).isoformat()
    elif schedule["review_count"] == 2:
        # After second review, schedule 1 week later
        schedule["next_review_date"] = (now + timedelta(days=7)).isoformat()
    else:
        # After third review, adjust based on performance
        if is_correct and confidence in ["high", "medium"]:
            # If correct, increase interval
            if schedule["mastery_level"] == "learning":
                schedule["mastery_level"] = "reviewing"
                schedule["next_review_date"] = (now + timedelta(days=14)).isoformat()
            elif schedule["mastery_level"] == "reviewing":
                schedule["mastery_level"] = "mastered"
                schedule["next_review_date"] = (now + timedelta(days=30)).isoformat()
        else:
            # If incorrect, reset to learning
            schedule["mastery_level"] = "learning"
            schedule["next_review_date"] = (now + timedelta(days=1)).isoformat()
    
    question_data["review_schedule"] = schedule
    questions_store[question_id] = question_data

async def generate_questions_with_ai(summary_content: str, question_types: List[str], 
                                     num_questions: int, difficulty: str) -> List[Dict]:
    """Generate questions using AI (Groq, Ollama, or LLaMA)"""
    questions = []
    
    # Create prompt for question generation
    question_type_names = {
        "qa": "Question-Answer pairs",
        "mcq": "Multiple Choice Questions (MCQs)",
        "flashcard": "Flashcards",
        "case_based": "Case-based questions"
    }
    
    types_str = ", ".join([question_type_names.get(qt, qt) for qt in question_types])
    
    prompt = f"""Generate {num_questions} {types_str} based on the following summary content.

Summary:
{summary_content}

Requirements:
- Difficulty level: {difficulty}
- Questions should test understanding of key concepts
- Answers should be clear and accurate
- For MCQs: Provide 4 options with only one correct answer
- For case-based: Create realistic scenarios related to the content
- For flashcards: Create concise Q&A pairs

Format your response as JSON with this structure:
{{
  "questions": [
    {{
      "type": "qa|mcq|flashcard|case_based",
      "question": "...",
      "answer": "...",
      "options": [{{"option": "...", "is_correct": true/false}}, ...],  // Only for MCQs
      "explanation": "..."  // Optional
    }}
  ]
}}

Generate the questions now:"""
    
    # Try to generate with AI
    messages = [{"role": "user", "content": prompt}]
    
    try:
        # Try Groq first
        if GROQ_API_KEY:
            response_text = await chat_with_groq(messages, use_rag=False, rag_context="")
            # Parse response (may need JSON extraction)
            # For now, we'll create questions from the response
            questions = parse_ai_questions(response_text, question_types, num_questions, difficulty)
            if questions:
                return questions
    except Exception as e:
        print(f"Groq question generation failed: {str(e)}")
    
    # Fallback: Generate simple questions from summary
    return generate_simple_questions(summary_content, question_types, num_questions, difficulty)

def parse_ai_questions(response_text: str, question_types: List[str], 
                      num_questions: int, difficulty: str) -> List[Dict]:
    """Parse AI-generated questions from response text"""
    # Try to extract JSON from response
    import re
    json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
    if json_match:
        try:
            import json
            data = json.loads(json_match.group())
            return data.get("questions", [])
        except:
            pass
    
    # If JSON parsing fails, create questions from text
    return generate_simple_questions(response_text, question_types, num_questions, difficulty)

def generate_simple_questions(summary_content: str, question_types: List[str], 
                              num_questions: int, difficulty: str) -> List[Dict]:
    """Generate simple questions from summary content (fallback)"""
    questions = []
    sentences = summary_content.split('. ')
    key_sentences = [s.strip() for s in sentences if len(s.strip()) > 20][:num_questions * 2]
    
    for i, qtype in enumerate(question_types):
        for j in range(num_questions):
            if i * num_questions + j >= len(key_sentences):
                break
            
            sentence = key_sentences[i * num_questions + j]
            question_id = str(uuid.uuid4())
            
            if qtype == "qa":
                questions.append({
                    "question_id": question_id,
                    "question_type": "qa",
                    "question": f"Explain: {sentence[:100]}...",
                    "answer": sentence,
                    "explanation": None,
                    "difficulty": difficulty
                })
            elif qtype == "mcq":
                # Create simple MCQ
                questions.append({
                    "question_id": question_id,
                    "question_type": "mcq",
                    "question": f"Which statement is true about: {sentence[:80]}...?",
                    "answer": sentence[:100],
                    "options": [
                        {"option": sentence[:100], "is_correct": True},
                        {"option": "Option 2 (incorrect)", "is_correct": False},
                        {"option": "Option 3 (incorrect)", "is_correct": False},
                        {"option": "Option 4 (incorrect)", "is_correct": False}
                    ],
                    "explanation": None,
                    "difficulty": difficulty
                })
            elif qtype == "flashcard":
                questions.append({
                    "question_id": question_id,
                    "question_type": "flashcard",
                    "question": f"What is: {sentence[:80]}...?",
                    "answer": sentence,
                    "explanation": None,
                    "difficulty": difficulty
                })
            elif qtype == "case_based":
                questions.append({
                    "question_id": question_id,
                    "question_type": "case_based",
                    "question": f"Case: A scenario related to {sentence[:60]}... How would you apply this concept?",
                    "answer": f"Based on the concept: {sentence}",
                    "explanation": None,
                    "difficulty": difficulty
                })
    
    return questions[:num_questions * len(question_types)]

# Saved Summarizer Endpoints

@app.post("/summaries/save", response_model=SavedSummary)
async def save_summary_endpoint(request: SaveSummaryRequest):
    """
    Save a summary for later review and question generation.
    Can save summaries from PDF, DOC, or text sources.
    """
    if not request.title or not request.content:
        raise HTTPException(status_code=400, detail="Title and content are required")
    
    if len(request.content.strip()) < 50:
        raise HTTPException(status_code=400, detail="Content must be at least 50 characters")
    
    summary_id = save_summary_to_store(
        title=request.title,
        content=request.content,
        source=request.source,
        source_type=request.source_type,
        metadata=request.metadata
    )
    
    summary_data = saved_summaries_store[summary_id]
    return SavedSummary(
        summary_id=summary_id,
        title=summary_data["title"],
        content=summary_data["content"],
        source=summary_data["source"],
        source_type=summary_data["source_type"],
        created_at=summary_data["created_at"],
        updated_at=summary_data["updated_at"],
        question_count=summary_data["question_count"],
        metadata=summary_data["metadata"]
    )

@app.get("/summaries", response_model=List[SavedSummary])
async def list_summaries():
    """Get all saved summaries"""
    summaries = []
    for summary_id, data in saved_summaries_store.items():
        summaries.append(SavedSummary(
            summary_id=summary_id,
            title=data["title"],
            content=data["content"][:200] + "..." if len(data["content"]) > 200 else data["content"],
            source=data["source"],
            source_type=data["source_type"],
            created_at=data["created_at"],
            updated_at=data["updated_at"],
            question_count=data["question_count"],
            metadata=data["metadata"]
        ))
    return summaries

@app.get("/summaries/{summary_id}", response_model=SavedSummary)
async def get_summary_endpoint(summary_id: str):
    """Get a specific saved summary"""
    summary_data = get_summary(summary_id)
    if not summary_data:
        raise HTTPException(status_code=404, detail="Summary not found")
    
    return SavedSummary(
        summary_id=summary_id,
        title=summary_data["title"],
        content=summary_data["content"],
        source=summary_data["source"],
        source_type=summary_data["source_type"],
        created_at=summary_data["created_at"],
        updated_at=summary_data["updated_at"],
        question_count=summary_data["question_count"],
        metadata=summary_data["metadata"]
    )

@app.post("/summaries/{summary_id}/questions/generate", response_model=QuestionGenerationResponse)
async def generate_questions_endpoint(summary_id: str, request: QuestionGenerationRequest):
    """
    Generate practice questions from a saved summary.
    Supports: Q&A, MCQs, Flashcards, Case-based questions.
    Questions are automatically scheduled for spaced repetition review.
    """
    summary_data = get_summary(summary_id)
    if not summary_data:
        raise HTTPException(status_code=404, detail="Summary not found")
    
    # Use request parameters or defaults
    question_types = request.question_types or ["qa", "mcq", "flashcard", "case_based"]
    num_questions = request.num_questions or 5
    difficulty = request.difficulty or "medium"
    
    # Generate questions
    generated_questions = await generate_questions_with_ai(
        summary_data["content"],
        question_types,
        num_questions,
        difficulty
    )
    
    # Save questions with review schedule
    saved_questions = []
    for q in generated_questions:
        question_id = q.get("question_id", str(uuid.uuid4()))
        review_schedule = create_review_schedule()
        
        question_data = {
            "summary_id": summary_id,
            "question_type": q["question_type"],
            "question": q["question"],
            "answer": q["answer"],
            "options": q.get("options"),
            "explanation": q.get("explanation"),
            "difficulty": q.get("difficulty", difficulty),
            "created_at": datetime.now().isoformat(),
            "review_schedule": review_schedule
        }
        
        questions_store[question_id] = question_data
        saved_questions.append(Question(
            question_id=question_id,
            summary_id=summary_id,
            question_type=q["question_type"],
            question=q["question"],
            answer=q["answer"],
            options=[MCQOption(**opt) for opt in q.get("options", [])] if q.get("options") else None,
            explanation=q.get("explanation"),
            difficulty=q.get("difficulty", difficulty),
            created_at=question_data["created_at"]
        ))
    
    # Update summary question count
    saved_summaries_store[summary_id]["question_count"] = len(saved_questions)
    saved_summaries_store[summary_id]["updated_at"] = datetime.now().isoformat()
    
    return QuestionGenerationResponse(
        summary_id=summary_id,
        questions=saved_questions,
        total_questions=len(saved_questions),
        success=True
    )

@app.get("/summaries/{summary_id}/questions", response_model=List[Question])
async def get_summary_questions(summary_id: str):
    """Get all questions for a specific summary"""
    if summary_id not in saved_summaries_store:
        raise HTTPException(status_code=404, detail="Summary not found")
    
    questions = []
    for question_id, question_data in questions_store.items():
        if question_data["summary_id"] == summary_id:
            questions.append(Question(
                question_id=question_id,
                summary_id=summary_id,
                question_type=question_data["question_type"],
                question=question_data["question"],
                answer=question_data["answer"],
                options=[MCQOption(**opt) for opt in question_data.get("options", [])] if question_data.get("options") else None,
                explanation=question_data.get("explanation"),
                difficulty=question_data.get("difficulty", "medium"),
                created_at=question_data["created_at"]
            ))
    
    return questions

@app.get("/reviews/pending", response_model=List[PendingReview])
async def get_pending_reviews():
    """Get all questions that are due for review (spaced repetition)"""
    now = datetime.now()
    pending = []
    
    for question_id, question_data in questions_store.items():
        schedule = question_data.get("review_schedule", {})
        next_review_str = schedule.get("next_review_date")
        
        if next_review_str:
            next_review = datetime.fromisoformat(next_review_str.replace('Z', '+00:00'))
            # Handle timezone-aware datetime
            if next_review.tzinfo:
                now_aware = now.replace(tzinfo=next_review.tzinfo)
                if now_aware >= next_review:
                    days_until = (now_aware - next_review).days
                    pending.append(PendingReview(
                        question_id=question_id,
                        summary_id=question_data["summary_id"],
                        question=question_data["question"],
                        question_type=question_data["question_type"],
                        answer=question_data["answer"],
                        next_review_date=next_review_str,
                        days_until_review=days_until
                    ))
            else:
                if now >= next_review:
                    days_until = (now - next_review).days
                    pending.append(PendingReview(
                        question_id=question_id,
                        summary_id=question_data["summary_id"],
                        question=question_data["question"],
                        question_type=question_data["question_type"],
                        answer=question_data["answer"],
                        next_review_date=next_review_str,
                        days_until_review=days_until
                    ))
    
    # Sort by days overdue (most overdue first)
    pending.sort(key=lambda x: x.days_until_review, reverse=True)
    return pending

@app.post("/reviews/submit", response_model=ReviewResponse)
async def submit_review(request: ReviewRequest):
    """
    Submit a review for a question (spaced repetition).
    Updates the review schedule based on performance.
    """
    if request.question_id not in questions_store:
        raise HTTPException(status_code=404, detail="Question not found")
    
    question_data = questions_store[request.question_id]
    update_review_schedule(request.question_id, request.is_correct, request.confidence)
    
    schedule = question_data["review_schedule"]
    
    # Determine message based on performance
    if request.is_correct:
        if schedule["mastery_level"] == "mastered":
            message = "Excellent! You've mastered this concept."
        elif schedule["mastery_level"] == "reviewing":
            message = "Great job! Keep reviewing to maintain mastery."
        else:
            message = "Well done! Continue practicing to reinforce learning."
    else:
        message = "Keep practicing! Review the answer and try again later."
    
    return ReviewResponse(
        question_id=request.question_id,
        correct_answer=question_data["answer"],
        is_correct=request.is_correct,
        next_review_date=schedule["next_review_date"],
        review_count=schedule["review_count"],
        mastery_level=schedule["mastery_level"],
        message=message
    )

@app.get("/reviews/stats")
async def get_review_stats():
    """Get statistics about reviews and learning progress"""
    total_questions = len(questions_store)
    learning = sum(1 for q in questions_store.values() 
                  if q.get("review_schedule", {}).get("mastery_level") == "learning")
    reviewing = sum(1 for q in questions_store.values() 
                   if q.get("review_schedule", {}).get("mastery_level") == "reviewing")
    mastered = sum(1 for q in questions_store.values() 
                  if q.get("review_schedule", {}).get("mastery_level") == "mastered")
    
    now = datetime.now()
    pending_count = 0
    for q in questions_store.values():
        schedule = q.get("review_schedule", {})
        next_review_str = schedule.get("next_review_date")
        if next_review_str:
            try:
                next_review = datetime.fromisoformat(next_review_str.replace('Z', '+00:00'))
                if next_review.tzinfo:
                    if now.replace(tzinfo=next_review.tzinfo) >= next_review:
                        pending_count += 1
                else:
                    if now >= next_review:
                        pending_count += 1
            except:
                pass
    
    return {
        "total_questions": total_questions,
        "learning": learning,
        "reviewing": reviewing,
        "mastered": mastered,
        "pending_reviews": pending_count,
        "total_summaries": len(saved_summaries_store)
    }

# ==================== QUIZ/TEST API ====================

async def generate_quiz_questions(subject: str, topic: str, provider: str = "auto", difficulty: str = "medium") -> List[Dict]:
    """Generate 10 quiz questions using AI models"""
    prompt = f"""Generate exactly 10 multiple-choice quiz questions (MCQs) about {subject} - {topic}.

Requirements:
- Difficulty level: {difficulty}
- Each question must have exactly 4 options (A, B, C, D)
- Only one option should be correct
- Questions should test understanding of key concepts
- Format your response as JSON with this structure:
{{
  "questions": [
    {{
      "question": "Question text here?",
      "options": ["Option A", "Option B", "Option C", "Option D"],
      "correct_answer": "Option A",  // The correct option text
      "correct_index": 0,  // Index of correct option (0-3)
      "explanation": "Brief explanation of why this is correct"
    }}
  ]
}}

Generate 10 questions now:"""
    
    messages = [{"role": "user", "content": prompt}]
    questions = []
    
    # Try to generate with AI
    try:
        if provider == "groq" or (provider == "auto" and GROQ_API_KEY):
            try:
                response_text = await chat_with_groq(messages, use_rag=False, rag_context="")
                questions = parse_quiz_questions(response_text, subject, topic, difficulty)
                if questions and len(questions) >= 10:
                    return questions[:10]
            except Exception as e:
                print(f"Groq quiz generation failed: {str(e)}")
                if provider == "groq":
                    raise
        
        if provider == "ollama" or (provider == "auto"):
            try:
                response_text = await chat_with_ollama(messages, use_rag=False, rag_context="")
                questions = parse_quiz_questions(response_text, subject, topic, difficulty)
                if questions and len(questions) >= 10:
                    return questions[:10]
            except Exception as e:
                print(f"Ollama quiz generation failed: {str(e)}")
                if provider == "ollama":
                    raise
        
        if provider == "llama" or (provider == "auto"):
            try:
                response_text = await chat_with_llama(messages, use_rag=False, rag_context="")
                questions = parse_quiz_questions(response_text, subject, topic, difficulty)
                if questions and len(questions) >= 10:
                    return questions[:10]
            except Exception as e:
                print(f"LLaMA quiz generation failed: {str(e)}")
                if provider == "llama":
                    raise
    except Exception as e:
        print(f"AI quiz generation failed: {str(e)}")
    
    # Fallback: Generate simple questions
    return generate_fallback_quiz_questions(subject, topic, difficulty)

def parse_quiz_questions(response_text: str, subject: str, topic: str, difficulty: str) -> List[Dict]:
    """Parse AI-generated quiz questions from response"""
    import re
    import json
    
    # Try to extract JSON from response
    json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
    if json_match:
        try:
            data = json.loads(json_match.group())
            questions_data = data.get("questions", [])
            
            questions = []
            for i, q_data in enumerate(questions_data[:10], 1):
                question_id = str(uuid.uuid4())
                options = q_data.get("options", [])
                correct_answer = q_data.get("correct_answer", "")
                correct_index = q_data.get("correct_index", 0)
                
                # Ensure we have 4 options
                while len(options) < 4:
                    options.append(f"Option {chr(68 + len(options))}")
                
                questions.append({
                    "question_id": question_id,
                    "question_number": i,
                    "question": q_data.get("question", f"Question {i} about {topic}?"),
                    "options": options[:4],
                    "correct_answer": correct_answer or options[correct_index] if correct_index < len(options) else options[0],
                    "correct_index": correct_index if correct_index < len(options) else 0,
                    "explanation": q_data.get("explanation", ""),
                    "question_type": "mcq"
                })
            
            return questions
        except Exception as e:
            print(f"Error parsing quiz JSON: {str(e)}")
    
    return []

def generate_fallback_quiz_questions(subject: str, topic: str, difficulty: str) -> List[Dict]:
    """Generate fallback quiz questions if AI fails"""
    questions = []
    base_questions = [
        f"What is the main concept of {topic} in {subject}?",
        f"Which of the following best describes {topic}?",
        f"What is a key characteristic of {topic}?",
        f"Which statement about {topic} is correct?",
        f"What is the primary purpose of {topic}?",
        f"Which of the following is related to {topic}?",
        f"What is an important aspect of {topic}?",
        f"Which concept is fundamental to {topic}?",
        f"What distinguishes {topic} from other concepts?",
        f"Which of the following applies to {topic}?"
    ]
    
    for i, q_text in enumerate(base_questions[:10], 1):
        question_id = str(uuid.uuid4())
        questions.append({
            "question_id": question_id,
            "question_number": i,
            "question": q_text,
            "options": [
                f"Option A: Correct answer about {topic}",
                f"Option B: Incorrect answer",
                f"Option C: Another incorrect answer",
                f"Option D: Yet another incorrect answer"
            ],
            "correct_answer": f"Option A: Correct answer about {topic}",
            "correct_index": 0,
            "explanation": f"This is the correct answer because it accurately describes {topic} in {subject}.",
            "question_type": "mcq"
        })
    
    return questions

@app.post("/quiz/generate", response_model=QuizResponse)
async def generate_quiz(request: QuizRequest):
    """
    Generate a quiz with 10 questions based on subject and topic.
    Uses AI models (Groq, Ollama, LLaMA) to create questions.
    """
    if not request.subject or not request.topic:
        raise HTTPException(status_code=400, detail="Subject and topic are required")
    
    try:
        # Generate questions
        provider = request.provider.lower() if request.provider else "auto"
        generated_questions = await generate_quiz_questions(
            request.subject,
            request.topic,
            provider,
            request.difficulty or "medium"
        )
        
        if len(generated_questions) < 10:
            # If we got fewer than 10, fill with fallback
            fallback = generate_fallback_quiz_questions(request.subject, request.topic, request.difficulty or "medium")
            generated_questions.extend(fallback[:10 - len(generated_questions)])
        
        # Create quiz
        quiz_id = str(uuid.uuid4())
        quiz_store[quiz_id] = {
            "subject": request.subject,
            "topic": request.topic,
            "questions": generated_questions,
            "created_at": datetime.now().isoformat(),
            "provider": provider,
            "difficulty": request.difficulty or "medium"
        }
        
        # Format questions for response (without answers)
        quiz_questions = []
        for q in generated_questions:
            quiz_questions.append(QuizQuestion(
                question_id=q["question_id"],
                question_number=q["question_number"],
                question=q["question"],
                options=q["options"],
                question_type=q["question_type"]
            ))
        
        return QuizResponse(
            quiz_id=quiz_id,
            subject=request.subject,
            topic=request.topic,
            questions=quiz_questions,
            total_questions=len(quiz_questions),
            created_at=quiz_store[quiz_id]["created_at"],
            provider=provider
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating quiz: {str(e)}")

@app.post("/quiz/submit", response_model=QuizResult)
async def submit_quiz(submission: QuizSubmission):
    """
    Submit quiz answers and get results.
    Returns detailed results showing correct/incorrect answers.
    """
    if submission.quiz_id not in quiz_store:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    quiz_data = quiz_store[submission.quiz_id]
    questions = quiz_data["questions"]
    
    correct_count = 0
    wrong_count = 0
    results = []
    
    for q in questions:
        question_id = q["question_id"]
        user_answer = submission.answers.get(question_id, "").strip()
        correct_answer = q["correct_answer"]
        correct_index = q.get("correct_index", 0)
        
        # Check if answer is correct (can be option text or index)
        is_correct = False
        if user_answer:
            # Check if user answered with option text
            if user_answer.lower() == correct_answer.lower():
                is_correct = True
            # Check if user answered with option letter (A, B, C, D)
            elif user_answer.upper() in ["A", "B", "C", "D"]:
                option_index = ord(user_answer.upper()) - ord("A")
                if option_index == correct_index:
                    is_correct = True
            # Check if user answered with option number (0, 1, 2, 3)
            elif user_answer.isdigit():
                if int(user_answer) == correct_index:
                    is_correct = True
        
        if is_correct:
            correct_count += 1
        else:
            wrong_count += 1
        
        results.append({
            "question_id": question_id,
            "question_number": q["question_number"],
            "question": q["question"],
            "user_answer": user_answer,
            "correct_answer": correct_answer,
            "correct_index": correct_index,
            "is_correct": is_correct,
            "explanation": q.get("explanation", "")
        })
    
    score_percentage = (correct_count / len(questions)) * 100 if questions else 0
    
    return QuizResult(
        quiz_id=submission.quiz_id,
        total_questions=len(questions),
        correct_answers=correct_count,
        wrong_answers=wrong_count,
        score_percentage=round(score_percentage, 2),
        results=results,
        submitted_at=datetime.now().isoformat()
    )

@app.get("/quiz/{quiz_id}")
async def get_quiz(quiz_id: str):
    """Get quiz details (without answers)"""
    if quiz_id not in quiz_store:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    quiz_data = quiz_store[quiz_id]
    questions = []
    
    for q in quiz_data["questions"]:
        questions.append(QuizQuestion(
            question_id=q["question_id"],
            question_number=q["question_number"],
            question=q["question"],
            options=q["options"],
            question_type=q["question_type"]
        ))
    
    return {
        "quiz_id": quiz_id,
        "subject": quiz_data["subject"],
        "topic": quiz_data["topic"],
        "questions": questions,
        "total_questions": len(questions),
        "created_at": quiz_data["created_at"],
        "difficulty": quiz_data.get("difficulty", "medium")
    }

# ==================== AGENT SIMULATORS ====================

# ==================== EduMentor Learning Assistant ====================

async def get_wikipedia_sources(subject: str, topic: str) -> List[Dict]:
    """Get Wikipedia sources for subject and topic"""
    try:
        import wikipedia
        wikipedia.set_lang("en")
        search_query = f"{subject} {topic}"
        results = wikipedia.search(search_query, results=3)
        
        sources = []
        for title in results[:3]:
            try:
                page = wikipedia.page(title, auto_suggest=False)
                sources.append({
                    "title": page.title,
                    "url": page.url,
                    "summary": page.summary[:200] + "..."
                })
            except:
                continue
        return sources
    except ImportError:
        print("[EduMentor] Wikipedia library not installed. Install with: pip install wikipedia")
        return []
    except Exception as e:
        print(f"[EduMentor] Wikipedia search failed: {str(e)}")
        return []

@app.post("/agent/edumentor/generate", response_model=EduMentorResponse)
async def edumentor_generate_lesson(request: EduMentorRequest):
    """
    EduMentor Learning Assistant - Generate personalized educational content.
    
    Features:
    - AI-powered lesson generation
    - Optional Wikipedia sources
    - Knowledge store integration
    - Orchestration engine support
    """
    if not request.subject or not request.topic:
        raise HTTPException(status_code=400, detail="Subject and topic are required")
    
    try:
        provider = request.provider.lower() if request.provider else "auto"
        
        # Build enhanced prompt
        prompt_parts = [
            f"You are EduMentor, a personal learning assistant specializing in {request.subject}.",
            f"Create a comprehensive, engaging lesson about '{request.topic}' in {request.subject}.",
            "",
            "**FORMATTING RULES:**",
            "- You MUST use Markdown for all formatting.",
            "- Use `#` for the Main Title.",
            "- Use `##` for Section Headers (Introduction, Key Concepts, etc.).",
            "- Use `###` for sub-sections if needed.",
            "- Use `**` for bold key terms and emphasis.",
            "- Use `-` or `1.` for lists. Do not use plain paragraphs for lists.",
            "- Use `>` for important callouts or summaries.",
            "",
            "The lesson content should include:",
            "1. **Introduction**: Clear overview and importance",
            "2. **Key Concepts**: Break down main ideas simply (use bullet points)",
            "3. **Detailed Explanation**: Thorough explanation with examples",
            "4. **Real-world Applications**: Practical uses",
            "5. **Practice Exercises**: 2-3 exercises to reinforce learning",
            "6. **Summary**: Key takeaways",
            "",
            "Write in a friendly, encouraging tone. Use analogies and examples.",
            f"Subject: {request.subject}",
            f"Topic: {request.topic}"
        ]
        
        # Add knowledge store context if enabled
        knowledge_context = ""
        if request.use_knowledge_store:
            knowledge_context = retrieve_rag_context(f"{request.subject} {request.topic}")
            if knowledge_context:
                prompt_parts.append(f"\n\nAdditional Context from Knowledge Store:\n{knowledge_context}")
        
        prompt = "\n".join(prompt_parts)
        
        # Get Wikipedia sources if requested
        wikipedia_sources = []
        if request.include_wikipedia:
            wikipedia_sources = await get_wikipedia_sources(request.subject, request.topic)
            if wikipedia_sources:
                sources_text = "\n".join([f"- {s['title']}: {s['url']}" for s in wikipedia_sources])
                prompt += f"\n\nWikipedia Sources:\n{sources_text}"
        
        # Generate lesson using AI
        messages = [{"role": "user", "content": prompt}]
        lesson_content = ""
        used_provider = provider
        
        if provider == "groq" or (provider == "auto" and GROQ_API_KEY):
            try:
                lesson_content = await chat_with_groq(messages, use_rag=request.use_knowledge_store, rag_context=knowledge_context)
                used_provider = "groq"
            except Exception as e:
                print(f"Groq failed: {str(e)}")
                if provider == "groq":
                    raise
        
        if not lesson_content and (provider == "ollama" or provider == "auto"):
            try:
                lesson_content = await chat_with_ollama(messages, use_rag=request.use_knowledge_store, rag_context=knowledge_context)
                used_provider = "ollama"
            except Exception as e:
                print(f"Ollama failed: {str(e)}")
                if provider == "ollama":
                    raise
        
        if not lesson_content and (provider == "llama" or provider == "auto"):
            try:
                lesson_content = await chat_with_llama(messages, use_rag=request.use_knowledge_store, rag_context=knowledge_context)
                used_provider = "llama"
            except Exception as e:
                print(f"LLaMA failed: {str(e)}")
                if provider == "llama":
                    raise
        
        if not lesson_content:
            # Fallback: Generate a basic lesson structure if all AI providers fail
            lesson_content = f"""# {request.topic} - {request.subject}

## Introduction
{request.topic} is an important concept in {request.subject}. Understanding this topic will help you build a strong foundation in the subject.

## Key Concepts
- Core principles of {request.topic}
- Fundamental relationships and patterns
- Essential terminology

## Detailed Explanation
{request.topic} involves several key aspects that are crucial to understand. Let's explore each one in detail.

## Real-world Applications
This concept has practical applications in various fields and everyday situations.

## Practice Exercises
1. Explain {request.topic} in your own words
2. Identify examples of {request.topic} in real life
3. Solve a problem related to {request.topic}

## Summary
{request.topic} is a fundamental concept in {request.subject} that requires careful study and practice.

**Note:** This is a basic lesson structure. For a complete AI-generated lesson, please ensure your AI provider (Groq, Ollama, or LLaMA) is properly configured."""
            used_provider = "fallback"
        
        # Use orchestration if enabled (enhance content)
        if request.use_orchestration and lesson_content:
            orchestration_prompt = f"""Review and enhance this educational content to make it more engaging and effective:

{lesson_content}

Enhancements needed:
- Improve clarity and flow
- Add more examples if needed
- Ensure all sections are well-structured
- Make it more interactive and engaging

Return the enhanced lesson:"""
            
            orchestration_messages = [{"role": "user", "content": orchestration_prompt}]
            try:
                if GROQ_API_KEY:
                    lesson_content = await chat_with_groq(orchestration_messages, use_rag=False, rag_context="")
                elif OLLAMA_BASE_URL:
                    lesson_content = await chat_with_ollama(orchestration_messages, use_rag=False, rag_context="")
            except:
                pass  # Keep original if orchestration fails
        
        return EduMentorResponse(
            subject=request.subject,
            topic=request.topic,
            lesson_content=lesson_content,
            wikipedia_sources=wikipedia_sources if request.include_wikipedia else None,
            knowledge_store_used=request.use_knowledge_store,
            orchestration_used=request.use_orchestration,
            provider=used_provider,
            created_at=datetime.now().isoformat(),
            success=True
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating lesson: {str(e)}")

# ==================== Financial Agent Simulator ====================

@app.post("/agent/financial/advice", response_model=FinancialAdviceResponse)
async def get_financial_advice(request: FinancialProfileRequest):
    """
    FinancialCrew - Generate personalized financial advice based on profile.
    
    Takes financial profile data and uses AI models to generate comprehensive
    financial advice, recommendations, and goal analysis.
    """
    if not request.name or request.monthly_income <= 0:
        raise HTTPException(status_code=400, detail="Name and valid monthly income are required")
    
    try:
        # Calculate financial metrics
        total_expenses = sum(exp.amount for exp in request.expense_categories)
        monthly_savings = request.monthly_income - total_expenses
        savings_rate = (monthly_savings / request.monthly_income * 100) if request.monthly_income > 0 else 0
        
        expense_breakdown = [{"name": exp.name, "amount": exp.amount, "percentage": round((exp.amount / request.monthly_income * 100), 2)} for exp in request.expense_categories]
        
        # Build comprehensive financial advice prompt
        expense_details = "\n".join([f"- {exp.name}: ₹{exp.amount} ({round((exp.amount / request.monthly_income * 100), 2)}% of income)" for exp in request.expense_categories])
        
        financial_prompt = f"""You are FinancialCrew, an expert financial advisor. Analyze this financial profile and provide comprehensive, personalized financial advice.

**Financial Profile:**
- Name: {request.name}
- Monthly Income: ₹{request.monthly_income}
- Monthly Expenses: ₹{total_expenses}
- Monthly Savings: ₹{monthly_savings}
- Savings Rate: {round(savings_rate, 2)}%

**Expense Breakdown:**
{expense_details}

**Financial Goals:**
- Primary Goal: {request.financial_goal}
- Financial Type: {request.financial_type}
- Risk Tolerance: {request.risk_level}

**Your Task:**
Provide detailed financial advice covering:
1. **Financial Health Assessment**: Evaluate their current financial situation
2. **Savings Analysis**: Comment on their savings rate and monthly savings amount
3. **Expense Management**: Analyze expense categories and suggest optimizations
4. **Goal Achievement Strategy**: Provide a clear plan to achieve "{request.financial_goal}"
5. **Investment Recommendations**: Suggest investment strategies based on their risk level ({request.risk_level})
6. **Budget Optimization**: Provide specific tips to improve their financial position
7. **Timeline Estimation**: Estimate how long it might take to achieve their goal

Write in a friendly, encouraging, and professional tone. Be specific with actionable advice. Use ₹ (rupee) symbol for all amounts."""

        # Generate financial advice using AI
        messages = [{"role": "user", "content": financial_prompt}]
        financial_advice = ""
        used_provider = "auto"
        
        if GROQ_API_KEY:
            try:
                financial_advice = await chat_with_groq(messages, use_rag=False, rag_context="")
                used_provider = "groq"
            except Exception as e:
                print(f"Groq failed: {str(e)}")
        
        if not financial_advice and OLLAMA_BASE_URL:
            try:
                financial_advice = await chat_with_ollama(messages, use_rag=False, rag_context="")
                used_provider = "ollama"
            except Exception as e:
                print(f"Ollama failed: {str(e)}")
        
        if not financial_advice and LOCAL_LLAMA_API_URL:
            try:
                financial_advice = await chat_with_llama(messages, use_rag=False, rag_context="")
                used_provider = "llama"
            except Exception as e:
                print(f"LLaMA failed: {str(e)}")
        
        # Fallback advice if all AI providers fail
        if not financial_advice:
            financial_advice = f"""**Financial Health Assessment for {request.name}**

**Current Situation:**
- Your monthly income of ₹{request.monthly_income} with expenses of ₹{total_expenses} leaves you with ₹{monthly_savings} in monthly savings.
- Your savings rate is {round(savings_rate, 2)}%, which is {'excellent' if savings_rate >= 20 else 'good' if savings_rate >= 10 else 'needs improvement'}.

**Expense Analysis:**
{expense_details}

**Goal Achievement Strategy:**
To achieve your goal of "{request.financial_goal}", focus on:
1. Maintaining consistent monthly savings of ₹{monthly_savings}
2. Reviewing and optimizing expense categories
3. Building an emergency fund first
4. Then directing savings toward your specific goal

**Recommendations:**
- Track all expenses carefully
- Consider automating savings
- Review your risk tolerance ({request.risk_level}) for investment decisions
- Set up a dedicated savings account for your goal"""
        
        # Generate actionable recommendations
        recommendations_prompt = f"""Based on this financial profile, provide 5-7 specific, actionable recommendations:

- Monthly Income: ₹{request.monthly_income}
- Monthly Savings: ₹{monthly_savings}
- Financial Goal: {request.financial_goal}
- Risk Level: {request.risk_level}

Provide recommendations as a numbered list, each being specific and actionable."""

        recommendations = []
        try:
            rec_messages = [{"role": "user", "content": recommendations_prompt}]
            if GROQ_API_KEY:
                rec_text = await chat_with_groq(rec_messages, use_rag=False, rag_context="")
                recommendations = [line.strip() for line in rec_text.split('\n') if line.strip() and (line.strip()[0].isdigit() or line.strip()[0] in ['-', '•'])]
                recommendations = [r for r in recommendations if len(r) > 10][:7]
            elif OLLAMA_BASE_URL:
                rec_text = await chat_with_ollama(rec_messages, use_rag=False, rag_context="")
                recommendations = [line.strip() for line in rec_text.split('\n') if line.strip() and (line.strip()[0].isdigit() or line.strip()[0] in ['-', '•'])]
                recommendations = [r for r in recommendations if len(r) > 10][:7]
        except:
            pass
        
        # Fallback recommendations
        if not recommendations:
            recommendations = [
                f"Maintain consistent monthly savings of ₹{monthly_savings}",
                f"Build an emergency fund of 3-6 months expenses (₹{total_expenses * 3} - ₹{total_expenses * 6})",
                f"Review and optimize your expense categories regularly",
                f"Set up automatic transfers to a dedicated savings account for '{request.financial_goal}'",
                f"Consider investment options matching your {request.risk_level} risk tolerance",
                f"Track your progress toward '{request.financial_goal}' monthly",
                f"Look for opportunities to increase income or reduce expenses"
            ]
        
        # Calculate goal analysis
        goal_amount = 0
        try:
            import re
            goal_match = re.search(r'₹?\s*(\d+(?:,\d+)*(?:\.\d+)?)', request.financial_goal)
            if goal_match:
                goal_amount = float(goal_match.group(1).replace(',', ''))
        except:
            pass
        
        months_to_goal = round((goal_amount / monthly_savings)) if monthly_savings > 0 and goal_amount > 0 else 0
        years_to_goal = round(months_to_goal / 12, 1) if months_to_goal > 0 else 0
        
        goal_analysis = {
            "goal": request.financial_goal,
            "estimated_cost": goal_amount if goal_amount > 0 else "Not specified",
            "current_savings": 0,
            "monthly_savings": round(monthly_savings, 2),
            "estimated_months": months_to_goal,
            "estimated_years": years_to_goal,
            "feasibility": "Achievable" if monthly_savings > 0 and (goal_amount == 0 or months_to_goal <= 120) else "Needs review"
        }
        
        return FinancialAdviceResponse(
            name=request.name,
            monthly_income=round(request.monthly_income, 2),
            monthly_expenses=round(total_expenses, 2),
            monthly_savings=round(monthly_savings, 2),
            expense_breakdown=expense_breakdown,
            financial_advice=financial_advice,
            recommendations=recommendations,
            goal_analysis=goal_analysis,
            provider=used_provider,
            created_at=datetime.now().isoformat()
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating financial advice: {str(e)}")

# ==================== Wellness Bot ====================

@app.post("/agent/wellness/support", response_model=WellnessSupportResponse)
async def get_wellness_support(request: WellnessSupportRequest):
    """
    Wellness Bot - Generate emotional support, motivational messages, and life guidance.
    
    Takes wellness scores (1-10 scale where 1=bad, 10=good) and uses AI models to provide:
    - Emotional support and encouragement
    - Motivational messages
    - Importance of life, study, and completing goals
    - Positive affirmations
    - Personalized recommendations
    """
    # Validate scores (1-10 range)
    if not (1 <= request.emotional_wellness_score <= 10):
        raise HTTPException(status_code=400, detail="emotional_wellness_score must be between 1 and 10")
    if not (1 <= request.financial_wellness_score <= 10):
        raise HTTPException(status_code=400, detail="financial_wellness_score must be between 1 and 10")
    if not (1 <= request.current_mood_score <= 10):
        raise HTTPException(status_code=400, detail="current_mood_score must be between 1 and 10")
    if not (1 <= request.stress_level <= 10):
        raise HTTPException(status_code=400, detail="stress_level must be between 1 and 10")
    
    try:
        # Determine overall wellness status
        avg_score = (request.emotional_wellness_score + request.financial_wellness_score + request.current_mood_score) / 3
        wellness_status = "excellent" if avg_score >= 8 else "good" if avg_score >= 6 else "moderate" if avg_score >= 4 else "needs_attention"
        
        # Generate each section separately with AI for unique, varied responses
        import random
        import time
        
        # Add randomness seed based on timestamp and scores for variety
        random_seed = int(time.time() * 1000) + sum([request.emotional_wellness_score, request.financial_wellness_score, request.current_mood_score, request.stress_level])
        
        wellness_context = f"""Wellness Assessment (Scale: 1=bad, 10=good):
- Emotional Wellness: {request.emotional_wellness_score}/10
- Financial Wellness: {request.financial_wellness_score}/10
- Current Mood: {request.current_mood_score}/10
- Stress Level: {request.stress_level}/10
- Concerns: {request.concerns if request.concerns else "None specified"}
- Overall Status: {wellness_status.title()}"""
        
        # Helper function to generate AI content with variety and higher temperature
        async def generate_ai_section(prompt_text: str, section_name: str) -> str:
            """Generate AI content for a specific section with variety"""
            full_prompt = f"""{wellness_context}

{prompt_text}

IMPORTANT: 
- Write in a warm, empathetic, POSITIVE, and encouraging tone
- Be genuine, specific, and personal
- Use different words and phrasing than previous responses
- Create unique, heartfelt content (not generic or placeholder text)
- Every sentence should be meaningful and uplifting
- Focus on hope, growth, and possibilities
- Be creative and vary your language each time"""
            
            messages = [{"role": "user", "content": full_prompt}]
            
            if GROQ_API_KEY:
                try:
                    # Use higher temperature for more variety in wellness responses
                    system_message = {
                        "role": "system",
                        "content": "You are a compassionate, wise wellness counselor. Provide warm, positive, and encouraging support. Always be genuine, specific, and uplifting."
                    }
                    api_messages = [system_message] + messages
                    
                    headers = {
                        "Authorization": f"Bearer {GROQ_API_KEY}",
                        "Content-Type": "application/json"
                    }
                    
                    payload = {
                        "model": GROQ_MODEL_NAME,
                        "messages": api_messages,
                        "temperature": 0.9,  # Higher temperature for more variety
                        "max_tokens": 1500
                    }
                    
                    response = requests.post(GROQ_API_ENDPOINT, headers=headers, json=payload, timeout=45)
                    response.raise_for_status()
                    data = response.json()
                    ai_response = data["choices"][0]["message"]["content"]
                    if ai_response and len(ai_response.strip()) > 50:
                        return ai_response.strip()
                except Exception as e:
                    print(f"Groq failed for {section_name}: {str(e)}")
            
            if OLLAMA_BASE_URL:
                try:
                    response = await chat_with_ollama(messages, use_rag=False, rag_context="")
                    if response and len(response.strip()) > 50:
                        return response.strip()
                except Exception as e:
                    print(f"Ollama failed for {section_name}: {str(e)}")
            
            if LOCAL_LLAMA_API_URL:
                try:
                    response = await chat_with_llama(messages, use_rag=False, rag_context="")
                    if response and len(response.strip()) > 50:
                        return response.strip()
                except Exception as e:
                    print(f"LLaMA failed for {section_name}: {str(e)}")
            
            return ""
        
        used_provider = "auto"
        
        # Generate Emotional Support
        emotional_prompt = """Write 2-3 heartfelt paragraphs providing emotional support. Acknowledge their current state with deep empathy, validate their feelings, offer comfort and reassurance. Remind them their feelings are valid and challenges are temporary. Be warm, understanding, and genuinely caring."""
        emotional_support = await generate_ai_section(emotional_prompt, "emotional_support")
        
        # Generate Motivational Message
        motivational_prompt = """Write 1-2 inspiring paragraphs that motivate and energize. Remind them of their inner strength, resilience, and potential. Use uplifting, empowering language. Encourage them to keep moving forward with hope and determination."""
        motivational_message = await generate_ai_section(motivational_prompt, "motivational")
        
        # Generate Life Importance
        life_prompt = """Write 1 meaningful paragraph about the importance and preciousness of life. Emphasize that every day is an opportunity for growth, highlight the beauty and potential in their journey, and encourage gratitude and appreciation for life. Be inspiring and positive."""
        life_importance = await generate_ai_section(life_prompt, "life")
        
        # Generate Study Importance
        study_prompt = """Write 1 motivating paragraph explaining why education and learning matter. Connect studying to personal growth and future opportunities. Motivate them to see study as an investment in themselves. Encourage persistence and dedication. Be encouraging and positive."""
        study_importance = await generate_ai_section(study_prompt, "study")
        
        # Generate Goal Importance
        goal_prompt = """Write 1 inspiring paragraph about the importance of completing goals. Explain the value of setting and achieving goals, connect goal completion to self-confidence and fulfillment. Encourage breaking goals into manageable steps. Remind them every small step counts. Be motivational and positive."""
        goal_importance = await generate_ai_section(goal_prompt, "goal")
        
        # Generate Positive Affirmations
        affirmations_prompt = """Create 5-7 unique, powerful positive affirmations tailored to their situation. Each should be:
- Short and impactful (one sentence)
- In present tense
- Specific and actionable
- Uplifting and empowering
- Different from generic affirmations

Format as a numbered or bulleted list."""
        affirmations_text = await generate_ai_section(affirmations_prompt, "affirmations")
        
        # Generate Recommendations
        recommendations_prompt = """Provide 5-7 specific, actionable wellness recommendations based on their scores. Each should be:
- Practical and immediately actionable
- Tailored to their specific wellness situation
- Positive and encouraging
- Specific (not vague)

Format as a numbered or bulleted list."""
        recommendations_text = await generate_ai_section(recommendations_prompt, "recommendations")
        
        # Generate Overall Assessment
        assessment_prompt = """Write 1 comprehensive paragraph that:
- Summarizes their current wellness state positively
- Provides hope and encouragement
- Reminds them improvement is always possible
- Ends on a positive, forward-looking note
- Is warm, genuine, and uplifting"""
        overall_assessment = await generate_ai_section(assessment_prompt, "assessment")
        
        # Determine which provider was used (check if any AI call succeeded)
        if emotional_support or motivational_message:
            if GROQ_API_KEY:
                used_provider = "groq"
            elif OLLAMA_BASE_URL:
                used_provider = "ollama"
            elif LOCAL_LLAMA_API_URL:
                used_provider = "llama"
        
        # Parse affirmations from AI response
        positive_affirmations = []
        if affirmations_text:
            lines = affirmations_text.split('\n')
            for line in lines:
                line = line.strip()
                # Remove numbering/bullets
                line = line.lstrip('0123456789.-•) ').strip()
                if line and len(line) > 15 and len(line) < 200:
                    # Ensure it's positive
                    if any(word in line.lower() for word in ['i am', 'i can', 'i will', 'i have', 'i believe', 'capable', 'strong', 'worthy', 'deserve', 'can', 'will']):
                        positive_affirmations.append(line)
        
        # Parse recommendations from AI response
        recommendations = []
        if recommendations_text:
            lines = recommendations_text.split('\n')
            for line in lines:
                line = line.strip()
                # Remove numbering/bullets
                line = line.lstrip('0123456789.-•) ').strip()
                if line and len(line) > 20 and len(line) < 300:
                    # Ensure it's actionable and positive
                    if any(word in line.lower() for word in ['practice', 'try', 'consider', 'engage', 'create', 'set', 'take', 'build', 'focus', 'develop']):
                        recommendations.append(line)
        
        # Varied fallback content based on scores (ensures different responses)
        fallback_variations = {
            'emotional': [
                f"Your emotional wellness score of {request.emotional_wellness_score}/10 shows you're on a meaningful journey of self-discovery. Every emotion you experience is valid and important. Remember, feeling deeply is a sign of being fully alive. You have incredible inner strength that will guide you through any challenge.",
                f"With an emotional wellness of {request.emotional_wellness_score}/10, you're demonstrating courage by acknowledging where you are. Your feelings matter, and taking care of your emotional health is one of the most important things you can do. Trust in your ability to heal and grow.",
                f"Your emotional wellness at {request.emotional_wellness_score}/10 reflects a person who is aware and growing. Every step you take toward emotional wellness is valuable. You deserve compassion, understanding, and support. Your journey toward emotional well-being is filled with potential."
            ],
            'motivational': [
                f"Your current mood score of {request.current_mood_score}/10 is just one moment in your beautiful journey. You possess incredible resilience and inner strength. Every day brings new possibilities for joy, growth, and positive change. Keep believing in yourself.",
                f"With a mood score of {request.current_mood_score}/10, remember that your potential is limitless. You have overcome challenges before, and you will again. Each new day is a fresh opportunity to create the life you want. Your strength is greater than any obstacle.",
                f"Your mood at {request.current_mood_score}/10 doesn't define your future—it's just where you are right now. You have the power to transform your life through small, consistent positive actions. Believe in your ability to create positive change."
            ],
            'life': [
                "Life is an extraordinary gift filled with endless possibilities. Every breath you take is an opportunity to experience something beautiful, learn something new, or make a positive impact. Your life has profound meaning and value, and the world is richer because you're in it.",
                "Each day of life is a precious opportunity to grow, love, learn, and contribute. Your existence matters deeply, and you have the power to create a meaningful, fulfilling life. Embrace the journey with gratitude and wonder.",
                "Life offers you countless moments of beauty, connection, and growth. Every experience, whether joyful or challenging, contributes to your unique story. Your life is valuable, meaningful, and full of potential waiting to be realized."
            ],
            'study': [
                "Education is a powerful catalyst for personal transformation. Every moment you invest in learning expands your mind, opens new doors, and builds the foundation for your dreams. Your commitment to studying is an investment in your brightest future.",
                "Learning is one of life's greatest adventures. Each concept you master, each skill you develop, brings you closer to your goals. Your dedication to education today shapes the opportunities you'll have tomorrow. Keep learning, keep growing.",
                "Studying is your pathway to empowerment and achievement. Every hour you spend learning builds your confidence, sharpens your mind, and prepares you for success. Your educational journey is creating the person you're meant to become."
            ],
            'goal': [
                "Achieving goals builds unshakeable confidence and creates powerful momentum in your life. Every goal you complete proves your capability and brings you closer to your dreams. Small steps lead to great achievements—keep moving forward.",
                "Completing goals transforms your potential into reality. Each achievement, no matter how small, strengthens your belief in yourself and opens new possibilities. Your dedication to finishing what you start is building an extraordinary life.",
                "Goals are the bridges between where you are and where you want to be. Every goal you complete is a victory that builds your confidence and creates positive momentum. Your commitment to achievement is shaping your future."
            ]
        }
        
        # Use score-based index to vary fallback content
        score_sum = request.emotional_wellness_score + request.financial_wellness_score + request.current_mood_score + request.stress_level
        variation_index = (score_sum + random_seed) % 3
        
        if not emotional_support:
            emotional_support = fallback_variations['emotional'][variation_index]
        
        if not motivational_message:
            motivational_message = fallback_variations['motivational'][variation_index]
        
        if not life_importance:
            life_importance = fallback_variations['life'][variation_index]
        
        if not study_importance:
            study_importance = fallback_variations['study'][variation_index]
        
        if not goal_importance:
            goal_importance = fallback_variations['goal'][variation_index]
        
        # Varied positive affirmations
        if not positive_affirmations or len(positive_affirmations) < 5:
            affirmation_sets = [
                ["I am capable of overcoming any challenge that comes my way", "My feelings are valid and I deserve compassion", "Every day I grow stronger and more resilient", "I have the power to create positive change", "I am worthy of happiness and fulfillment", "Small progress is still meaningful progress", "I believe in my ability to achieve my dreams"],
                ["I am strong enough to handle whatever life brings", "I deserve love, care, and understanding", "Each moment is an opportunity for growth", "I have inner strength that guides me forward", "I am capable of achieving great things", "Progress happens one step at a time", "I trust in my ability to overcome obstacles"],
                ["I am resilient and can bounce back from challenges", "My journey matters and I am making progress", "I have the courage to face difficult moments", "I am deserving of peace, joy, and success", "Every day I become a better version of myself", "I have the strength to pursue my goals", "I am capable of creating the life I want"]
            ]
            selected_affirmations = affirmation_sets[variation_index]
            if positive_affirmations:
                positive_affirmations.extend(selected_affirmations[:7-len(positive_affirmations)])
            else:
                positive_affirmations = selected_affirmations[:7]
        
        # Varied recommendations
        if not recommendations or len(recommendations) < 5:
            rec_sets = [
                ["Practice deep breathing for 5 minutes when feeling stressed", "Engage in one activity you enjoy each day", "Write down three things you're grateful for daily", "Connect with someone supportive this week", "Take a 10-minute walk in nature", "Practice self-compassion in difficult moments", "Set one small, achievable goal for this week"],
                ["Try meditation or mindfulness for 10 minutes daily", "Do something creative that brings you joy", "Spend quality time with people who uplift you", "Create a peaceful morning or evening routine", "Focus on one positive action you can take today", "Practice positive self-talk throughout the day", "Celebrate small wins and progress you've made"],
                ["Take regular breaks to rest and recharge", "Engage in physical activity that feels good", "Express your feelings through journaling or talking", "Practice gratitude by noticing good moments", "Set boundaries to protect your energy", "Do something kind for yourself each day", "Remember that asking for help is a sign of strength"]
            ]
            selected_recs = rec_sets[variation_index]
            if recommendations:
                recommendations.extend(selected_recs[:7-len(recommendations)])
            else:
                recommendations = selected_recs[:7]
        
        # Varied overall assessment
        if not overall_assessment:
            assessment_variations = [
                f"Your wellness journey shows you're at a {wellness_status} level, which is a meaningful starting point for growth. With emotional wellness at {request.emotional_wellness_score}/10, financial wellness at {request.financial_wellness_score}/10, and mood at {request.current_mood_score}/10, you have a clear path forward. Every positive choice you make builds toward greater well-being. You have the power to improve, and I believe in your ability to create positive change.",
                f"Currently at a {wellness_status} wellness level, you're on a path of self-discovery and improvement. Your scores—emotional {request.emotional_wellness_score}/10, financial {request.financial_wellness_score}/10, mood {request.current_mood_score}/10—reflect where you are now, not where you'll always be. Wellness is a journey of small, consistent steps. You're capable of remarkable growth and positive transformation.",
                f"At a {wellness_status} wellness level, you're taking important steps toward better well-being. Your emotional wellness ({request.emotional_wellness_score}/10), financial wellness ({request.financial_wellness_score}/10), and mood ({request.current_mood_score}/10) are areas of opportunity for growth. Remember, every positive action matters. You have the strength and ability to improve your wellness and create a more fulfilling life."
            ]
            overall_assessment = assessment_variations[variation_index]
        
        return WellnessSupportResponse(
            emotional_support=emotional_support,
            motivational_message=motivational_message,
            life_importance=life_importance,
            study_importance=study_importance,
            goal_importance=goal_importance,
            positive_affirmations=positive_affirmations[:7],
            recommendations=recommendations[:7],
            overall_assessment=overall_assessment,
            provider=used_provider,
            created_at=datetime.now().isoformat()
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating wellness support: {str(e)}")

if __name__ == "__main__":
    uvicorn.run("main:app", host=HOST, port=PORT, reload=RELOAD)
