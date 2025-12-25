from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List, Tuple, Dict
import uvicorn
import os
import requests
import json
import io
import base64
from dotenv import load_dotenv
from datetime import datetime, timedelta
import uuid
from collections import defaultdict
import hashlib

# Try to import optional libraries
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from PIL import Image
    import pytesseract
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_SUPPORT = True
except ImportError:
    PDF2IMAGE_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# Import PDF Summarizer (from TextSummarizer folder - using transformers)
try:
    from pdf_summarizer import PDFSummarizer
    PDF_SUMMARIZER_SUPPORT = True
except ImportError:
    PDF_SUMMARIZER_SUPPORT = False
    print("PDF Summarizer not available. Install: pip install transformers torch sentencepiece")

# Import DOC Summarizer (from TextSummarizer folder - using transformers)
try:
    from doc_summarizer import DOCSummarizer
    DOC_SUMMARIZER_SUPPORT = True
except ImportError:
    DOC_SUMMARIZER_SUPPORT = False
    print("DOC Summarizer not available. Install: pip install transformers torch sentencepiece")

# Load environment variables
load_dotenv()

# Configuration
API_TITLE = os.getenv("API_TITLE", "Gurukul Backend API")
HOST = os.getenv("HOST", "0.0.0.0")
PORT = int(os.getenv("PORT", "3000"))
RELOAD = os.getenv("RELOAD", "True").lower() == "true"

# API Keys
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_API_ENDPOINT = os.getenv("GROQ_API_ENDPOINT", "https://api.groq.com/openai/v1/chat/completions")
GROQ_MODEL_NAME = os.getenv("GROQ_MODEL_NAME", "llama-3.3-70b-versatile")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL_PRIMARY = os.getenv("OLLAMA_MODEL_PRIMARY", "llama2")
LOCAL_LLAMA_API_URL = os.getenv("LOCAL_LLAMA_API_URL", "http://localhost:8080/v1/chat/completions")

YOUTUBE_API_KEY = os.getenv("YOUTUBE_API_KEY")
YOUTUBE_API_BASE_URL = "https://www.googleapis.com/youtube/v3/search"

# Initialize FastAPI
app = FastAPI(title=API_TITLE)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Conversation Memory Storage (In-memory, can be upgraded to database)
# Structure: {conversation_id: {"messages": [...], "created_at": "...", "updated_at": "..."}}
conversation_store: Dict[str, Dict] = {}

# RAG Knowledge Store (Simple vector store for context retrieval)
# Structure: {hash: {"text": "...", "metadata": {...}}}
rag_knowledge_store: Dict[str, Dict] = {}

# Saved Summaries Store
# Structure: {summary_id: {"title": "...", "content": "...", "source": "...", "created_at": "...", "questions": [...]}}
saved_summaries_store: Dict[str, Dict] = {}

# Questions Store (linked to summaries)
# Structure: {question_id: {"summary_id": "...", "type": "...", "question": "...", "answer": "...", "options": [...], "review_schedule": {...}}}
questions_store: Dict[str, Dict] = {}

# Quiz Store
# Structure: {quiz_id: {"subject": "...", "topic": "...", "questions": [...], "created_at": "...", "provider": "..."}}
quiz_store: Dict[str, Dict] = {}

# Agent Simulator Stores
# Financial Profiles Store
# Structure: {profile_id: {"name": "...", "income": ..., "expenses": ..., "goal": "...", "created_at": "..."}}
financial_profiles_store: Dict[str, Dict] = {}

# Wellness Profiles Store
# Structure: {profile_id: {"emotional_score": ..., "financial_score": ..., "mood": ..., "stress": ..., "created_at": "..."}}
wellness_profiles_store: Dict[str, Dict] = {}

# Pydantic Models
class YouTubeVideo(BaseModel):
    title: str
    video_id: str
    url: str
    thumbnail: str
    channel_title: str
    duration: Optional[str] = None
    view_count: Optional[str] = None

class SubjectExplorerRequest(BaseModel):
    subject: str
    topic: str
    provider: Optional[str] = "groq"  # groq, ollama, llama
    use_knowledge_store: Optional[bool] = False

class SubjectExplorerResponse(BaseModel):
    subject: str
    topic: str
    notes: str
    provider: str
    youtube_recommendations: List[YouTubeVideo]
    success: bool

class SummarizerRequest(BaseModel):
    text: Optional[str] = None
    provider: Optional[str] = "groq"  # groq, ollama, llama, textrank
    summary_type: Optional[str] = "concise"  # concise, detailed, bullet_points

class SummarizerResponse(BaseModel):
    original_length: int
    summary: str
    summary_length: int
    provider: str
    summary_type: str
    success: bool

class PDFPageSummary(BaseModel):
    page_number: int
    summary: str
    original_length: int
    summary_length: int

class PDFSummarizerResponse(BaseModel):
    total_pages: int
    pages_summarized: int
    page_summaries: List[PDFPageSummary]
    overall_summary: str
    summary_type: str
    provider: str
    success: bool

class DOCSectionSummary(BaseModel):
    section_number: int
    summary: str
    original_length: int
    summary_length: int

class DOCSummarizerResponse(BaseModel):
    total_sections: int
    sections_summarized: int
    section_summaries: List[DOCSectionSummary]
    overall_summary: str
    summary_type: str
    provider: str
    success: bool

# Chatbot Models
class ChatMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str
    timestamp: Optional[str] = None

class ChatRequest(BaseModel):
    message: str
    conversation_id: Optional[str] = None  # If None, creates new conversation
    provider: Optional[str] = "auto"  # auto, groq, ollama, llama
    use_rag: Optional[bool] = True  # Enable RAG for context retrieval
    max_history: Optional[int] = 10  # Maximum number of previous messages to include

class ChatResponse(BaseModel):
    response: str
    conversation_id: str
    provider: str
    message_count: int
    timestamp: str
    success: bool

class ConversationHistory(BaseModel):
    conversation_id: str
    messages: List[ChatMessage]
    created_at: str
    updated_at: str
    message_count: int

# Saved Summarizer Models
class SaveSummaryRequest(BaseModel):
    title: str
    content: str  # Summary content
    source: Optional[str] = None  # Original file name or source
    source_type: Optional[str] = None  # "pdf", "doc", "text", etc.
    metadata: Optional[Dict] = None

class SavedSummary(BaseModel):
    summary_id: str
    title: str
    content: str
    source: Optional[str]
    source_type: Optional[str]
    created_at: str
    updated_at: str
    question_count: int
    metadata: Optional[Dict]

class QuestionGenerationRequest(BaseModel):
    summary_id: str
    question_types: Optional[List[str]] = ["qa", "mcq", "flashcard", "case_based"]  # Types to generate
    num_questions: Optional[int] = 5  # Number of questions per type
    difficulty: Optional[str] = "medium"  # easy, medium, hard

class MCQOption(BaseModel):
    option: str
    is_correct: bool

class Question(BaseModel):
    question_id: str
    summary_id: str
    question_type: str  # "qa", "mcq", "flashcard", "case_based"
    question: str
    answer: str
    options: Optional[List[MCQOption]] = None  # For MCQs
    explanation: Optional[str] = None
    difficulty: str
    created_at: str

class QuestionGenerationResponse(BaseModel):
    summary_id: str
    questions: List[Question]
    total_questions: int
    success: bool

class ReviewSchedule(BaseModel):
    question_id: str
    summary_id: str
    first_review: str  # Same day
    second_review: str  # Next day
    third_review: str  # 1 week later
    last_review_date: Optional[str] = None
    next_review_date: str
    review_count: int = 0
    mastery_level: str = "learning"  # learning, reviewing, mastered

class ReviewRequest(BaseModel):
    question_id: str
    user_answer: Optional[str] = None
    is_correct: bool
    confidence: Optional[str] = "medium"  # low, medium, high

class ReviewResponse(BaseModel):
    question_id: str
    correct_answer: str
    is_correct: bool
    next_review_date: str
    review_count: int
    mastery_level: str
    message: str

class PendingReview(BaseModel):
    question_id: str
    summary_id: str
    question: str
    question_type: str
    answer: str
    next_review_date: str
    days_until_review: int

# Quiz/Test Models
class QuizRequest(BaseModel):
    subject: str
    topic: str
    provider: Optional[str] = "auto"  # auto, groq, ollama, llama
    difficulty: Optional[str] = "medium"  # easy, medium, hard

class QuizQuestion(BaseModel):
    question_id: str
    question_number: int
    question: str
    options: List[str]  # For MCQs
    question_type: str  # "mcq" or "qa"

class QuizResponse(BaseModel):
    quiz_id: str
    subject: str
    topic: str
    questions: List[QuizQuestion]
    total_questions: int
    created_at: str
    provider: str

class QuizSubmission(BaseModel):
    quiz_id: str
    answers: Dict[str, str]  # {question_id: user_answer}

class QuizResult(BaseModel):
    quiz_id: str
    total_questions: int
    correct_answers: int
    wrong_answers: int
    score_percentage: float
    results: List[Dict]  # Detailed results for each question
    submitted_at: str

# Agent Simulator Models
# EduMentor Models
class EduMentorRequest(BaseModel):
    subject: str
    topic: str
    include_wikipedia: Optional[bool] = False
    use_knowledge_store: Optional[bool] = False
    use_orchestration: Optional[bool] = False
    provider: Optional[str] = "auto"  # auto, groq, ollama, llama

class EduMentorResponse(BaseModel):
    subject: str
    topic: str
    lesson_content: str
    wikipedia_sources: Optional[List[Dict]] = None
    knowledge_store_used: bool
    orchestration_used: bool
    provider: str
    created_at: str
    success: bool

# Financial Agent Models
class ExpenseCategory(BaseModel):
    name: str
    amount: float

class FinancialProfileRequest(BaseModel):
    name: str
    monthly_income: float
    monthly_expenses: float
    expense_categories: List[ExpenseCategory]
    financial_goal: str
    financial_type: str  # Conservative, Moderate, Aggressive
    risk_level: str  # Low, Moderate, High

class FinancialSimulationRequest(BaseModel):
    profile_id: str
    months_to_simulate: Optional[int] = 12

class FinancialSimulationResponse(BaseModel):
    profile_id: str
    simulation_period: int  # months
    initial_balance: float
    final_balance: float
    total_savings: float
    monthly_breakdown: List[Dict]
    recommendations: List[str]
    goal_progress: Dict
    provider: str
    created_at: str

# Wellness Bot Models
class WellnessProfileRequest(BaseModel):
    emotional_wellness_score: Optional[int] = 5  # 1-10
    financial_wellness_score: Optional[int] = 5  # 1-10
    current_mood_score: Optional[int] = 5  # 1-10
    stress_level: Optional[int] = 3  # 1-10
    concerns: Optional[str] = None

class WellnessSessionRequest(BaseModel):
    profile_id: str
    message: str
    session_type: Optional[str] = "general"  # general, emotional, financial

class WellnessSessionResponse(BaseModel):
    profile_id: str
    session_id: str
    response: str
    mood_analysis: Optional[Dict] = None
    stress_analysis: Optional[Dict] = None
    recommendations: List[str]
    provider: str
    timestamp: str

# Helper function to create teacher-like prompt
def create_teaching_prompt(subject: str, topic: str) -> str:
    return f"""You are an expert teacher specializing in {subject}. Your task is to teach the topic "{topic}" to a student.

Please provide comprehensive, clear, and engaging educational notes that include:

1. **Introduction**: A brief overview of the topic and why it's important
2. **Key Concepts**: Break down the main concepts in simple, understandable terms
3. **Detailed Explanation**: Provide a thorough explanation with examples
4. **Real-world Applications**: Show how this topic applies in real life
5. **Summary**: A concise recap of the key points
6. **Practice Questions**: 2-3 questions to help reinforce learning

Write in a friendly, encouraging tone as if you're speaking directly to a student. Use simple language, analogies, and examples to make complex concepts easy to understand. Structure the content with clear headings and bullet points for better readability.

Subject: {subject}
Topic: {topic}

Now, provide the educational notes:"""

# Groq API call
async def call_groq_api(subject: str, topic: str) -> str:
    if not GROQ_API_KEY:
        raise HTTPException(status_code=500, detail="Groq API key not configured")
    
    prompt = create_teaching_prompt(subject, topic)
    
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": GROQ_MODEL_NAME,
        "messages": [
            {
                "role": "system",
                "content": "You are an expert teacher who explains concepts clearly and simply."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.7,
        "max_tokens": 2000
    }
    
    try:
        response = requests.post(GROQ_API_ENDPOINT, headers=headers, json=payload, timeout=30)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Groq API error: {str(e)}")

# Ollama API call
async def call_ollama_api(subject: str, topic: str) -> str:
    prompt = create_teaching_prompt(subject, topic)
    
    url = f"{OLLAMA_BASE_URL}/api/generate"
    
    payload = {
        "model": OLLAMA_MODEL_PRIMARY,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.7,
            "num_predict": 2000
        }
    }
    
    try:
        response = requests.post(url, json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
        return data.get("response", "Error generating response from Ollama")
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Ollama service not available. Make sure Ollama is running on localhost:11434")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ollama API error: {str(e)}")

# Local LLaMA API call
async def call_llama_api(subject: str, topic: str) -> str:
    prompt = create_teaching_prompt(subject, topic)
    
    headers = {
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "llama3.1",
        "messages": [
            {
                "role": "system",
                "content": "You are an expert teacher who explains concepts clearly and simply."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.7,
        "max_tokens": 2000
    }
    
    try:
        response = requests.post(LOCAL_LLAMA_API_URL, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Local LLaMA service not available. Make sure the local LLaMA server is running")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLaMA API error: {str(e)}")

# YouTube API call to get recommendations
async def get_youtube_recommendations(subject: str, topic: str, max_results: int = 5) -> List[YouTubeVideo]:
    """Get YouTube video recommendations based on subject and topic"""
    if not YOUTUBE_API_KEY:
        return []  # Return empty list if API key is not configured
    
    # Create search query
    search_query = f"{subject} {topic} tutorial explanation"
    
    params = {
        "part": "snippet",
        "q": search_query,
        "type": "video",
        "maxResults": max_results,
        "key": YOUTUBE_API_KEY,
        "order": "relevance",
        "videoCategoryId": "27",  # Education category
        "safeSearch": "strict"
    }
    
    try:
        response = requests.get(YOUTUBE_API_BASE_URL, params=params, timeout=10)
        response.raise_for_status()
        data = response.json()
        
        videos = []
        for item in data.get("items", []):
            video_id = item["id"]["videoId"]
            snippet = item["snippet"]
            
            video = YouTubeVideo(
                title=snippet["title"],
                video_id=video_id,
                url=f"https://www.youtube.com/watch?v={video_id}",
                thumbnail=snippet["thumbnails"]["high"]["url"],
                channel_title=snippet["channelTitle"]
            )
            videos.append(video)
        
        return videos
    except requests.exceptions.RequestException as e:
        # If YouTube API fails, return empty list (don't break the main functionality)
        print(f"YouTube API error: {str(e)}")
        return []
    except Exception as e:
        print(f"Error fetching YouTube recommendations: {str(e)}")
        return []

# Text extraction functions
async def extract_pages_from_pdf(file: UploadFile) -> List[str]:
    """
    Extract text from PDF file with multi-page support.
    Returns list of page texts (one per page).
    Tries text extraction first, falls back to OCR for scanned PDFs.
    """
    if not PDF_SUPPORT:
        raise HTTPException(status_code=400, detail="PDF support not available. Install PyPDF2: pip install PyPDF2")
    
    try:
        content = await file.read()
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        print(f"[PDF] Processing PDF with {len(pdf_reader.pages)} pages...")
        
        # Try text extraction first
        text_pages = []
        for page_num, page in enumerate(pdf_reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text and len(page_text.strip()) > 10:
                    text_pages.append(page_text)
                    print(f"[PDF] Page {page_num}: Extracted {len(page_text)} characters")
                else:
                    print(f"[PDF] Page {page_num}: No text found, will use OCR")
                    text_pages.append(None)  # Mark for OCR
            except Exception as e:
                print(f"[PDF] Page {page_num}: Text extraction failed: {str(e)}, will use OCR")
                text_pages.append(None)
        
        # Check if we need OCR (scanned PDFs)
        needs_ocr = any(page is None for page in text_pages)
        
        if needs_ocr and OCR_SUPPORT and PDF2IMAGE_SUPPORT:
            print("[PDF] Using OCR for scanned PDF pages...")
            # Convert PDF pages to images and use OCR
            images = convert_from_bytes(content, dpi=300)
            
            ocr_text_pages = []
            for page_num, image in enumerate(images, 1):
                try:
                    page_text = pytesseract.image_to_string(image)
                    if page_text and len(page_text.strip()) > 10:
                        ocr_text_pages.append(page_text)
                        print(f"[PDF] Page {page_num} (OCR): Extracted {len(page_text)} characters")
                    else:
                        ocr_text_pages.append("")
                except Exception as e:
                    print(f"[PDF] Page {page_num} (OCR): Failed - {str(e)}")
                    ocr_text_pages.append("")
            
            # Combine OCR results with text extraction results
            final_pages = []
            for i, (text_page, ocr_page) in enumerate(zip(text_pages, ocr_text_pages), 1):
                if text_page:
                    final_pages.append(text_page)
                elif ocr_page:
                    final_pages.append(ocr_page)
                else:
                    final_pages.append("")  # Empty page
            
            print(f"[PDF] Total extracted: {len(final_pages)} pages")
            return final_pages
        
        elif needs_ocr:
            raise HTTPException(
                status_code=400, 
                detail="PDF appears to be scanned (image-based). OCR support required. Install: pip install pdf2image pytesseract Pillow"
            )
        else:
            # All pages have text, return as list
            print(f"[PDF] Total extracted: {len(text_pages)} pages")
            return text_pages
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from PDF: {str(e)}")

async def extract_text_from_pdf(file: UploadFile) -> str:
    """
    Extract text from PDF file with multi-page support.
    Tries text extraction first, falls back to OCR for scanned PDFs.
    """
    if not PDF_SUPPORT:
        raise HTTPException(status_code=400, detail="PDF support not available. Install PyPDF2: pip install PyPDF2")
    
    try:
        content = await file.read()
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        print(f"[PDF] Processing PDF with {len(pdf_reader.pages)} pages...")
        
        # Try text extraction first
        text_pages = []
        for page_num, page in enumerate(pdf_reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text and len(page_text.strip()) > 10:
                    text_pages.append(page_text)
                    print(f"[PDF] Page {page_num}: Extracted {len(page_text)} characters")
                else:
                    print(f"[PDF] Page {page_num}: No text found, will use OCR")
                    text_pages.append(None)  # Mark for OCR
            except Exception as e:
                print(f"[PDF] Page {page_num}: Text extraction failed: {str(e)}, will use OCR")
                text_pages.append(None)
        
        # Check if we need OCR (scanned PDFs)
        needs_ocr = any(page is None for page in text_pages)
        
        if needs_ocr and OCR_SUPPORT and PDF2IMAGE_SUPPORT:
            print("[PDF] Using OCR for scanned PDF pages...")
            # Convert PDF pages to images and use OCR
            images = convert_from_bytes(content, dpi=300)
            
            ocr_text_pages = []
            for page_num, image in enumerate(images, 1):
                try:
                    page_text = pytesseract.image_to_string(image)
                    if page_text and len(page_text.strip()) > 10:
                        ocr_text_pages.append(page_text)
                        print(f"[PDF] Page {page_num} (OCR): Extracted {len(page_text)} characters")
                    else:
                        ocr_text_pages.append("")
                except Exception as e:
                    print(f"[PDF] Page {page_num} (OCR): Failed - {str(e)}")
                    ocr_text_pages.append("")
            
            # Combine OCR results with text extraction results
            final_text = ""
            for i, (text_page, ocr_page) in enumerate(zip(text_pages, ocr_text_pages), 1):
                if text_page:
                    final_text += f"\n[Page {i}]\n{text_page}\n"
                elif ocr_page:
                    final_text += f"\n[Page {i} - OCR]\n{ocr_page}\n"
            
            result = final_text.strip()
            print(f"[PDF] Total extracted text: {len(result)} characters from {len(images)} pages")
            return result
        
        elif needs_ocr:
            raise HTTPException(
                status_code=400, 
                detail="PDF appears to be scanned (image-based). OCR support required. Install: pip install pdf2image pytesseract Pillow"
            )
        else:
            # All pages have text, combine them
            final_text = "\n\n".join([f"[Page {i+1}]\n{page_text}" for i, page_text in enumerate(text_pages) if page_text])
            print(f"[PDF] Total extracted text: {len(final_text)} characters from {len(text_pages)} pages")
            return final_text.strip()
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from PDF: {str(e)}")

async def extract_text_from_image(file: UploadFile) -> str:
    """Extract text from image using OCR"""
    if not OCR_SUPPORT:
        raise HTTPException(status_code=400, detail="OCR support not available. Install pytesseract and Pillow")
    
    try:
        content = await file.read()
        image = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from image: {str(e)}")

async def extract_sections_from_docx(file: UploadFile) -> List[str]:
    """Extract text from DOCX file with multi-section support. Returns list of sections."""
    if not DOCX_SUPPORT:
        raise HTTPException(status_code=400, detail="DOCX support not available. Install: pip install python-docx")
    
    try:
        content = await file.read()
        doc_file = io.BytesIO(content)
        doc = Document(doc_file)
        
        print(f"[DOCX] Processing DOCX with {len(doc.paragraphs)} paragraphs...")
        
        # Extract text from all paragraphs as separate sections
        sections = []
        current_section = []
        
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                # Check if this looks like a heading (bold, larger font, etc.)
                is_heading = False
                if para.runs:
                    for run in para.runs:
                        if run.bold or (run.font and run.font.size and run.font.size.pt > 12):
                            is_heading = True
                            break
                
                if is_heading and current_section:
                    # Start new section
                    sections.append("\n".join(current_section))
                    current_section = [para_text]
                else:
                    current_section.append(para_text)
        
        # Add last section
        if current_section:
            sections.append("\n".join(current_section))
        
        # Also extract text from tables as separate sections
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    table_text.append(row_text)
            if table_text:
                sections.append("\n".join(table_text))
        
        # If no sections found, combine all as one
        if not sections:
            all_text = "\n\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
            if all_text:
                sections = [all_text]
        
        print(f"[DOCX] Extracted {len(sections)} sections")
        return sections
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from DOCX: {str(e)}")

async def extract_text_from_docx(file: UploadFile) -> str:
    """Extract text from DOCX file with multi-page support (legacy function)"""
    sections = await extract_sections_from_docx(file)
    return "\n\n".join(sections)

async def improve_grammar_and_clarity(text: str, provider: str = "groq") -> str:
    """
    Improve grammar, fix errors, and enhance clarity of text using AI models.
    Uses Groq by default, with fallback to other models.
    """
    if not text or len(text.strip()) < 10:
        return text
    
    prompt = f"""Please improve the following text by:
1. Fixing all grammar errors
2. Correcting spelling mistakes
3. Improving sentence structure and clarity
4. Ensuring proper punctuation
5. Making sentences more meaningful and coherent
6. Maintaining the original meaning and context

Text to improve:
{text}

Improved text:"""
    
    # Try Groq first
    if provider == "groq" and GROQ_API_KEY:
        try:
            headers = {
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": GROQ_MODEL_NAME,
                "messages": [
                    {
                        "role": "system",
                        "content": "You are an expert editor who improves text quality, fixes grammar errors, and enhances clarity while maintaining the original meaning."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "temperature": 0.3,
                "max_tokens": 2000
            }
            
            response = requests.post(GROQ_API_ENDPOINT, headers=headers, json=payload, timeout=60)
            response.raise_for_status()
            data = response.json()
            improved_text = data["choices"][0]["message"]["content"]
            print(f"[Grammar] Text improved using Groq. Original: {len(text)} chars, Improved: {len(improved_text)} chars")
            return improved_text
        except Exception as e:
            print(f"[Grammar] Groq failed: {str(e)}, trying fallback...")
    
    # Try Ollama as fallback
    try:
        url = f"{OLLAMA_BASE_URL}/api/generate"
        payload = {
            "model": OLLAMA_MODEL_PRIMARY,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": 0.3,
                "num_predict": 2000
            }
        }
        response = requests.post(url, json=payload, timeout=120)
        response.raise_for_status()
        data = response.json()
        improved_text = data.get("response", text)
        print(f"[Grammar] Text improved using Ollama. Original: {len(text)} chars, Improved: {len(improved_text)} chars")
        return improved_text
    except Exception as e:
        print(f"[Grammar] Ollama failed: {str(e)}, trying LLaMA...")
    
    # Try Local LLaMA as last fallback
    try:
        headers = {"Content-Type": "application/json"}
        payload = {
            "model": "llama3.1",
            "messages": [
                {
                    "role": "system",
                    "content": "You are an expert editor who improves text quality, fixes grammar errors, and enhances clarity while maintaining the original meaning."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.3,
            "max_tokens": 2000
        }
        response = requests.post(LOCAL_LLAMA_API_URL, headers=headers, json=payload, timeout=120)
        response.raise_for_status()
        data = response.json()
        improved_text = data["choices"][0]["message"]["content"]
        print(f"[Grammar] Text improved using LLaMA. Original: {len(text)} chars, Improved: {len(improved_text)} chars")
        return improved_text
    except Exception as e:
        print(f"[Grammar] All grammar improvement methods failed: {str(e)}, returning original text")
        return text

def create_summarization_prompt(text: str, summary_type: str) -> str:
    """Create prompt for summarization based on type"""
    base_prompt = f"""Please summarize the following text. The text is:"""
    
    if summary_type == "concise":
        instruction = "Provide a brief, concise summary (2-3 sentences) capturing the main points."
    elif summary_type == "detailed":
        instruction = "Provide a comprehensive, detailed summary covering all important aspects and key points."
    elif summary_type == "bullet_points":
        instruction = "Provide a summary in bullet point format, highlighting the main ideas and key information."
    else:
        instruction = "Provide a clear and comprehensive summary."
    
    return f"""{base_prompt}

{instruction}

Text to summarize:
{text}

Summary:"""

# Summarization functions using different providers
async def summarize_with_groq(text: str, summary_type: str) -> str:
    """Summarize text using Groq API"""
    if not GROQ_API_KEY:
        raise HTTPException(status_code=500, detail="Groq API key not configured")
    
    prompt = create_summarization_prompt(text, summary_type)
    
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": GROQ_MODEL_NAME,
        "messages": [
            {
                "role": "system",
                "content": "You are an expert at summarizing documents and extracting key information."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.3,
        "max_tokens": 1000
    }
    
    try:
        response = requests.post(GROQ_API_ENDPOINT, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Groq API error: {str(e)}")

async def summarize_with_ollama(text: str, summary_type: str) -> str:
    """Summarize text using Ollama API"""
    prompt = create_summarization_prompt(text, summary_type)
    
    url = f"{OLLAMA_BASE_URL}/api/generate"
    
    payload = {
        "model": OLLAMA_MODEL_PRIMARY,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.3,
            "num_predict": 1000
        }
    }
    
    try:
        response = requests.post(url, json=payload, timeout=120)
        response.raise_for_status()
        data = response.json()
        return data.get("response", "Error generating summary from Ollama")
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Ollama service not available. Make sure Ollama is running on localhost:11434")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ollama API error: {str(e)}")

async def summarize_with_llama(text: str, summary_type: str) -> str:
    """Summarize text using Local LLaMA API"""
    prompt = create_summarization_prompt(text, summary_type)
    
    headers = {
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "llama3.1",
        "messages": [
            {
                "role": "system",
                "content": "You are an expert at summarizing documents and extracting key information."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.3,
        "max_tokens": 1000
    }
    
    try:
        response = requests.post(LOCAL_LLAMA_API_URL, headers=headers, json=payload, timeout=120)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Local LLaMA service not available. Make sure the local LLaMA server is running")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLaMA API error: {str(e)}")

async def summarize_with_bart(text: str, summary_type: str) -> str:
    """Summarize text using BART model (from TextSummarizer folder - transformers)"""
    if not PDF_SUMMARIZER_SUPPORT:
        raise HTTPException(status_code=500, detail="BART summarizer not available. Install: pip install transformers torch sentencepiece")
    
    try:
        print(f"[BART] Starting summarization. Text length: {len(text)}, Type: {summary_type}")
        from pdf_summarizer import PDFSummarizer
        summarizer = PDFSummarizer()
        
        # Check if text contains multiple pages (marked with [Page X])
        if "[Page" in text or "\n[Page" in text:
            print("[BART] Detected multi-page document, using multi-page summarization...")
            # Split by page markers
            pages = []
            current_page = ""
            for line in text.split('\n'):
                if line.strip().startswith('[Page'):
                    if current_page:
                        pages.append(current_page.strip())
                    current_page = ""
                else:
                    current_page += line + "\n"
            if current_page:
                pages.append(current_page.strip())
            
            print(f"[BART] Processing {len(pages)} pages...")
            
            # Determine summary length based on type
            if summary_type == "concise":
                max_length = 80
                min_length = 30
            elif summary_type == "detailed":
                max_length = 250
                min_length = 100
            else:
                max_length = 150
                min_length = 40
            
            # Use multi-page summarization
            summary = summarizer.multi_page_summarize(pages, max_length=max_length, min_length=min_length)
        else:
            # Single document summarization
            if summary_type == "concise":
                max_length = 80
                min_length = 30
            elif summary_type == "detailed":
                max_length = 250
                min_length = 100
            else:
                max_length = 150
                min_length = 40
            
            print(f"[BART] Using max_length: {max_length}, min_length: {min_length}")
            summary = summarizer.summarize(text, max_length=max_length, min_length=min_length)
        
        print(f"[BART] Summary generated. Length: {len(summary)}")
        
        # Format based on summary type
        if summary_type == "bullet_points":
            sentences = summary.split('. ')
            summary = '\n• ' + '\n• '.join([s.strip() for s in sentences if s.strip()])
        
        if not summary or len(summary.strip()) == 0:
            raise ValueError("BART returned empty summary")
        
        print(f"[BART] Success! Final summary length: {len(summary)}")
        return summary
    except Exception as e:
        print(f"[BART] ERROR: {str(e)}")
        raise HTTPException(status_code=500, detail=f"BART summarization error: {str(e)}")

async def summarize_with_fallback(text: str, summary_type: str, preferred_provider: str = None) -> Tuple[str, str]:
    """
    Summarize text with fallback logic:
    1. Try BART first (from TextSummarizer folder - transformers) - ALWAYS TRY FIRST
    2. If BART fails or not available, try Groq
    3. If Groq fails, try Ollama
    4. If Ollama fails, try LLaMA
    
    Returns: (summary, provider_used)
    """
    # Priority 1: ALWAYS Try BART first (from TextSummarizer folder - transformers)
    if PDF_SUMMARIZER_SUPPORT:
        try:
            print("Attempting BART summarization (first priority - from TextSummarizer)...")
            summary = await summarize_with_bart(text, summary_type)
            if summary and len(summary.strip()) > 0:
                print("BART summarization successful!")
                return summary, "bart"
            else:
                print("BART returned empty summary, trying fallback...")
        except Exception as e:
            print(f"BART failed: {str(e)}, trying fallback...")
    else:
        print("BART not available, trying fallback...")
    
    # Priority 2: Try Groq (only if TextRank failed)
    if GROQ_API_KEY:
        try:
            print("Attempting Groq summarization (fallback)...")
            summary = await summarize_with_groq(text, summary_type)
            return summary, "groq"
        except Exception as e:
            print(f"Groq failed, trying fallback: {str(e)}")
    
    # Priority 3: Try Ollama
    try:
        print("Attempting Ollama summarization (fallback)...")
        summary = await summarize_with_ollama(text, summary_type)
        return summary, "ollama"
    except Exception as e:
        print(f"Ollama failed, trying fallback: {str(e)}")
    
    # Priority 4: Try LLaMA
    try:
        print("Attempting LLaMA summarization (fallback)...")
        summary = await summarize_with_llama(text, summary_type)
        return summary, "llama"
    except Exception as e:
        raise HTTPException(
            status_code=500, 
            detail=f"All summarization methods failed. BART (first priority) and all fallbacks failed. Last error: {str(e)}"
        )

# Root endpoint
@app.get("/")
async def root():
    return {
        "message": "Welcome to Gurukul Backend API",
        "status": "running"
    }

# Health check
@app.get("/health")
async def health_check():
    return {"status": "healthy"}

# Test BART endpoint
@app.get("/test-bart")
async def test_bart():
    """Test endpoint to verify BART summarizer is working"""
    test_text = """
    Artificial intelligence is transforming the world. Machine learning algorithms can learn from data.
    Deep learning uses neural networks with multiple layers. Natural language processing helps computers understand text.
    Computer vision enables machines to see and interpret images. AI is being used in healthcare, finance, and education.
    The future of AI looks promising with continued research and development.
    """
    
    try:
        if not PDF_SUMMARIZER_SUPPORT:
            return {
                "status": "error",
                "message": "BART not available. Install: pip install transformers torch sentencepiece",
                "bart_support": False
            }
        
        from pdf_summarizer import PDFSummarizer
        summarizer = PDFSummarizer()
        summary = summarizer.summarize(test_text, max_length=100, min_length=30)
        
        return {
            "status": "success",
            "message": "BART is working!",
            "bart_support": True,
            "original_length": len(test_text),
            "summary": summary,
            "summary_length": len(summary),
            "test_passed": len(summary) > 0 and len(summary) < len(test_text)
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"BART test failed: {str(e)}",
            "bart_support": PDF_SUMMARIZER_SUPPORT,
            "error": str(e)
        }

# Subject Explorer endpoint
@app.post("/subject-explorer", response_model=SubjectExplorerResponse)
async def subject_explorer(request: SubjectExplorerRequest):
    """
    Generate educational notes for a subject and topic using AI models.
    Supports Groq, Ollama, and Local LLaMA providers.
    Also provides YouTube video recommendations related to the subject and topic.
    """
    if not request.subject or not request.topic:
        raise HTTPException(status_code=400, detail="Subject and Topic are required")
    
    provider = request.provider.lower()
    
    try:
        # Generate notes
        if provider == "groq":
            notes = await call_groq_api(request.subject, request.topic)
        elif provider == "ollama":
            notes = await call_ollama_api(request.subject, request.topic)
        elif provider == "llama":
            notes = await call_llama_api(request.subject, request.topic)
        else:
            raise HTTPException(status_code=400, detail="Invalid provider. Use 'groq', 'ollama', or 'llama'")
        
        # Get YouTube recommendations
        youtube_videos = await get_youtube_recommendations(request.subject, request.topic)
        
        return SubjectExplorerResponse(
            subject=request.subject,
            topic=request.topic,
            notes=notes,
            provider=provider,
            youtube_recommendations=youtube_videos,
            success=True
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating notes: {str(e)}")

# Summarizer endpoint - Text input
@app.post("/summarize", response_model=SummarizerResponse)
async def summarize_text(request: SummarizerRequest):
    """
    Summarize text input using AI models with fallback logic.
    Priority: BART (from TextSummarizer folder - transformers) -> Groq -> Ollama -> LLaMA
    """
    if not request.text:
        raise HTTPException(status_code=400, detail="Text is required")
    
    if len(request.text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Text must be at least 50 characters long")
    
    provider = request.provider.lower() if request.provider else None
    summary_type = request.summary_type.lower()
    
    try:
        # If specific provider is requested, use it directly
        if provider == "groq":
            summary = await summarize_with_groq(request.text, summary_type)
            used_provider = "groq"
        elif provider == "ollama":
            summary = await summarize_with_ollama(request.text, summary_type)
            used_provider = "ollama"
        elif provider == "llama":
            summary = await summarize_with_llama(request.text, summary_type)
            used_provider = "llama"
        elif provider == "bart":
            summary = await summarize_with_bart(request.text, summary_type)
            used_provider = "bart"
        elif provider is None or provider == "auto":
            # Use fallback logic: BART first, then others
            summary, used_provider = await summarize_with_fallback(request.text, summary_type)
        else:
            raise HTTPException(status_code=400, detail="Invalid provider. Use 'auto', 'bart', 'groq', 'ollama', or 'llama'")
        
        return SummarizerResponse(
            original_length=len(request.text),
            summary=summary,
            summary_length=len(summary),
            provider=used_provider,
            summary_type=summary_type,
            success=True
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating summary: {str(e)}")

# Summarizer endpoint - DOC/DOCX file upload
@app.post("/summarize-doc", response_model=DOCSummarizerResponse)
async def summarize_doc(
    file: UploadFile = File(...),
    summary_type: str = Form("comprehensive"),
    save_summary: bool = Form(False),
    summary_title: Optional[str] = Form(None)
):
    """
    Comprehensive DOC/DOCX summarizer that processes ALL sections individually and creates an overall summary.
    This endpoint is specifically designed for multi-section documents and ensures all points are covered.
    
    Features:
    - Processes each section individually
    - Creates section-by-section summaries
    - Generates comprehensive overall summary covering all sections
    - Uses BART model for contextual, abstractive summarization
    
    Summary types:
    - "concise": Brief summaries (2-3 sentences per section)
    - "detailed": Comprehensive summaries with all key points
    - "comprehensive": Most detailed summaries covering all aspects (default)
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="File is required")
    
    filename_lower = file.filename.lower()
    if not (filename_lower.endswith('.docx') or filename_lower.endswith('.doc')):
        raise HTTPException(status_code=400, detail="Only DOC/DOCX files are supported for this endpoint")
    
    if filename_lower.endswith('.doc'):
        raise HTTPException(status_code=400, detail="Old .doc format not supported. Please convert to .docx")
    
    if not DOCX_SUPPORT:
        raise HTTPException(status_code=400, detail="DOCX support not available. Install: pip install python-docx")
    
    if not DOC_SUMMARIZER_SUPPORT:
        raise HTTPException(
            status_code=400, 
            detail="DOC Summarizer not available. Install: pip install transformers torch sentencepiece"
        )
    
    summary_type = summary_type.lower()
    if summary_type not in ["concise", "detailed", "comprehensive"]:
        summary_type = "comprehensive"
    
    try:
        # Extract sections from DOCX
        sections = await extract_sections_from_docx(file)
        
        if not sections or all(not s or len(s.strip()) < 50 for s in sections):
            raise HTTPException(
                status_code=400, 
                detail="Could not extract sufficient text from DOCX. File may be empty or corrupted."
            )
        
        print(f"[DOC Summarizer] Processing {len(sections)} sections with {summary_type} summarization...")
        
        # Use comprehensive DOC summarizer
        doc_summarizer = DOCSummarizer()
        result = doc_summarizer.summarize_all_sections(sections, summary_type=summary_type)
        
        # Convert to response format
        section_summaries = [
            DOCSectionSummary(
                section_number=ss["section_number"],
                summary=ss["summary"],
                original_length=ss["original_length"],
                summary_length=ss["summary_length"]
            )
            for ss in result["section_summaries"]
        ]
        
        response = DOCSummarizerResponse(
            total_sections=result["total_sections"],
            sections_summarized=result["sections_summarized"],
            section_summaries=section_summaries,
            overall_summary=result["overall_summary"],
            summary_type=summary_type,
            provider="doc_summarizer",
            success=True
        )
        
        # Optionally save summary for later review
        if save_summary:
            title = summary_title or file.filename or f"DOC Summary - {datetime.now().strftime('%Y-%m-%d')}"
            save_summary_to_store(
                title=title,
                content=result["overall_summary"],
                source=file.filename,
                source_type="docx",
                metadata={"summary_type": summary_type, "total_sections": result["total_sections"]}
            )
            print(f"[DOC Summarizer] Summary saved for later review")
        
        return response
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing DOCX: {str(e)}")

# Dedicated PDF Summarizer endpoint - Comprehensive multi-page summarization
@app.post("/summarize-pdf", response_model=PDFSummarizerResponse)
async def summarize_pdf(
    file: UploadFile = File(...),
    summary_type: str = Form("comprehensive"),
    save_summary: bool = Form(False),
    summary_title: Optional[str] = Form(None)
):
    """
    Comprehensive PDF summarizer that processes ALL pages individually and creates an overall summary.
    This endpoint is specifically designed for multi-page PDFs and ensures all points are covered.
    
    Features:
    - Processes each page individually
    - Creates page-by-page summaries
    - Generates comprehensive overall summary covering all pages
    - Uses BART model for contextual, abstractive summarization
    - Supports OCR for scanned PDFs
    
    Summary types:
    - "concise": Brief summaries (2-3 sentences per page)
    - "detailed": Comprehensive summaries with all key points
    - "comprehensive": Most detailed summaries covering all aspects (default)
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="File is required")
    
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported for this endpoint")
    
    if not PDF_SUPPORT:
        raise HTTPException(status_code=400, detail="PDF support not available. Install PyPDF2: pip install PyPDF2")
    
    if not PDF_SUMMARIZER_SUPPORT:
        raise HTTPException(
            status_code=400, 
            detail="PDF Summarizer not available. Install: pip install transformers torch sentencepiece"
        )
    
    summary_type = summary_type.lower()
    if summary_type not in ["concise", "detailed", "comprehensive"]:
        summary_type = "comprehensive"
    
    try:
        # Extract pages from PDF
        pages = await extract_pages_from_pdf(file)
        
        if not pages or all(not p or len(p.strip()) < 50 for p in pages):
            raise HTTPException(
                status_code=400, 
                detail="Could not extract sufficient text from PDF. File may be empty, corrupted, or all pages are images."
            )
        
        print(f"[PDF Summarizer] Processing {len(pages)} pages with {summary_type} summarization...")
        
        # Use comprehensive PDF summarizer
        pdf_summarizer = PDFSummarizer()
        result = pdf_summarizer.summarize_all_pages(pages, summary_type=summary_type, improve_grammar=False)
        
        # Improve grammar of all page summaries
        print("[PDF Summarizer] Improving grammar of page summaries...")
        for ps in result["page_summaries"]:
            if ps["summary_length"] > 0:
                try:
                    ps["summary"] = await improve_grammar_and_clarity(ps["summary"])
                except Exception as e:
                    print(f"[PDF Summarizer] Error improving grammar for page {ps['page_number']}: {str(e)}")
        
        # Improve grammar of overall summary
        if result["overall_summary"]:
            try:
                result["overall_summary"] = await improve_grammar_and_clarity(result["overall_summary"])
            except Exception as e:
                print(f"[PDF Summarizer] Error improving overall summary grammar: {str(e)}")
        
        # Convert to response format
        page_summaries = [
            PDFPageSummary(
                page_number=ps["page_number"],
                summary=ps["summary"],
                original_length=ps["original_length"],
                summary_length=ps["summary_length"]
            )
            for ps in result["page_summaries"]
        ]
        
        response = PDFSummarizerResponse(
            total_pages=result["total_pages"],
            pages_summarized=result["pages_summarized"],
            page_summaries=page_summaries,
            overall_summary=result["overall_summary"],
            summary_type=summary_type,
            provider="pdf_summarizer",
            success=True
        )
        
        # Optionally save summary for later review
        if save_summary:
            title = summary_title or file.filename or f"PDF Summary - {datetime.now().strftime('%Y-%m-%d')}"
            save_summary_to_store(
                title=title,
                content=result["overall_summary"],
                source=file.filename,
                source_type="pdf",
                metadata={"summary_type": summary_type, "total_pages": result["total_pages"]}
            )
            print(f"[PDF Summarizer] Summary saved for later review")
        
        return response
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing PDF: {str(e)}")

# ==================== CHATBOT FUNCTIONALITY ====================

# Helper Functions for Chatbot

def get_or_create_conversation(conversation_id: Optional[str] = None) -> str:
    """Get existing conversation or create a new one with the provided ID"""
    if conversation_id:
        if conversation_id in conversation_store:
            return conversation_id
        else:
            # Create new conversation with the provided ID
            conversation_store[conversation_id] = {
                "messages": [],
                "created_at": datetime.now().isoformat(),
                "updated_at": datetime.now().isoformat()
            }
            return conversation_id
    
    # Create new conversation with UUID if no ID provided
    new_id = str(uuid.uuid4())
    conversation_store[new_id] = {
        "messages": [],
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat()
    }
    return new_id

def add_message_to_conversation(conversation_id: str, role: str, content: str):
    """Add a message to conversation history"""
    if conversation_id not in conversation_store:
        get_or_create_conversation(conversation_id)
    
    message = {
        "role": role,
        "content": content,
        "timestamp": datetime.now().isoformat()
    }
    conversation_store[conversation_id]["messages"].append(message)
    conversation_store[conversation_id]["updated_at"] = datetime.now().isoformat()

def get_conversation_history(conversation_id: str, max_messages: int = 10) -> List[Dict]:
    """Get conversation history (limited to max_messages)"""
    if conversation_id not in conversation_store:
        return []
    
    messages = conversation_store[conversation_id]["messages"]
    # Return last max_messages messages
    return messages[-max_messages:] if len(messages) > max_messages else messages

def retrieve_rag_context(query: str, top_k: int = 3) -> str:
    """Retrieve relevant context from RAG knowledge store using simple keyword matching"""
    if not rag_knowledge_store:
        return ""
    
    query_lower = query.lower()
    query_words = set(query_lower.split())
    
    # Simple keyword-based retrieval (can be upgraded to vector similarity)
    scored_items = []
    for item_id, item in rag_knowledge_store.items():
        text_lower = item["text"].lower()
        text_words = set(text_lower.split())
        
        # Calculate simple overlap score
        overlap = len(query_words.intersection(text_words))
        if overlap > 0:
            scored_items.append((overlap, item["text"]))
    
    # Sort by score and return top_k
    scored_items.sort(reverse=True, key=lambda x: x[0])
    context_parts = [text for _, text in scored_items[:top_k]]
    
    if context_parts:
        return "\n\n".join([f"Context: {text}" for text in context_parts])
    return ""

def add_to_rag_store(text: str, metadata: Optional[Dict] = None):
    """Add text to RAG knowledge store"""
    text_hash = hashlib.md5(text.encode()).hexdigest()
    rag_knowledge_store[text_hash] = {
        "text": text,
        "metadata": metadata or {},
        "added_at": datetime.now().isoformat()
    }

async def chat_with_groq(messages: List[Dict], use_rag: bool = False, rag_context: str = "") -> str:
    """Chat with Groq API"""
    if not GROQ_API_KEY:
        raise HTTPException(status_code=500, detail="Groq API key not configured")
    
    # Prepare messages
    system_message = {
        "role": "system",
        "content": "You are a helpful, knowledgeable, and friendly AI assistant. You help users learn, solve problems, and answer questions clearly and accurately."
    }
    
    if use_rag and rag_context:
        system_message["content"] += f"\n\nRelevant context from knowledge base:\n{rag_context}\n\nUse this context to provide accurate and helpful responses."
    
    api_messages = [system_message] + messages
    
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": GROQ_MODEL_NAME,
        "messages": api_messages,
        "temperature": 0.7,
        "max_tokens": 2000
    }
    
    try:
        response = requests.post(GROQ_API_ENDPOINT, headers=headers, json=payload, timeout=45)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Groq API error: {str(e)}")

async def chat_with_ollama(messages: List[Dict], use_rag: bool = False, rag_context: str = "") -> str:
    """Chat with Ollama API"""
    # Convert messages to prompt format
    prompt_parts = []
    
    if use_rag and rag_context:
        prompt_parts.append(f"Relevant context:\n{rag_context}\n\n")
    
    for msg in messages:
        role = msg["role"]
        content = msg["content"]
        if role == "user":
            prompt_parts.append(f"User: {content}")
        elif role == "assistant":
            prompt_parts.append(f"Assistant: {content}")
    
    prompt_parts.append("Assistant:")
    prompt = "\n".join(prompt_parts)
    
    url = f"{OLLAMA_BASE_URL}/api/generate"
    payload = {
        "model": OLLAMA_MODEL_PRIMARY,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.7,
            "num_predict": 2000
        }
    }
    
    try:
        response = requests.post(url, json=payload, timeout=90)
        response.raise_for_status()
        data = response.json()
        return data.get("response", "Error generating response from Ollama")
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Ollama service not available. Make sure Ollama is running on localhost:11434")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ollama API error: {str(e)}")

async def chat_with_llama(messages: List[Dict], use_rag: bool = False, rag_context: str = "") -> str:
    """Chat with Local LLaMA API"""
    system_message = {
        "role": "system",
        "content": "You are a helpful, knowledgeable, and friendly AI assistant."
    }
    
    if use_rag and rag_context:
        system_message["content"] += f"\n\nRelevant context:\n{rag_context}\n\nUse this context to provide accurate responses."
    
    api_messages = [system_message] + messages
    
    headers = {"Content-Type": "application/json"}
    payload = {
        "model": "llama3.1",
        "messages": api_messages,
        "temperature": 0.7,
        "max_tokens": 2000
    }
    
    try:
        response = requests.post(LOCAL_LLAMA_API_URL, headers=headers, json=payload, timeout=90)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Local LLaMA service not available. Make sure the local LLaMA server is running")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLaMA API error: {str(e)}")

async def chat_with_fallback(messages: List[Dict], use_rag: bool = False, rag_context: str = "") -> Tuple[str, str]:
    """Chat with fallback logic: Groq -> Ollama -> LLaMA"""
    # Try Groq first
    if GROQ_API_KEY:
        try:
            response = await chat_with_groq(messages, use_rag, rag_context)
            return response, "groq"
        except Exception as e:
            print(f"Groq failed: {str(e)}, trying fallback...")
    
    # Try Ollama
    try:
        response = await chat_with_ollama(messages, use_rag, rag_context)
        return response, "ollama"
    except Exception as e:
        print(f"Ollama failed: {str(e)}, trying fallback...")
    
    # Try LLaMA
    try:
        response = await chat_with_llama(messages, use_rag, rag_context)
        return response, "llama"
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"All chat providers failed. Last error: {str(e)}")

# Chatbot Endpoint
@app.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """
    Chatbot endpoint with conversation memory and RAG support.
    
    Features:
    - Maintains conversation history across messages
    - Supports multiple AI providers (Groq, Ollama, LLaMA) with auto-fallback
    - RAG (Retrieval Augmented Generation) for context-aware responses
    - Remembers previous conversations using conversation_id
    
    Usage:
    - First message: Don't provide conversation_id (creates new conversation)
    - Subsequent messages: Use the conversation_id from previous response
    - Set use_rag=True to enable context retrieval from knowledge base
    """
    if not request.message or len(request.message.strip()) < 1:
        raise HTTPException(status_code=400, detail="Message is required")
    
    try:
        # Get or create conversation
        conversation_id = get_or_create_conversation(request.conversation_id)
        
        # Get conversation history
        history = get_conversation_history(conversation_id, request.max_history)
        
        # Convert history to API format
        api_messages = []
        for msg in history:
            api_messages.append({
                "role": msg["role"],
                "content": msg["content"]
            })
        
        # Add current user message
        api_messages.append({
            "role": "user",
            "content": request.message
        })
        
        # Retrieve RAG context if enabled
        rag_context = ""
        if request.use_rag:
            rag_context = retrieve_rag_context(request.message)
            print(f"[Chat] RAG context retrieved: {len(rag_context)} characters")
        
        # Store user message
        add_message_to_conversation(conversation_id, "user", request.message)
        
        # Get response from AI provider
        provider = request.provider.lower() if request.provider else "auto"
        
        if provider == "groq":
            response_text = await chat_with_groq(api_messages, request.use_rag, rag_context)
            used_provider = "groq"
        elif provider == "ollama":
            response_text = await chat_with_ollama(api_messages, request.use_rag, rag_context)
            used_provider = "ollama"
        elif provider == "llama":
            response_text = await chat_with_llama(api_messages, request.use_rag, rag_context)
            used_provider = "llama"
        elif provider == "auto":
            response_text, used_provider = await chat_with_fallback(api_messages, request.use_rag, rag_context)
        else:
            raise HTTPException(status_code=400, detail="Invalid provider. Use 'auto', 'groq', 'ollama', or 'llama'")
        
        # Store assistant response
        add_message_to_conversation(conversation_id, "assistant", response_text)
        
        # Add conversation to RAG store (for future context retrieval)
        if request.use_rag:
            conversation_text = f"User: {request.message}\nAssistant: {response_text}"
            add_to_rag_store(conversation_text, {"conversation_id": conversation_id, "type": "conversation"})
        
        return ChatResponse(
            response=response_text,
            conversation_id=conversation_id,
            provider=used_provider,
            message_count=len(conversation_store[conversation_id]["messages"]),
            timestamp=datetime.now().isoformat(),
            success=True
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error in chat: {str(e)}")

# Get Conversation History Endpoint
@app.get("/chat/history/{conversation_id}", response_model=ConversationHistory)
async def get_conversation_history_endpoint(conversation_id: str):
    """Get full conversation history for a given conversation_id"""
    if conversation_id not in conversation_store:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    conv_data = conversation_store[conversation_id]
    messages = [ChatMessage(**msg) for msg in conv_data["messages"]]
    
    return ConversationHistory(
        conversation_id=conversation_id,
        messages=messages,
        created_at=conv_data["created_at"],
        updated_at=conv_data["updated_at"],
        message_count=len(messages)
    )

# Delete Conversation Endpoint
@app.delete("/chat/history/{conversation_id}")
async def delete_conversation(conversation_id: str):
    """Delete a conversation and its history"""
    if conversation_id not in conversation_store:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    del conversation_store[conversation_id]
    return {"message": "Conversation deleted successfully", "conversation_id": conversation_id}

# RAG Knowledge Store Endpoint (renamed from /chat/knowledge)
@app.post("/rag/knowledge")
async def add_knowledge(text: str, metadata: Optional[Dict] = None):
    """Add knowledge to RAG store for context retrieval in chat"""
    if not text or len(text.strip()) < 10:
        raise HTTPException(status_code=400, detail="Text must be at least 10 characters")
    
    add_to_rag_store(text, metadata)
    return {"message": "Knowledge added to RAG store successfully", "text_length": len(text)}

# ==================== SAVED SUMMARIZER WITH QUESTION GENERATION ====================

# Helper Functions for Saved Summarizer

def save_summary_to_store(title: str, content: str, source: Optional[str] = None, 
                 source_type: Optional[str] = None, metadata: Optional[Dict] = None) -> str:
    """Save a summary to the store"""
    summary_id = str(uuid.uuid4())
    saved_summaries_store[summary_id] = {
        "title": title,
        "content": content,
        "source": source,
        "source_type": source_type,
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat(),
        "question_count": 0,
        "metadata": metadata or {}
    }
    return summary_id

def get_summary(summary_id: str) -> Optional[Dict]:
    """Get a saved summary"""
    return saved_summaries_store.get(summary_id)

def create_review_schedule() -> Dict:
    """Create spaced repetition schedule: same day, next day, 1 week later"""
    now = datetime.now()
    return {
        "first_review": now.isoformat(),  # Same day
        "second_review": (now + timedelta(days=1)).isoformat(),  # Next day
        "third_review": (now + timedelta(days=7)).isoformat(),  # 1 week later
        "next_review_date": now.isoformat(),  # Start with first review
        "last_review_date": None,
        "review_count": 0,
        "mastery_level": "learning"
    }

def update_review_schedule(question_id: str, is_correct: bool, confidence: str = "medium"):
    """Update review schedule based on performance"""
    if question_id not in questions_store:
        return
    
    question_data = questions_store[question_id]
    schedule = question_data.get("review_schedule", create_review_schedule())
    
    now = datetime.now()
    schedule["last_review_date"] = now.isoformat()
    schedule["review_count"] += 1
    
    # Calculate next review based on performance and current review count
    if schedule["review_count"] == 1:
        # After first review, schedule next day
        schedule["next_review_date"] = (now + timedelta(days=1)).isoformat()
    elif schedule["review_count"] == 2:
        # After second review, schedule 1 week later
        schedule["next_review_date"] = (now + timedelta(days=7)).isoformat()
    else:
        # After third review, adjust based on performance
        if is_correct and confidence in ["high", "medium"]:
            # If correct, increase interval
            if schedule["mastery_level"] == "learning":
                schedule["mastery_level"] = "reviewing"
                schedule["next_review_date"] = (now + timedelta(days=14)).isoformat()
            elif schedule["mastery_level"] == "reviewing":
                schedule["mastery_level"] = "mastered"
                schedule["next_review_date"] = (now + timedelta(days=30)).isoformat()
        else:
            # If incorrect, reset to learning
            schedule["mastery_level"] = "learning"
            schedule["next_review_date"] = (now + timedelta(days=1)).isoformat()
    
    question_data["review_schedule"] = schedule
    questions_store[question_id] = question_data

async def generate_questions_with_ai(summary_content: str, question_types: List[str], 
                                     num_questions: int, difficulty: str) -> List[Dict]:
    """Generate questions using AI (Groq, Ollama, or LLaMA)"""
    questions = []
    
    # Create prompt for question generation
    question_type_names = {
        "qa": "Question-Answer pairs",
        "mcq": "Multiple Choice Questions (MCQs)",
        "flashcard": "Flashcards",
        "case_based": "Case-based questions"
    }
    
    types_str = ", ".join([question_type_names.get(qt, qt) for qt in question_types])
    
    prompt = f"""Generate {num_questions} {types_str} based on the following summary content.

Summary:
{summary_content}

Requirements:
- Difficulty level: {difficulty}
- Questions should test understanding of key concepts
- Answers should be clear and accurate
- For MCQs: Provide 4 options with only one correct answer
- For case-based: Create realistic scenarios related to the content
- For flashcards: Create concise Q&A pairs

Format your response as JSON with this structure:
{{
  "questions": [
    {{
      "type": "qa|mcq|flashcard|case_based",
      "question": "...",
      "answer": "...",
      "options": [{{"option": "...", "is_correct": true/false}}, ...],  // Only for MCQs
      "explanation": "..."  // Optional
    }}
  ]
}}

Generate the questions now:"""
    
    # Try to generate with AI
    messages = [{"role": "user", "content": prompt}]
    
    try:
        # Try Groq first
        if GROQ_API_KEY:
            response_text = await chat_with_groq(messages, use_rag=False, rag_context="")
            # Parse response (may need JSON extraction)
            # For now, we'll create questions from the response
            questions = parse_ai_questions(response_text, question_types, num_questions, difficulty)
            if questions:
                return questions
    except Exception as e:
        print(f"Groq question generation failed: {str(e)}")
    
    # Fallback: Generate simple questions from summary
    return generate_simple_questions(summary_content, question_types, num_questions, difficulty)

def parse_ai_questions(response_text: str, question_types: List[str], 
                      num_questions: int, difficulty: str) -> List[Dict]:
    """Parse AI-generated questions from response text"""
    # Try to extract JSON from response
    import re
    json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
    if json_match:
        try:
            import json
            data = json.loads(json_match.group())
            return data.get("questions", [])
        except:
            pass
    
    # If JSON parsing fails, create questions from text
    return generate_simple_questions(response_text, question_types, num_questions, difficulty)

def generate_simple_questions(summary_content: str, question_types: List[str], 
                              num_questions: int, difficulty: str) -> List[Dict]:
    """Generate simple questions from summary content (fallback)"""
    questions = []
    sentences = summary_content.split('. ')
    key_sentences = [s.strip() for s in sentences if len(s.strip()) > 20][:num_questions * 2]
    
    for i, qtype in enumerate(question_types):
        for j in range(num_questions):
            if i * num_questions + j >= len(key_sentences):
                break
            
            sentence = key_sentences[i * num_questions + j]
            question_id = str(uuid.uuid4())
            
            if qtype == "qa":
                questions.append({
                    "question_id": question_id,
                    "question_type": "qa",
                    "question": f"Explain: {sentence[:100]}...",
                    "answer": sentence,
                    "explanation": None,
                    "difficulty": difficulty
                })
            elif qtype == "mcq":
                # Create simple MCQ
                questions.append({
                    "question_id": question_id,
                    "question_type": "mcq",
                    "question": f"Which statement is true about: {sentence[:80]}...?",
                    "answer": sentence[:100],
                    "options": [
                        {"option": sentence[:100], "is_correct": True},
                        {"option": "Option 2 (incorrect)", "is_correct": False},
                        {"option": "Option 3 (incorrect)", "is_correct": False},
                        {"option": "Option 4 (incorrect)", "is_correct": False}
                    ],
                    "explanation": None,
                    "difficulty": difficulty
                })
            elif qtype == "flashcard":
                questions.append({
                    "question_id": question_id,
                    "question_type": "flashcard",
                    "question": f"What is: {sentence[:80]}...?",
                    "answer": sentence,
                    "explanation": None,
                    "difficulty": difficulty
                })
            elif qtype == "case_based":
                questions.append({
                    "question_id": question_id,
                    "question_type": "case_based",
                    "question": f"Case: A scenario related to {sentence[:60]}... How would you apply this concept?",
                    "answer": f"Based on the concept: {sentence}",
                    "explanation": None,
                    "difficulty": difficulty
                })
    
    return questions[:num_questions * len(question_types)]

# Saved Summarizer Endpoints

@app.post("/summaries/save", response_model=SavedSummary)
async def save_summary_endpoint(request: SaveSummaryRequest):
    """
    Save a summary for later review and question generation.
    Can save summaries from PDF, DOC, or text sources.
    """
    if not request.title or not request.content:
        raise HTTPException(status_code=400, detail="Title and content are required")
    
    if len(request.content.strip()) < 50:
        raise HTTPException(status_code=400, detail="Content must be at least 50 characters")
    
    summary_id = save_summary_to_store(
        title=request.title,
        content=request.content,
        source=request.source,
        source_type=request.source_type,
        metadata=request.metadata
    )
    
    summary_data = saved_summaries_store[summary_id]
    return SavedSummary(
        summary_id=summary_id,
        title=summary_data["title"],
        content=summary_data["content"],
        source=summary_data["source"],
        source_type=summary_data["source_type"],
        created_at=summary_data["created_at"],
        updated_at=summary_data["updated_at"],
        question_count=summary_data["question_count"],
        metadata=summary_data["metadata"]
    )

@app.get("/summaries", response_model=List[SavedSummary])
async def list_summaries():
    """Get all saved summaries"""
    summaries = []
    for summary_id, data in saved_summaries_store.items():
        summaries.append(SavedSummary(
            summary_id=summary_id,
            title=data["title"],
            content=data["content"][:200] + "..." if len(data["content"]) > 200 else data["content"],
            source=data["source"],
            source_type=data["source_type"],
            created_at=data["created_at"],
            updated_at=data["updated_at"],
            question_count=data["question_count"],
            metadata=data["metadata"]
        ))
    return summaries

@app.get("/summaries/{summary_id}", response_model=SavedSummary)
async def get_summary_endpoint(summary_id: str):
    """Get a specific saved summary"""
    summary_data = get_summary(summary_id)
    if not summary_data:
        raise HTTPException(status_code=404, detail="Summary not found")
    
    return SavedSummary(
        summary_id=summary_id,
        title=summary_data["title"],
        content=summary_data["content"],
        source=summary_data["source"],
        source_type=summary_data["source_type"],
        created_at=summary_data["created_at"],
        updated_at=summary_data["updated_at"],
        question_count=summary_data["question_count"],
        metadata=summary_data["metadata"]
    )

@app.post("/summaries/{summary_id}/questions/generate", response_model=QuestionGenerationResponse)
async def generate_questions_endpoint(summary_id: str, request: QuestionGenerationRequest):
    """
    Generate practice questions from a saved summary.
    Supports: Q&A, MCQs, Flashcards, Case-based questions.
    Questions are automatically scheduled for spaced repetition review.
    """
    summary_data = get_summary(summary_id)
    if not summary_data:
        raise HTTPException(status_code=404, detail="Summary not found")
    
    # Use request parameters or defaults
    question_types = request.question_types or ["qa", "mcq", "flashcard", "case_based"]
    num_questions = request.num_questions or 5
    difficulty = request.difficulty or "medium"
    
    # Generate questions
    generated_questions = await generate_questions_with_ai(
        summary_data["content"],
        question_types,
        num_questions,
        difficulty
    )
    
    # Save questions with review schedule
    saved_questions = []
    for q in generated_questions:
        question_id = q.get("question_id", str(uuid.uuid4()))
        review_schedule = create_review_schedule()
        
        question_data = {
            "summary_id": summary_id,
            "question_type": q["question_type"],
            "question": q["question"],
            "answer": q["answer"],
            "options": q.get("options"),
            "explanation": q.get("explanation"),
            "difficulty": q.get("difficulty", difficulty),
            "created_at": datetime.now().isoformat(),
            "review_schedule": review_schedule
        }
        
        questions_store[question_id] = question_data
        saved_questions.append(Question(
            question_id=question_id,
            summary_id=summary_id,
            question_type=q["question_type"],
            question=q["question"],
            answer=q["answer"],
            options=[MCQOption(**opt) for opt in q.get("options", [])] if q.get("options") else None,
            explanation=q.get("explanation"),
            difficulty=q.get("difficulty", difficulty),
            created_at=question_data["created_at"]
        ))
    
    # Update summary question count
    saved_summaries_store[summary_id]["question_count"] = len(saved_questions)
    saved_summaries_store[summary_id]["updated_at"] = datetime.now().isoformat()
    
    return QuestionGenerationResponse(
        summary_id=summary_id,
        questions=saved_questions,
        total_questions=len(saved_questions),
        success=True
    )

@app.get("/summaries/{summary_id}/questions", response_model=List[Question])
async def get_summary_questions(summary_id: str):
    """Get all questions for a specific summary"""
    if summary_id not in saved_summaries_store:
        raise HTTPException(status_code=404, detail="Summary not found")
    
    questions = []
    for question_id, question_data in questions_store.items():
        if question_data["summary_id"] == summary_id:
            questions.append(Question(
                question_id=question_id,
                summary_id=summary_id,
                question_type=question_data["question_type"],
                question=question_data["question"],
                answer=question_data["answer"],
                options=[MCQOption(**opt) for opt in question_data.get("options", [])] if question_data.get("options") else None,
                explanation=question_data.get("explanation"),
                difficulty=question_data.get("difficulty", "medium"),
                created_at=question_data["created_at"]
            ))
    
    return questions

@app.get("/reviews/pending", response_model=List[PendingReview])
async def get_pending_reviews():
    """Get all questions that are due for review (spaced repetition)"""
    now = datetime.now()
    pending = []
    
    for question_id, question_data in questions_store.items():
        schedule = question_data.get("review_schedule", {})
        next_review_str = schedule.get("next_review_date")
        
        if next_review_str:
            next_review = datetime.fromisoformat(next_review_str.replace('Z', '+00:00'))
            # Handle timezone-aware datetime
            if next_review.tzinfo:
                now_aware = now.replace(tzinfo=next_review.tzinfo)
                if now_aware >= next_review:
                    days_until = (now_aware - next_review).days
                    pending.append(PendingReview(
                        question_id=question_id,
                        summary_id=question_data["summary_id"],
                        question=question_data["question"],
                        question_type=question_data["question_type"],
                        answer=question_data["answer"],
                        next_review_date=next_review_str,
                        days_until_review=days_until
                    ))
            else:
                if now >= next_review:
                    days_until = (now - next_review).days
                    pending.append(PendingReview(
                        question_id=question_id,
                        summary_id=question_data["summary_id"],
                        question=question_data["question"],
                        question_type=question_data["question_type"],
                        answer=question_data["answer"],
                        next_review_date=next_review_str,
                        days_until_review=days_until
                    ))
    
    # Sort by days overdue (most overdue first)
    pending.sort(key=lambda x: x.days_until_review, reverse=True)
    return pending

@app.post("/reviews/submit", response_model=ReviewResponse)
async def submit_review(request: ReviewRequest):
    """
    Submit a review for a question (spaced repetition).
    Updates the review schedule based on performance.
    """
    if request.question_id not in questions_store:
        raise HTTPException(status_code=404, detail="Question not found")
    
    question_data = questions_store[request.question_id]
    update_review_schedule(request.question_id, request.is_correct, request.confidence)
    
    schedule = question_data["review_schedule"]
    
    # Determine message based on performance
    if request.is_correct:
        if schedule["mastery_level"] == "mastered":
            message = "Excellent! You've mastered this concept."
        elif schedule["mastery_level"] == "reviewing":
            message = "Great job! Keep reviewing to maintain mastery."
        else:
            message = "Well done! Continue practicing to reinforce learning."
    else:
        message = "Keep practicing! Review the answer and try again later."
    
    return ReviewResponse(
        question_id=request.question_id,
        correct_answer=question_data["answer"],
        is_correct=request.is_correct,
        next_review_date=schedule["next_review_date"],
        review_count=schedule["review_count"],
        mastery_level=schedule["mastery_level"],
        message=message
    )

@app.get("/reviews/stats")
async def get_review_stats():
    """Get statistics about reviews and learning progress"""
    total_questions = len(questions_store)
    learning = sum(1 for q in questions_store.values() 
                  if q.get("review_schedule", {}).get("mastery_level") == "learning")
    reviewing = sum(1 for q in questions_store.values() 
                   if q.get("review_schedule", {}).get("mastery_level") == "reviewing")
    mastered = sum(1 for q in questions_store.values() 
                  if q.get("review_schedule", {}).get("mastery_level") == "mastered")
    
    now = datetime.now()
    pending_count = 0
    for q in questions_store.values():
        schedule = q.get("review_schedule", {})
        next_review_str = schedule.get("next_review_date")
        if next_review_str:
            try:
                next_review = datetime.fromisoformat(next_review_str.replace('Z', '+00:00'))
                if next_review.tzinfo:
                    if now.replace(tzinfo=next_review.tzinfo) >= next_review:
                        pending_count += 1
                else:
                    if now >= next_review:
                        pending_count += 1
            except:
                pass
    
    return {
        "total_questions": total_questions,
        "learning": learning,
        "reviewing": reviewing,
        "mastered": mastered,
        "pending_reviews": pending_count,
        "total_summaries": len(saved_summaries_store)
    }

# ==================== QUIZ/TEST API ====================

async def generate_quiz_questions(subject: str, topic: str, provider: str = "auto", difficulty: str = "medium") -> List[Dict]:
    """Generate 10 quiz questions using AI models"""
    prompt = f"""Generate exactly 10 multiple-choice quiz questions (MCQs) about {subject} - {topic}.

Requirements:
- Difficulty level: {difficulty}
- Each question must have exactly 4 options (A, B, C, D)
- Only one option should be correct
- Questions should test understanding of key concepts
- Format your response as JSON with this structure:
{{
  "questions": [
    {{
      "question": "Question text here?",
      "options": ["Option A", "Option B", "Option C", "Option D"],
      "correct_answer": "Option A",  // The correct option text
      "correct_index": 0,  // Index of correct option (0-3)
      "explanation": "Brief explanation of why this is correct"
    }}
  ]
}}

Generate 10 questions now:"""
    
    messages = [{"role": "user", "content": prompt}]
    questions = []
    
    # Try to generate with AI
    try:
        if provider == "groq" or (provider == "auto" and GROQ_API_KEY):
            try:
                response_text = await chat_with_groq(messages, use_rag=False, rag_context="")
                questions = parse_quiz_questions(response_text, subject, topic, difficulty)
                if questions and len(questions) >= 10:
                    return questions[:10]
            except Exception as e:
                print(f"Groq quiz generation failed: {str(e)}")
                if provider == "groq":
                    raise
        
        if provider == "ollama" or (provider == "auto"):
            try:
                response_text = await chat_with_ollama(messages, use_rag=False, rag_context="")
                questions = parse_quiz_questions(response_text, subject, topic, difficulty)
                if questions and len(questions) >= 10:
                    return questions[:10]
            except Exception as e:
                print(f"Ollama quiz generation failed: {str(e)}")
                if provider == "ollama":
                    raise
        
        if provider == "llama" or (provider == "auto"):
            try:
                response_text = await chat_with_llama(messages, use_rag=False, rag_context="")
                questions = parse_quiz_questions(response_text, subject, topic, difficulty)
                if questions and len(questions) >= 10:
                    return questions[:10]
            except Exception as e:
                print(f"LLaMA quiz generation failed: {str(e)}")
                if provider == "llama":
                    raise
    except Exception as e:
        print(f"AI quiz generation failed: {str(e)}")
    
    # Fallback: Generate simple questions
    return generate_fallback_quiz_questions(subject, topic, difficulty)

def parse_quiz_questions(response_text: str, subject: str, topic: str, difficulty: str) -> List[Dict]:
    """Parse AI-generated quiz questions from response"""
    import re
    import json
    
    # Try to extract JSON from response
    json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
    if json_match:
        try:
            data = json.loads(json_match.group())
            questions_data = data.get("questions", [])
            
            questions = []
            for i, q_data in enumerate(questions_data[:10], 1):
                question_id = str(uuid.uuid4())
                options = q_data.get("options", [])
                correct_answer = q_data.get("correct_answer", "")
                correct_index = q_data.get("correct_index", 0)
                
                # Ensure we have 4 options
                while len(options) < 4:
                    options.append(f"Option {chr(68 + len(options))}")
                
                questions.append({
                    "question_id": question_id,
                    "question_number": i,
                    "question": q_data.get("question", f"Question {i} about {topic}?"),
                    "options": options[:4],
                    "correct_answer": correct_answer or options[correct_index] if correct_index < len(options) else options[0],
                    "correct_index": correct_index if correct_index < len(options) else 0,
                    "explanation": q_data.get("explanation", ""),
                    "question_type": "mcq"
                })
            
            return questions
        except Exception as e:
            print(f"Error parsing quiz JSON: {str(e)}")
    
    return []

def generate_fallback_quiz_questions(subject: str, topic: str, difficulty: str) -> List[Dict]:
    """Generate fallback quiz questions if AI fails"""
    questions = []
    base_questions = [
        f"What is the main concept of {topic} in {subject}?",
        f"Which of the following best describes {topic}?",
        f"What is a key characteristic of {topic}?",
        f"Which statement about {topic} is correct?",
        f"What is the primary purpose of {topic}?",
        f"Which of the following is related to {topic}?",
        f"What is an important aspect of {topic}?",
        f"Which concept is fundamental to {topic}?",
        f"What distinguishes {topic} from other concepts?",
        f"Which of the following applies to {topic}?"
    ]
    
    for i, q_text in enumerate(base_questions[:10], 1):
        question_id = str(uuid.uuid4())
        questions.append({
            "question_id": question_id,
            "question_number": i,
            "question": q_text,
            "options": [
                f"Option A: Correct answer about {topic}",
                f"Option B: Incorrect answer",
                f"Option C: Another incorrect answer",
                f"Option D: Yet another incorrect answer"
            ],
            "correct_answer": f"Option A: Correct answer about {topic}",
            "correct_index": 0,
            "explanation": f"This is the correct answer because it accurately describes {topic} in {subject}.",
            "question_type": "mcq"
        })
    
    return questions

@app.post("/quiz/generate", response_model=QuizResponse)
async def generate_quiz(request: QuizRequest):
    """
    Generate a quiz with 10 questions based on subject and topic.
    Uses AI models (Groq, Ollama, LLaMA) to create questions.
    """
    if not request.subject or not request.topic:
        raise HTTPException(status_code=400, detail="Subject and topic are required")
    
    try:
        # Generate questions
        provider = request.provider.lower() if request.provider else "auto"
        generated_questions = await generate_quiz_questions(
            request.subject,
            request.topic,
            provider,
            request.difficulty or "medium"
        )
        
        if len(generated_questions) < 10:
            # If we got fewer than 10, fill with fallback
            fallback = generate_fallback_quiz_questions(request.subject, request.topic, request.difficulty or "medium")
            generated_questions.extend(fallback[:10 - len(generated_questions)])
        
        # Create quiz
        quiz_id = str(uuid.uuid4())
        quiz_store[quiz_id] = {
            "subject": request.subject,
            "topic": request.topic,
            "questions": generated_questions,
            "created_at": datetime.now().isoformat(),
            "provider": provider,
            "difficulty": request.difficulty or "medium"
        }
        
        # Format questions for response (without answers)
        quiz_questions = []
        for q in generated_questions:
            quiz_questions.append(QuizQuestion(
                question_id=q["question_id"],
                question_number=q["question_number"],
                question=q["question"],
                options=q["options"],
                question_type=q["question_type"]
            ))
        
        return QuizResponse(
            quiz_id=quiz_id,
            subject=request.subject,
            topic=request.topic,
            questions=quiz_questions,
            total_questions=len(quiz_questions),
            created_at=quiz_store[quiz_id]["created_at"],
            provider=provider
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating quiz: {str(e)}")

@app.post("/quiz/submit", response_model=QuizResult)
async def submit_quiz(submission: QuizSubmission):
    """
    Submit quiz answers and get results.
    Returns detailed results showing correct/incorrect answers.
    """
    if submission.quiz_id not in quiz_store:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    quiz_data = quiz_store[submission.quiz_id]
    questions = quiz_data["questions"]
    
    correct_count = 0
    wrong_count = 0
    results = []
    
    for q in questions:
        question_id = q["question_id"]
        user_answer = submission.answers.get(question_id, "").strip()
        correct_answer = q["correct_answer"]
        correct_index = q.get("correct_index", 0)
        
        # Check if answer is correct (can be option text or index)
        is_correct = False
        if user_answer:
            # Check if user answered with option text
            if user_answer.lower() == correct_answer.lower():
                is_correct = True
            # Check if user answered with option letter (A, B, C, D)
            elif user_answer.upper() in ["A", "B", "C", "D"]:
                option_index = ord(user_answer.upper()) - ord("A")
                if option_index == correct_index:
                    is_correct = True
            # Check if user answered with option number (0, 1, 2, 3)
            elif user_answer.isdigit():
                if int(user_answer) == correct_index:
                    is_correct = True
        
        if is_correct:
            correct_count += 1
        else:
            wrong_count += 1
        
        results.append({
            "question_id": question_id,
            "question_number": q["question_number"],
            "question": q["question"],
            "user_answer": user_answer,
            "correct_answer": correct_answer,
            "correct_index": correct_index,
            "is_correct": is_correct,
            "explanation": q.get("explanation", "")
        })
    
    score_percentage = (correct_count / len(questions)) * 100 if questions else 0
    
    return QuizResult(
        quiz_id=submission.quiz_id,
        total_questions=len(questions),
        correct_answers=correct_count,
        wrong_answers=wrong_count,
        score_percentage=round(score_percentage, 2),
        results=results,
        submitted_at=datetime.now().isoformat()
    )

@app.get("/quiz/{quiz_id}")
async def get_quiz(quiz_id: str):
    """Get quiz details (without answers)"""
    if quiz_id not in quiz_store:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    quiz_data = quiz_store[quiz_id]
    questions = []
    
    for q in quiz_data["questions"]:
        questions.append(QuizQuestion(
            question_id=q["question_id"],
            question_number=q["question_number"],
            question=q["question"],
            options=q["options"],
            question_type=q["question_type"]
        ))
    
    return {
        "quiz_id": quiz_id,
        "subject": quiz_data["subject"],
        "topic": quiz_data["topic"],
        "questions": questions,
        "total_questions": len(questions),
        "created_at": quiz_data["created_at"],
        "difficulty": quiz_data.get("difficulty", "medium")
    }

# ==================== AGENT SIMULATORS ====================

# ==================== EduMentor Learning Assistant ====================

async def get_wikipedia_sources(subject: str, topic: str) -> List[Dict]:
    """Get Wikipedia sources for subject and topic"""
    try:
        import wikipedia
        wikipedia.set_lang("en")
        search_query = f"{subject} {topic}"
        results = wikipedia.search(search_query, results=3)
        
        sources = []
        for title in results[:3]:
            try:
                page = wikipedia.page(title, auto_suggest=False)
                sources.append({
                    "title": page.title,
                    "url": page.url,
                    "summary": page.summary[:200] + "..."
                })
            except:
                continue
        return sources
    except ImportError:
        print("[EduMentor] Wikipedia library not installed. Install with: pip install wikipedia")
        return []
    except Exception as e:
        print(f"[EduMentor] Wikipedia search failed: {str(e)}")
        return []

@app.post("/agent/edumentor/generate", response_model=EduMentorResponse)
async def edumentor_generate_lesson(request: EduMentorRequest):
    """
    EduMentor Learning Assistant - Generate personalized educational content.
    
    Features:
    - AI-powered lesson generation
    - Optional Wikipedia sources
    - Knowledge store integration
    - Orchestration engine support
    """
    if not request.subject or not request.topic:
        raise HTTPException(status_code=400, detail="Subject and topic are required")
    
    try:
        provider = request.provider.lower() if request.provider else "auto"
        
        # Build enhanced prompt
        prompt_parts = [
            f"You are EduMentor, a personal learning assistant specializing in {request.subject}.",
            f"Create a comprehensive, engaging lesson about '{request.topic}' in {request.subject}.",
            "",
            "The lesson should include:",
            "1. **Introduction**: Clear overview and importance",
            "2. **Key Concepts**: Break down main ideas simply",
            "3. **Detailed Explanation**: Thorough explanation with examples",
            "4. **Real-world Applications**: Practical uses",
            "5. **Practice Exercises**: 2-3 exercises to reinforce learning",
            "6. **Summary**: Key takeaways",
            "",
            "Write in a friendly, encouraging tone. Use analogies and examples.",
            f"Subject: {request.subject}",
            f"Topic: {request.topic}"
        ]
        
        # Add knowledge store context if enabled
        knowledge_context = ""
        if request.use_knowledge_store:
            knowledge_context = retrieve_rag_context(f"{request.subject} {request.topic}")
            if knowledge_context:
                prompt_parts.append(f"\n\nAdditional Context from Knowledge Store:\n{knowledge_context}")
        
        prompt = "\n".join(prompt_parts)
        
        # Get Wikipedia sources if requested
        wikipedia_sources = []
        if request.include_wikipedia:
            wikipedia_sources = await get_wikipedia_sources(request.subject, request.topic)
            if wikipedia_sources:
                sources_text = "\n".join([f"- {s['title']}: {s['url']}" for s in wikipedia_sources])
                prompt += f"\n\nWikipedia Sources:\n{sources_text}"
        
        # Generate lesson using AI
        messages = [{"role": "user", "content": prompt}]
        lesson_content = ""
        used_provider = provider
        
        if provider == "groq" or (provider == "auto" and GROQ_API_KEY):
            try:
                lesson_content = await chat_with_groq(messages, use_rag=request.use_knowledge_store, rag_context=knowledge_context)
                used_provider = "groq"
            except Exception as e:
                print(f"Groq failed: {str(e)}")
                if provider == "groq":
                    raise
        
        if not lesson_content and (provider == "ollama" or provider == "auto"):
            try:
                lesson_content = await chat_with_ollama(messages, use_rag=request.use_knowledge_store, rag_context=knowledge_context)
                used_provider = "ollama"
            except Exception as e:
                print(f"Ollama failed: {str(e)}")
                if provider == "ollama":
                    raise
        
        if not lesson_content and (provider == "llama" or provider == "auto"):
            try:
                lesson_content = await chat_with_llama(messages, use_rag=request.use_knowledge_store, rag_context=knowledge_context)
                used_provider = "llama"
            except Exception as e:
                print(f"LLaMA failed: {str(e)}")
                if provider == "llama":
                    raise
        
        if not lesson_content:
            # Fallback: Generate a basic lesson structure if all AI providers fail
            lesson_content = f"""# {request.topic} - {request.subject}

## Introduction
{request.topic} is an important concept in {request.subject}. Understanding this topic will help you build a strong foundation in the subject.

## Key Concepts
- Core principles of {request.topic}
- Fundamental relationships and patterns
- Essential terminology

## Detailed Explanation
{request.topic} involves several key aspects that are crucial to understand. Let's explore each one in detail.

## Real-world Applications
This concept has practical applications in various fields and everyday situations.

## Practice Exercises
1. Explain {request.topic} in your own words
2. Identify examples of {request.topic} in real life
3. Solve a problem related to {request.topic}

## Summary
{request.topic} is a fundamental concept in {request.subject} that requires careful study and practice.

**Note:** This is a basic lesson structure. For a complete AI-generated lesson, please ensure your AI provider (Groq, Ollama, or LLaMA) is properly configured."""
            used_provider = "fallback"
        
        # Use orchestration if enabled (enhance content)
        if request.use_orchestration and lesson_content:
            orchestration_prompt = f"""Review and enhance this educational content to make it more engaging and effective:

{lesson_content}

Enhancements needed:
- Improve clarity and flow
- Add more examples if needed
- Ensure all sections are well-structured
- Make it more interactive and engaging

Return the enhanced lesson:"""
            
            orchestration_messages = [{"role": "user", "content": orchestration_prompt}]
            try:
                if GROQ_API_KEY:
                    lesson_content = await chat_with_groq(orchestration_messages, use_rag=False, rag_context="")
                elif OLLAMA_BASE_URL:
                    lesson_content = await chat_with_ollama(orchestration_messages, use_rag=False, rag_context="")
            except:
                pass  # Keep original if orchestration fails
        
        return EduMentorResponse(
            subject=request.subject,
            topic=request.topic,
            lesson_content=lesson_content,
            wikipedia_sources=wikipedia_sources if request.include_wikipedia else None,
            knowledge_store_used=request.use_knowledge_store,
            orchestration_used=request.use_orchestration,
            provider=used_provider,
            created_at=datetime.now().isoformat(),
            success=True
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating lesson: {str(e)}")

# ==================== Financial Agent Simulator ====================

@app.post("/agent/financial/profile", response_model=Dict)
async def create_financial_profile(request: FinancialProfileRequest):
    """
    Create a financial profile for simulation.
    """
    if not request.name or request.monthly_income <= 0:
        raise HTTPException(status_code=400, detail="Name and valid monthly income are required")
    
    profile_id = str(uuid.uuid4())
    total_expenses = sum(exp.amount for exp in request.expense_categories)
    
    financial_profiles_store[profile_id] = {
        "name": request.name,
        "monthly_income": request.monthly_income,
        "monthly_expenses": total_expenses,
        "expense_categories": [{"name": exp.name, "amount": exp.amount} for exp in request.expense_categories],
        "financial_goal": request.financial_goal,
        "financial_type": request.financial_type,
        "risk_level": request.risk_level,
        "created_at": datetime.now().isoformat(),
        "current_balance": 0.0
    }
    
    return {
        "profile_id": profile_id,
        "message": "Financial profile created successfully",
        "monthly_savings": request.monthly_income - total_expenses,
        "created_at": financial_profiles_store[profile_id]["created_at"]
    }

@app.post("/agent/financial/simulate", response_model=FinancialSimulationResponse)
async def simulate_financial_life(request: FinancialSimulationRequest):
    """
    Simulate months of financial life and provide guidance.
    """
    if request.profile_id not in financial_profiles_store:
        raise HTTPException(status_code=404, detail="Financial profile not found")
    
    profile = financial_profiles_store[request.profile_id]
    months = request.months_to_simulate or 12
    
    # Calculate monthly savings
    monthly_savings = profile["monthly_income"] - profile["monthly_expenses"]
    initial_balance = profile.get("current_balance", 0.0)
    
    # Simulate months
    monthly_breakdown = []
    current_balance = initial_balance
    
    for month in range(1, months + 1):
        current_balance += monthly_savings
        monthly_breakdown.append({
            "month": month,
            "income": profile["monthly_income"],
            "expenses": profile["monthly_expenses"],
            "savings": monthly_savings,
            "balance": round(current_balance, 2)
        })
    
    final_balance = current_balance
    total_savings = final_balance - initial_balance
    
    # Generate AI recommendations
    recommendations_prompt = f"""You are a financial advisor. Provide personalized financial guidance based on:

Profile:
- Name: {profile['name']}
- Monthly Income: ₹{profile['monthly_income']}
- Monthly Expenses: ₹{profile['monthly_expenses']}
- Monthly Savings: ₹{monthly_savings}
- Financial Goal: {profile['financial_goal']}
- Financial Type: {profile['financial_type']}
- Risk Level: {profile['risk_level']}
- Projected Savings (after {months} months): ₹{total_savings}

Provide 5-7 actionable recommendations to help achieve the financial goal. Be specific and practical."""

    recommendations = []
    try:
        messages = [{"role": "user", "content": recommendations_prompt}]
        if GROQ_API_KEY:
            rec_text = await chat_with_groq(messages, use_rag=False, rag_context="")
            recommendations = [line.strip() for line in rec_text.split('\n') if line.strip() and line.strip()[0] in ['-', '•', '1', '2', '3', '4', '5']][:7]
        elif OLLAMA_BASE_URL:
            rec_text = await chat_with_ollama(messages, use_rag=False, rag_context="")
            recommendations = [line.strip() for line in rec_text.split('\n') if line.strip() and line.strip()[0] in ['-', '•', '1', '2', '3', '4', '5']][:7]
    except:
        recommendations = [
            f"Save ₹{monthly_savings} per month consistently",
            f"Review and reduce expenses in high-cost categories",
            f"Consider investing based on your {profile['risk_level']} risk tolerance",
            f"Track progress toward goal: {profile['financial_goal']}",
            f"Build emergency fund of 3-6 months expenses"
        ]
    
    # Calculate goal progress
    goal_amount = 0
    try:
        # Try to extract amount from goal text
        import re
        goal_match = re.search(r'₹?\s*(\d+(?:,\d+)*(?:\.\d+)?)', profile['financial_goal'])
        if goal_match:
            goal_amount = float(goal_match.group(1).replace(',', ''))
    except:
        pass
    
    goal_progress = {
        "goal": profile['financial_goal'],
        "current_savings": round(total_savings, 2),
        "target_amount": goal_amount,
        "progress_percentage": round((total_savings / goal_amount * 100) if goal_amount > 0 else 0, 2),
        "months_remaining": round((goal_amount - total_savings) / monthly_savings) if monthly_savings > 0 and goal_amount > 0 else 0
    }
    
    # Update profile balance
    financial_profiles_store[request.profile_id]["current_balance"] = final_balance
    
    return FinancialSimulationResponse(
        profile_id=request.profile_id,
        simulation_period=months,
        initial_balance=round(initial_balance, 2),
        final_balance=round(final_balance, 2),
        total_savings=round(total_savings, 2),
        monthly_breakdown=monthly_breakdown,
        recommendations=recommendations,
        goal_progress=goal_progress,
        provider="financial_simulator",
        created_at=datetime.now().isoformat()
    )

@app.get("/agent/financial/profile/{profile_id}")
async def get_financial_profile(profile_id: str):
    """Get financial profile details"""
    if profile_id not in financial_profiles_store:
        raise HTTPException(status_code=404, detail="Financial profile not found")
    
    return financial_profiles_store[profile_id]

# ==================== Wellness Bot ====================

@app.post("/agent/wellness/profile", response_model=Dict)
async def create_wellness_profile(request: WellnessProfileRequest):
    """
    Create a wellness profile for tracking emotional and financial wellness.
    """
    profile_id = str(uuid.uuid4())
    
    wellness_profiles_store[profile_id] = {
        "emotional_wellness_score": request.emotional_wellness_score or 5,
        "financial_wellness_score": request.financial_wellness_score or 5,
        "current_mood_score": request.current_mood_score or 5,
        "stress_level": request.stress_level or 3,
        "concerns": request.concerns,
        "created_at": datetime.now().isoformat(),
        "sessions": []
    }
    
    return {
        "profile_id": profile_id,
        "message": "Wellness profile created successfully",
        "created_at": wellness_profiles_store[profile_id]["created_at"]
    }

@app.post("/agent/wellness/session", response_model=WellnessSessionResponse)
async def wellness_session(request: WellnessSessionRequest):
    """
    Wellness Bot - Chat session for emotional and financial wellness guidance.
    """
    if request.profile_id not in wellness_profiles_store:
        raise HTTPException(status_code=404, detail="Wellness profile not found")
    
    if not request.message or len(request.message.strip()) < 5:
        raise HTTPException(status_code=400, detail="Message is required (minimum 5 characters)")
    
    profile = wellness_profiles_store[request.profile_id]
    session_id = str(uuid.uuid4())
    
    # Build wellness-aware prompt
    wellness_context = f"""Wellness Profile:
- Emotional Wellness: {profile['emotional_wellness_score']}/10
- Financial Wellness: {profile['financial_wellness_score']}/10
- Current Mood: {profile['current_mood_score']}/10
- Stress Level: {profile['stress_level']}/10
- Previous Concerns: {profile.get('concerns', 'None')}

User Message: {request.message}

You are a compassionate wellness assistant. Provide empathetic, supportive guidance that addresses:
1. Emotional support and validation
2. Practical wellness strategies
3. Stress management techniques
4. Financial wellness if relevant
5. Actionable steps for improvement

Be warm, understanding, and encouraging. Provide specific, helpful advice."""

    # Generate response
    messages = [{"role": "user", "content": wellness_context}]
    response_text = ""
    used_provider = "auto"
    
    try:
        if GROQ_API_KEY:
            response_text = await chat_with_groq(messages, use_rag=False, rag_context="")
            used_provider = "groq"
        elif OLLAMA_BASE_URL:
            response_text = await chat_with_ollama(messages, use_rag=False, rag_context="")
            used_provider = "ollama"
        elif LOCAL_LLAMA_API_URL:
            response_text = await chat_with_llama(messages, use_rag=False, rag_context="")
            used_provider = "llama"
    except Exception as e:
        response_text = f"I understand you're feeling this way. Let's work through this together. Your wellness scores suggest we should focus on {request.session_type} wellness. Would you like to talk more about what's on your mind?"
        print(f"Wellness AI failed: {str(e)}")
    
    # Analyze mood and stress from response
    mood_analysis = {
        "detected_mood": "neutral",
        "mood_score_estimate": profile['current_mood_score'],
        "sentiment": "neutral"
    }
    
    stress_analysis = {
        "current_stress_level": profile['stress_level'],
        "stress_indicators": [],
        "recommendations": []
    }
    
    # Generate recommendations
    recommendations = []
    if profile['stress_level'] >= 7:
        recommendations.append("Consider deep breathing exercises or meditation")
        recommendations.append("Take regular breaks and practice self-care")
    if profile['current_mood_score'] <= 4:
        recommendations.append("Engage in activities you enjoy")
        recommendations.append("Connect with friends or family")
    if profile['financial_wellness_score'] <= 4:
        recommendations.append("Review your financial situation and create a budget")
        recommendations.append("Seek financial counseling if needed")
    
    # Save session
    session_data = {
        "session_id": session_id,
        "message": request.message,
        "response": response_text,
        "session_type": request.session_type,
        "timestamp": datetime.now().isoformat()
    }
    wellness_profiles_store[request.profile_id]["sessions"].append(session_data)
    
    return WellnessSessionResponse(
        profile_id=request.profile_id,
        session_id=session_id,
        response=response_text,
        mood_analysis=mood_analysis,
        stress_analysis=stress_analysis,
        recommendations=recommendations,
        provider=used_provider,
        timestamp=datetime.now().isoformat()
    )

@app.get("/agent/wellness/profile/{profile_id}")
async def get_wellness_profile(profile_id: str):
    """Get wellness profile with session history"""
    if profile_id not in wellness_profiles_store:
        raise HTTPException(status_code=404, detail="Wellness profile not found")
    
    profile = wellness_profiles_store[profile_id].copy()
    return profile

@app.put("/agent/wellness/profile/{profile_id}/update")
async def update_wellness_scores(profile_id: str, 
                                 emotional_score: Optional[int] = None,
                                 financial_score: Optional[int] = None,
                                 mood_score: Optional[int] = None,
                                 stress_level: Optional[int] = None):
    """Update wellness scores"""
    if profile_id not in wellness_profiles_store:
        raise HTTPException(status_code=404, detail="Wellness profile not found")
    
    if emotional_score is not None:
        wellness_profiles_store[profile_id]["emotional_wellness_score"] = max(1, min(10, emotional_score))
    if financial_score is not None:
        wellness_profiles_store[profile_id]["financial_wellness_score"] = max(1, min(10, financial_score))
    if mood_score is not None:
        wellness_profiles_store[profile_id]["current_mood_score"] = max(1, min(10, mood_score))
    if stress_level is not None:
        wellness_profiles_store[profile_id]["stress_level"] = max(1, min(10, stress_level))
    
    return {
        "profile_id": profile_id,
        "message": "Wellness scores updated successfully",
        "updated_scores": wellness_profiles_store[profile_id]
    }

if __name__ == "__main__":
    uvicorn.run("main:app", host=HOST, port=PORT, reload=RELOAD)
